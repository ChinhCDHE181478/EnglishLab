# -*- coding: utf-8 -*-
"""
Word: Bộ TEST CASE Integration Test từng module (111 case)
Dạng viết test case bình thường + tích hợp hàm Controller→Service→Repository
Cho người mới đọc được và mang giải thích với cô.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from full_modules import MODULES  # noqa: E402

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Bo_TEST_CASE_Integration_Test_tung_module_EnglishLab.docx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.1


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    font(run, size=11)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    font(run, size=11)


def field(doc, label, value):
    para = doc.add_paragraph()
    r1 = para.add_run(label + " ")
    font(r1, size=11, bold=True)
    r2 = para.add_run(value)
    font(r2, size=11)
    para.paragraph_format.space_after = Pt(2)


def multiline_field(doc, label, value):
    p(doc, label, bold=True)
    for line in str(value).split("\n"):
        line = line.strip()
        if line:
            bullet(doc, line)


# Mục tiêu tiếng Việt ngắn theo mã (để người mới hiểu ngay)
VN_GOAL = {
    "IT_AUTH_01": "Kiểm tra đăng ký tài khoản mới: Controller → Service → lưu users + OTP.",
    "IT_AUTH_02": "Kiểm tra từ chối đăng ký trùng email (không tạo user thứ 2).",
    "IT_AUTH_03": "Kiểm tra xác thực email bằng OTP: kích hoạt tài khoản.",
    "IT_AUTH_04": "Kiểm tra OTP sai: không kích hoạt tài khoản.",
    "IT_AUTH_05": "Kiểm tra login cấp JWT và /me đọc đúng user qua Security Filter.",
    "IT_AUTH_06": "Kiểm tra login sai mật khẩu: không cấp token.",
    "IT_AUTH_07": "Kiểm tra gọi /me không JWT: Security chặn trước Controller.",
    "IT_AUTH_08": "Kiểm tra quên mật khẩu tạo token reset.",
    "IT_AUTH_09": "Kiểm tra OTP reset sai: không đổi mật khẩu.",
    "IT_AUTH_10": "Kiểm tra resend verification tạo/xoay OTP.",
    "IT_USER_01": "Kiểm tra lấy hồ sơ hiện tại theo JWT.",
    "IT_USER_02": "Kiểm tra cập nhật hồ sơ được lưu DB.",
    "IT_USER_03": "Kiểm tra đổi mật khẩu (sai current bị từ chối; đúng thì đổi hash).",
    "IT_USER_04": "Kiểm tra upload avatar.",
    "IT_USER_05": "Kiểm tra sửa hồ sơ không token bị chặn.",
    "IT_NOTIF_01": "Kiểm tra đọc preference thông báo.",
    "IT_NOTIF_02": "Kiểm tra cập nhật preference thông báo.",
    "IT_NOTIF_03": "Kiểm tra body preference thiếu → validation lỗi.",
    "IT_NOTIF_04": "Kiểm tra danh sách thông báo học viên.",
    "IT_NOTIF_05": "Kiểm tra đếm / đánh dấu đã đọc thông báo.",
    "IT_COMMERCE_01": "Kiểm tra thêm khóa vào giỏ hàng.",
    "IT_COMMERCE_02": "Kiểm tra chuyển wishlist sang giỏ.",
    "IT_COMMERCE_03": "Kiểm tra xóa giỏ hàng.",
    "IT_COMMERCE_04": "Kiểm tra thêm lại khóa vào giỏ sau khi xóa.",
    "IT_PAYMENT_01": "Kiểm tra tạo link thanh toán PayOS.",
    "IT_PAYMENT_02": "Kiểm tra quote giá đơn hàng.",
    "IT_PAYMENT_03": "Kiểm tra webhook thiếu chữ ký bị từ chối.",
    "IT_PAYMENT_04": "Kiểm tra Manager xem danh sách orders.",
    "IT_PAYMENT_05": "Kiểm tra chi tiết / lọc orders (nếu có).",
    "IT_COURSE_01": "Kiểm tra list khóa học public.",
    "IT_COURSE_02": "Kiểm tra chi tiết khóa public.",
    "IT_COURSE_03": "Kiểm tra học viên xem content (cần enroll).",
    "IT_COURSE_04": "Kiểm tra cập nhật progress bài học.",
    "IT_COURSE_05": "Kiểm tra rating khóa học.",
    "IT_COURSE_06": "Kiểm tra từ chối content khi chưa enroll.",
    "IT_DISCUSS_01": "Kiểm tra tạo thảo luận (cần enroll).",
    "IT_DISCUSS_02": "Kiểm tra list thảo luận.",
    "IT_DISCUSS_03": "Kiểm tra report thread.",
    "IT_DISCUSS_04": "Kiểm tra report khi thiếu thread / quyền.",
    "IT_DISCUSS_05": "Kiểm tra CM xem discussion reports.",
    "IT_CONTENT_01": "Kiểm tra CM list khóa online.",
    "IT_CONTENT_02": "Kiểm tra CM tạo/cập nhật khóa (nếu case tạo).",
    "IT_CONTENT_03": "Kiểm tra CM publish/unpublish.",
    "IT_CONTENT_04": "Kiểm tra CM xem chi tiết khóa quản trị.",
    "IT_PACKAGE_01": "Kiểm tra CM list packages.",
    "IT_PACKAGE_02": "Kiểm tra CM chi tiết / tạo package.",
    "IT_PACKAGE_03": "Kiểm tra gắn khóa vào package.",
    "IT_CURRICULUM_01": "Kiểm tra list curriculum programs.",
    "IT_CURRICULUM_02": "Kiểm tra exercise/assessment bank.",
    "IT_CURRICULUM_03": "Kiểm tra learning paths.",
    "IT_CURRICULUM_04": "Kiểm tra rubrics.",
    "IT_CURRICULUM_05": "Kiểm tra chi tiết chương trình / liên kết bank.",
    "IT_ENROLLREQ_01": "Kiểm tra HV tạo enrollment request.",
    "IT_ENROLLREQ_02": "Kiểm tra Staff list enrollment requests.",
    "IT_ENROLLREQ_03": "Kiểm tra Staff xử lý / cập nhật request.",
    "IT_ENROLLREQ_04": "Kiểm tra HV xem request của mình.",
    "IT_ENROLLREQ_05": "Kiểm tra từ chối tạo trùng / validation.",
    "IT_CLASS_01": "Kiểm tra public list classroom offerings.",
    "IT_CLASS_02": "Kiểm tra TM list classrooms.",
    "IT_CLASS_03": "Kiểm tra TM xem chi tiết lớp.",
    "IT_CLASS_04": "Kiểm tra TM xem registrations.",
    "IT_CLASS_05": "Kiểm tra reorder waitlist.",
    "IT_CLASS_06": "Kiểm tra lọc registrations theo status.",
    "IT_CLASS_07": "Kiểm tra xem lớp trước/không phụ thuộc gán GV.",
    "IT_CLASS_08": "Kiểm tra TM list có phân trang/filter.",
    "IT_LEARNERCLS_01": "Kiểm tra HV my-classrooms.",
    "IT_LEARNERCLS_02": "Kiểm tra HV xem sessions.",
    "IT_LEARNERCLS_03": "Kiểm tra HV xem homework.",
    "IT_LEARNERCLS_04": "Kiểm tra HV xem materials.",
    "IT_LEARNERCLS_05": "Kiểm tra HV thao tác homework liên quan.",
    "IT_LEARNERCLS_06": "Kiểm tra HV xem gradebook/me.",
    "IT_TEACH_01": "Kiểm tra GV list lớp assigned.",
    "IT_TEACH_02": "Kiểm tra GV xem homework lớp.",
    "IT_TEACH_03": "Kiểm tra GV xem attendance theo session.",
    "IT_TEACH_04": "Kiểm tra GV xem gradebook lớp.",
    "IT_TEACH_05": "Kiểm tra GV xem change requests của mình.",
    "IT_TEACH_06": "Kiểm tra GV truy cập lớp không được assign (negative nếu có).",
    "IT_QUIZ_01": "Kiểm tra GV list quiz theo lớp.",
    "IT_QUIZ_02": "Kiểm tra GV tạo/xem chi tiết quiz.",
    "IT_QUIZ_03": "Kiểm tra HV list quiz.",
    "IT_QUIZ_04": "Kiểm tra xóa quiz (destructive / N/A trên demo).",
    "IT_ASSESS_01": "Kiểm tra lấy placement test hiện tại.",
    "IT_ASSESS_02": "Kiểm tra submit placement thiếu đáp án.",
    "IT_ASSESS_03": "Kiểm tra list assessments theo khóa.",
    "IT_ASSESS_04": "Kiểm tra list mock tests.",
    "IT_ASSESS_05": "Kiểm tra assessments khi chưa enroll (negative/N/A).",
    "IT_ASSESS_06": "Kiểm tra chi tiết / start mock test.",
    "IT_SUPPORT_01": "Kiểm tra HV tạo support ticket.",
    "IT_SUPPORT_02": "Kiểm tra HV list ticket của mình.",
    "IT_SUPPORT_03": "Kiểm tra Manager list tickets.",
    "IT_SUPPORT_04": "Kiểm tra tạo ticket body rỗng → validation.",
    "IT_ADMIN_01": "Kiểm tra Admin list users.",
    "IT_ADMIN_02": "Kiểm tra Admin xem/lọc user.",
    "IT_ADMIN_03": "Kiểm tra Admin audit logs.",
    "IT_ADMIN_04": "Kiểm tra Admin system config.",
    "IT_LARK_01": "Kiểm tra webhook Lark url_verification.",
    "IT_LARK_02": "Kiểm tra webhook thiếu cấu hình/chữ ký.",
    "IT_LARK_03": "Kiểm tra sync recording session Lark.",
    "IT_INFRA_01": "Kiểm tra TM list campuses.",
    "IT_INFRA_02": "Kiểm tra TM list rooms.",
    "IT_INFRA_03": "Kiểm tra TM list session templates.",
    "IT_REPORT_01": "Kiểm tra TM dashboard.",
    "IT_REPORT_02": "Kiểm tra CM revenue analytics.",
    "IT_PROPOSAL_01": "Kiểm tra Staff list classroom proposals.",
    "IT_PROPOSAL_02": "Kiểm tra tạo/xem chi tiết proposal.",
    "IT_PROPOSAL_03": "Kiểm tra cập nhật trạng thái proposal.",
    "IT_DISPUTE_01": "Kiểm tra HV list attendance disputes.",
    "IT_DISPUTE_02": "Kiểm tra GV list disputes pending.",
    "IT_DISPUTE_03": "Kiểm tra GV xử lý dispute.",
    "IT_NOTES_01": "Kiểm tra HV list learning notes.",
    "IT_NOTES_02": "Kiểm tra tạo/cập nhật note (nếu có).",
}


def vn_goal(case_id: str, fallback_desc: str) -> str:
    if case_id in VN_GOAL:
        return VN_GOAL[case_id]
    # fallback: giữ desc gốc, rút gọn
    return fallback_desc


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "BỘ TEST CASE INTEGRATION TEST\n"
        "TỪNG MODULE – DỰ ÁN ENGLISHLAB (SEP490_G23)\n"
        "Định dạng test case bình thường + tích hợp hàm"
    )
    font(r, size=16, bold=True)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Dùng để học, nộp/giải thích với giảng viên · 24 module · 111 test case")
    font(r, size=11)

    p(doc, "Người lập: phongdx", bold=True)
    p(doc, "Loại kiểm thử: Integration Test (tích hợp Controller → Service → Repository/DB).")
    p(doc, "Công cụ thực thi (không phải test case): Postman hoặc JUnit/MockMvc — chỉ để lấy Actual Result.")

    h(doc, "0. Hướng dẫn đọc file này (cho người mới)", 1)
    numbered(doc, "Mỗi MODULE = 1 sheet Excel (IT - Auth, IT - User, …).")
    numbered(doc, "Mỗi TEST CASE có đủ: Mã, Mục tiêu, Thành phần tích hợp, Tiền điều kiện, Các bước, Kết quả mong đợi.")
    numbered(doc, "Đọc “Các bước” để thấy hàm nào gọi hàm nào — đây là phần cô muốn.")
    numbered(doc, "Khi chạy thật: ghi Actual = Passed / Failed / N/A vào Excel Round.")
    bullet(doc, "Passed: đúng Expected (kể cả negative bị từ chối đúng).")
    bullet(doc, "Failed: sai Expected hoặc lỗi 500.")
    bullet(doc, "N/A: thiếu tiền điều kiện (chưa enroll, thiếu OTP, thiếu data…).")

    p(doc, "Câu nói với cô:", bold=True)
    p(
        doc,
        "Em viết bộ test case Integration theo từng module: mỗi case mô tả tích hợp các hàm "
        "Controller–Service–Repository, kèm precondition và expected. Tool chỉ dùng để thực thi lấy actual.",
    )

    # Mục lục ngắn
    h(doc, "1. Danh mục module", 1)
    rows = []
    total = 0
    for i, m in enumerate(MODULES, 1):
        n = sum(len(g["cases"]) for g in m["groups"])
        total += n
        rows.append([str(i), m["sheet"], m["name"], m["function"], str(n)])
    rows.append(["", "TỔNG", "", "", str(total)])
    # simple table
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Table Grid"
    headers = ["STT", "Sheet Excel", "Module", "Controller chính", "Số case"]
    for i, hd in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(hd)
        font(rr, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(str(val))
            font(rr, size=9)
    doc.add_paragraph()

    # Toàn bộ test case
    h(doc, "2. Chi tiết TEST CASE theo từng MODULE", 1)

    for mi, m in enumerate(MODULES, 1):
        h(doc, f"MODULE {mi}. {m['name']} ({m['sheet']})", 2)
        field(doc, "Yêu cầu tích hợp:", m.get("requirement", ""))
        field(doc, "Thành phần chính:", m.get("components", ""))
        field(doc, "Điểm tích hợp:", m.get("integrations", ""))
        if m.get("srs"):
            field(doc, "Tham chiếu SRS:", m["srs"])

        case_no = 0
        for g in m["groups"]:
            h(doc, f"Nhóm: {g['name']}", 3)
            for c in g["cases"]:
                case_no += 1
                cid = c["id"]
                # Case box header
                para = doc.add_paragraph()
                rr = para.add_run(f"TEST CASE {case_no}: {cid}")
                font(rr, size=12, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

                field(doc, "1. Mục tiêu (tiếng Việt):", vn_goal(cid, c["desc"]))
                field(doc, "2. Mô tả kỹ thuật:", c["desc"])
                field(doc, "3. Module / Sheet:", m["sheet"])
                field(doc, "4. Thành phần tích hợp (components):", m.get("components", ""))
                multiline_field(doc, "5. Tiền điều kiện (Pre-condition):", c["pre"])
                multiline_field(
                    doc,
                    "6. Các bước thực hiện – Procedure (tích hợp hàm):",
                    c["proc"],
                )
                multiline_field(doc, "7. Kết quả mong đợi (Expected):", c["exp"])
                field(
                    doc,
                    "8. Actual Result (điền khi chạy):",
                    "☐ Passed    ☐ Failed    ☐ N/A     Note: ........................",
                )
                p(doc, "—" * 40)

    h(doc, "3. Cách dùng bộ test case này với cô", 1)
    numbered(doc, "Mang file Word này + file Excel cùng mã IT_*.")
    numbered(doc, "Giải thích: đây là test case viết sẵn theo module, mỗi case là một kịch bản tích hợp hàm.")
    numbered(doc, "Chọn 1–2 case Auth demo: đọc mục 6 (Procedure) và 7 (Expected).")
    numbered(doc, "Nói rõ Postman/JUnit chỉ là tool chạy, không thay thế việc viết test case.")
    numbered(doc, "Nếu đã chạy: mở Excel Round đã ghi Passed/Failed/N/A.")

    h(doc, "4. Checklist nộp / thuyết trình", 1)
    bullet(doc, "Biết mở đúng MODULE và đúng mã IT_*.")
    bullet(doc, "Giải thích được 1 case theo Controller → Service → Repository.")
    bullet(doc, "Phân biệt được Passed negative và N/A.")
    bullet(doc, "Không mở đầu bằng “em test Postman”.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    PROJ.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out, PROJ / OUT_NAME)
    print(f"DOCX {out}")
    print(f"cases {total}")
    return out


if __name__ == "__main__":
    build()
