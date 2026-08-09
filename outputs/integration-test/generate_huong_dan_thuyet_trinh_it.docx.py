# -*- coding: utf-8 -*-
"""Tao Word huong dan thuyet trinh Integration Test (bo AUTH) — kich ban noi tung TC."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(
    r"D:\EngLishLab\EnglishLab\outputs\integration-test"
    r"\Huong_dan_THUYET_TRINH_Integration_Test_tung_Test_Case.docx"
)
OUT_DL = Path(
    r"C:\Users\phong\Downloads\intergration test"
    r"\Huong_dan_THUYET_TRINH_Integration_Test_tung_Test_Case.docx"
)


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def add_p(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_say(doc, text):
    """Khoi 'NOI:' — cau noi to."""
    p = doc.add_paragraph()
    tag = p.add_run("NÓI: ")
    set_run_font(tag, bold=True, color=RGBColor(0x0B, 0x5F, 0xA5))
    body = p.add_run(text)
    set_run_font(body)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(12)
    return p


def add_tip(doc, text):
    p = doc.add_paragraph()
    tag = p.add_run("MẸO / NẾU CÔ HỎI: ")
    set_run_font(tag, bold=True, color=RGBColor(0xC0, 0x56, 0x00))
    body = p.add_run(text)
    set_run_font(body)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_font(run)
    return p


# (sheet, feature_vi, open_script, demo_ids, cases[])
# case = id, muc_dich, noi (full script), luong, ket_qua, neu_hoi

SHEETS = [
    {
        "sheet": "IT_COURSE",
        "feature": "Xem khóa học công khai (View public courses)",
        "open": (
            "Em xin trình bày sheet IT_COURSE — chức năng xem khóa học công khai. "
            "Đây là cửa vào hệ thống: ai cũng xem được catalog, không cần đăng nhập. "
            "Em kiểm tra luồng Controller → Service → Repository: chỉ trả khóa PUBLISHED, "
            "không lộ khóa DRAFT."
        ),
        "demo": ["IT_COURSE_01", "IT_COURSE_03"],
        "cases": [
            {
                "id": "IT_COURSE_01",
                "title": "Lấy danh sách khóa học công khai",
                "muc_dich": "Chứng minh API catalog chỉ trả khóa đã xuất bản.",
                "noi": (
                    "Test case IT_COURSE_01 kiểm tra lấy danh sách khóa học công khai. "
                    "Em gọi GET /api/online-courses bằng MockMvc, không cần xác thực vì đây là API public. "
                    "Request vào PublicOnlineCourseController, controller gọi tiếp OnlineCourseService, "
                    "service truy vấn OnlineCourseRepository chỉ lấy khóa PUBLISHED. "
                    "Kết quả mong đợi HTTP 200, danh sách chỉ có khóa đã xuất bản, không có khóa DRAFT."
                ),
                "luong": "MockMvc → PublicOnlineCourseController → OnlineCourseService → OnlineCourseRepository → DB",
                "ket_qua": "200 OK; chỉ PUBLISHED; DRAFT không xuất hiện.",
                "neu_hoi": (
                    "Vì sao không cần token? Vì UC xem catalog là public. "
                    "Vì sao phải so với DB? Để chứng minh service filter đúng, không chỉ 'API trả 200'."
                ),
            },
            {
                "id": "IT_COURSE_02",
                "title": "Xem chi tiết một khóa học",
                "muc_dich": "Chi tiết khóa tồn tại trả 200; slug sai không được trả 500.",
                "noi": (
                    "IT_COURSE_02 kiểm tra chi tiết khóa. Em chọn một khóa PUBLISHED trong DB, "
                    "gọi GET /api/online-courses/{slugOrId}, không cần xác thực. "
                    "Controller gọi service getPublicCourse, service load từ repository. "
                    "Khóa tồn tại: 200, có title, price, status khớp DB. "
                    "Slug không tồn tại: 404 hoặc lỗi nghiệp vụ, không được 500."
                ),
                "luong": "PublicOnlineCourseController → OnlineCourseService.getPublicCourse → Repository",
                "ket_qua": "200 + đúng field; unknown slug → 404/not-found; read-only.",
                "neu_hoi": "Không ghi DB vì đây là API đọc. 404 chứng minh xử lý not-found đúng.",
            },
            {
                "id": "IT_COURSE_03",
                "title": "Tìm kiếm / lọc khóa học",
                "muc_dich": "Filter theo keyword/category hoạt động đúng.",
                "noi": (
                    "IT_COURSE_03 kiểm tra tìm kiếm. Em seed ít nhất 2 khóa PUBLISHED khác tiêu đề, "
                    "gọi GET /api/online-courses?keyword=... không cần xác thực. "
                    "Service lọc repository theo keyword. Mong đợi 200 và chỉ khóa khớp bộ lọc, "
                    "vẫn chỉ PUBLISHED."
                ),
                "luong": "Controller → Service.getPublicCourses(keyword) → Repository filter",
                "ket_qua": "200; đúng filter; không khớp thì bị loại.",
                "neu_hoi": "Case này chứng minh không chỉ list all mà còn search nghiệp vụ.",
            },
            {
                "id": "IT_COURSE_04",
                "title": "Tìm kiếm không có kết quả",
                "muc_dich": "Keyword không khớp vẫn 200 list rỗng, không crash.",
                "noi": (
                    "IT_COURSE_04 là luồng phụ: keyword chắc chắn không có. "
                    "Vẫn gọi API public, mong đợi 200 với list rỗng hoặc totalElements=0, không 5xx. "
                    "Em nhấn mạnh hệ thống xử lý 'không có dữ liệu' đúng cách."
                ),
                "luong": "Controller → Service → Repository (0 row)",
                "ket_qua": "200 + rỗng; không 500.",
                "neu_hoi": "Cô hỏi negative case: đây là empty result, không phải lỗi hệ thống.",
            },
            {
                "id": "IT_COURSE_05",
                "title": "Liệt kê danh mục khóa học",
                "muc_dich": "Chỉ trả category đang active.",
                "noi": (
                    "IT_COURSE_05 gọi GET /api/online-courses/categories, không cần xác thực. "
                    "Service lấy danh mục active. Mong đợi 200, chỉ category active, không ghi DB."
                ),
                "luong": "Controller → CourseCategoryManagementService → category repository",
                "ket_qua": "200; chỉ active categories.",
                "neu_hoi": "Danh mục phục vụ UI filter catalog.",
            },
        ],
    },
    {
        "sheet": "IT_ENROLL",
        "feature": "Đăng ký khóa học (Enroll in Course)",
        "open": (
            "Sheet IT_ENROLL — học viên gửi yêu cầu ghi danh và xem danh sách khóa đã ghi danh. "
            "Điểm quan trọng: dữ liệu phải theo đúng user đang đăng nhập, không lộ dữ liệu người khác."
        ),
        "demo": ["IT_ENROLL_01", "IT_ENROLL_02"],
        "cases": [
            {
                "id": "IT_ENROLL_01",
                "title": "Học viên gửi yêu cầu ghi danh",
                "muc_dich": "LEARNER tạo enrollment request hợp lệ; thiếu field / thiếu token bị từ chối.",
                "noi": (
                    "IT_ENROLL_01: học viên đăng nhập, gọi POST tạo enrollment request với JSON bắt buộc. "
                    "StudentEnrollmentRequestController gọi tiếp EnrollmentRequestService, service lưu qua repository. "
                    "Hợp lệ: 200/201. Thiếu field hoặc không cần xác thực (không gửi token): bị từ chối."
                ),
                "luong": "StudentEnrollmentRequestController → EnrollmentRequestService → Repository",
                "ket_qua": "Tạo thành công khi đủ điều kiện; thiếu dữ liệu/token thì 4xx.",
                "neu_hoi": "Đây là bước học viên xin vào khóa trước khi staff xếp lớp (IT_ASSIGN).",
            },
            {
                "id": "IT_ENROLL_02",
                "title": "Xem my-enrollments (chỉ của mình)",
                "muc_dich": "Chứng minh phân quyền dữ liệu theo user — case rất thuyết phục.",
                "noi": (
                    "IT_ENROLL_02 rất quan trọng. Em seed learner A và B ghi danh khác khóa. "
                    "Đăng nhập A, gọi GET /api/student/online-courses/my-enrollments kèm Bearer token. "
                    "StudentOnlineCourseController gọi OnlineCourseService.listMyEnrollments, "
                    "service chỉ query theo userId hiện tại. "
                    "Mong đợi 200, chỉ khóa của A, không có khóa của B — không bị lộ dữ liệu giữa user."
                ),
                "luong": "StudentOnlineCourseController → OnlineCourseService.listMyEnrollments → DB scoped by user",
                "ket_qua": "200; chỉ enrollment của A; B không xuất hiện.",
                "neu_hoi": (
                    "Cô hỏi 'test gì?': Em test data scoping / ownership. "
                    "Không chỉ API chạy, mà còn đúng quyền dữ liệu."
                ),
            },
        ],
    },
    {
        "sheet": "IT_CHECKOUT",
        "feature": "Thanh toán / Checkout",
        "open": (
            "Sheet IT_CHECKOUT — luồng thanh toán: tạo đơn, xem đơn, webhook PayOS, và chặn truy cập trái phép. "
            "Đây là feature tiền thật, cô thường hỏi kỹ."
        ),
        "demo": ["IT_CHECKOUT_01", "IT_CHECKOUT_04", "IT_CHECKOUT_05"],
        "cases": [
            {
                "id": "IT_CHECKOUT_01",
                "title": "Tạo đơn / payment link",
                "muc_dich": "Learner có giỏ hàng tạo được quote/order.",
                "noi": (
                    "IT_CHECKOUT_01: đăng nhập LEARNER có item trong cart, gọi API tạo payment link/order. "
                    "Controller gọi service thanh toán, service tạo bản ghi order và liên kết PayOS (stub/sandbox). "
                    "Mong đợi 200/201, có order id hoặc payment URL."
                ),
                "luong": "Payment/Checkout Controller → Payment Service → Order repository (+ PayOS stub)",
                "ket_qua": "Tạo order thành công; có thông tin thanh toán trả về.",
                "neu_hoi": "Môi trường test dùng stub/sandbox PayOS, không cần thanh toán thật.",
            },
            {
                "id": "IT_CHECKOUT_02",
                "title": "Xem danh sách đơn của tôi",
                "muc_dich": "Chỉ thấy order của chính mình.",
                "noi": (
                    "IT_CHECKOUT_02: GET danh sách đơn của learner đang login. "
                    "Service filter theo user. 200 và chỉ order của user đó."
                ),
                "luong": "Controller → Service list my orders → Repository by user",
                "ket_qua": "200; đúng ownership.",
                "neu_hoi": "Giống tinh thần my-enrollments: dữ liệu theo user.",
            },
            {
                "id": "IT_CHECKOUT_03",
                "title": "Checkout khi giỏ trống bị từ chối",
                "muc_dich": "Negative: không cho tạo đơn khi cart empty.",
                "noi": (
                    "IT_CHECKOUT_03: learner giỏ trống vẫn gọi checkout. "
                    "Service validate và từ chối. Mong đợi 4xx, không tạo order."
                ),
                "luong": "Controller → Service validate cart → reject",
                "ket_qua": "4xx; không insert order.",
                "neu_hoi": "Đây là business validation, không phải lỗi server.",
            },
            {
                "id": "IT_CHECKOUT_04",
                "title": "Webhook PayOS cập nhật trạng thái đơn",
                "muc_dich": "Tích hợp cổng thanh toán: callback đổi status order.",
                "noi": (
                    "IT_CHECKOUT_04 kiểm tra webhook: giả lập callback PayOS sau thanh toán. "
                    "Endpoint webhook nhận request, service cập nhật status order trong DB "
                    "(ví dụ PAID/SUCCESS). Mong đợi xử lý thành công và DB đổi đúng."
                ),
                "luong": "Webhook Controller → Payment Service → cập nhật Order repository",
                "ket_qua": "Order status đổi đúng theo callback.",
                "neu_hoi": (
                    "Webhook là server PayOS gọi lại hệ thống mình. "
                    "IT chứng minh Controller-Service-Repository cập nhật đơn đúng."
                ),
            },
            {
                "id": "IT_CHECKOUT_05",
                "title": "Checkout không xác thực bị chặn",
                "muc_dich": "Không token thì không checkout được.",
                "noi": (
                    "IT_CHECKOUT_05: gọi API checkout không cần xác thực (không gửi token). "
                    "Security filter chặn trước khi vào business. Mong đợi 401/403, không tạo đơn."
                ),
                "luong": "SecurityFilterChain chặn → không vào Controller nghiệp vụ",
                "ket_qua": "401/403; không tạo order.",
                "neu_hoi": "Chứng minh endpoint tiền bạc được bảo vệ.",
            },
        ],
    },
    {
        "sheet": "IT_CLASS",
        "feature": "Quản lý lớp học (Manage Classrooms)",
        "open": (
            "Sheet IT_CLASS — STAFF quản lý lớp: xem danh sách, tạo đề xuất lớp, cập nhật lớp. "
            "Lưu ý API thật dùng /api/staff/... không phải training-manager."
        ),
        "demo": ["IT_CLASS_01", "IT_CLASS_02"],
        "cases": [
            {
                "id": "IT_CLASS_01",
                "title": "STAFF xem danh sách lớp",
                "muc_dich": "STAFF list classrooms thành công.",
                "noi": (
                    "IT_CLASS_01: đăng nhập STAFF, GET /api/staff/classrooms kèm token. "
                    "Controller staff gọi service, service đọc repository. Mong đợi 200 và danh sách lớp."
                ),
                "luong": "Staff Classroom Controller → Service → Repository",
                "ket_qua": "200 + list classrooms.",
                "neu_hoi": "Role STAFF mới được vào; LEARNER sẽ 403.",
            },
            {
                "id": "IT_CLASS_02",
                "title": "Tạo đề xuất lớp (classroom proposal)",
                "muc_dich": "Luồng tạo lớp qua proposal, không POST create thẳng lung tung.",
                "noi": (
                    "IT_CLASS_02: STAFF tạo classroom proposal bằng POST /api/staff/classroom-proposals "
                    "với title, courseOfferingId, capacity, lịch học... "
                    "Controller gọi service tạo bản ghi proposal. Mong đợi 200/201, có id proposal trong DB."
                ),
                "luong": "Staff proposal Controller → Service → Repository save",
                "ket_qua": "Tạo proposal thành công; dữ liệu khớp request.",
                "neu_hoi": (
                    "Trong hệ thống thật, tạo lớp đi qua proposal/duyệt chứ không phải create tùy tiện. "
                    "Em test đúng flow đó."
                ),
            },
            {
                "id": "IT_CLASS_03",
                "title": "Cập nhật lớp học",
                "muc_dich": "STAFF cập nhật thông tin lớp và DB phản ánh đúng.",
                "noi": (
                    "IT_CLASS_03: STAFF gọi API update classroom, service cập nhật repository. "
                    "200 và field mới khớp DB."
                ),
                "luong": "Controller → Service update → Repository",
                "ket_qua": "200; DB updated.",
                "neu_hoi": "Case update chứng minh không chỉ đọc mà ghi nghiệp vụ đúng.",
            },
        ],
    },
    {
        "sheet": "IT_ASSIGN",
        "feature": "Xếp học viên vào lớp (Assign Learner to Classroom)",
        "open": (
            "Sheet IT_ASSIGN — sau khi học viên gửi enrollment request, STAFF xem/lọc request "
            "và xếp vào lớp còn chỗ, hoặc từ chối."
        ),
        "demo": ["IT_ASSIGN_01", "IT_ASSIGN_03"],
        "cases": [
            {
                "id": "IT_ASSIGN_01",
                "title": "STAFF xem danh sách enrollment requests",
                "muc_dich": "STAFF xem được hàng đợi; LEARNER không xem được hàng đợi staff.",
                "noi": (
                    "IT_ASSIGN_01: STAFF GET /api/staff/enrollment-requests. "
                    "Controller gọi EnrollmentRequestService.listForStaff. "
                    "STAFF được 200 có id, status, learner. Gọi bằng LEARNER thì 403."
                ),
                "luong": "StaffEnrollmentRequestController → EnrollmentRequestService → Repository",
                "ket_qua": "STAFF 200; LEARNER 403.",
                "neu_hoi": "Phân quyền theo role, không chỉ 'API có data'.",
            },
            {
                "id": "IT_ASSIGN_02",
                "title": "Lọc request theo status",
                "muc_dich": "Filter status hoạt động (ví dụ WAITING_FOR_CLASS).",
                "noi": (
                    "IT_ASSIGN_02: seed nhiều status, STAFF gọi GET kèm ?status=WAITING_FOR_CLASS. "
                    "Mọi dòng trả về phải đúng status đó."
                ),
                "luong": "Controller → Service filter by status → Repository",
                "ket_qua": "200; mọi row đúng status filter.",
                "neu_hoi": "Staff cần lọc để xử lý đúng hàng đợi.",
            },
            {
                "id": "IT_ASSIGN_03",
                "title": "Xếp học viên vào lớp",
                "muc_dich": "Happy path: request WAITING_FOR_CLASS → gán classroom còn chỗ.",
                "noi": (
                    "IT_ASSIGN_03 là case trung tâm. Có request WAITING_FOR_CLASS và lớp còn chỗ. "
                    "STAFF PATCH .../assign-class với classroomId. "
                    "Service cập nhật request thành CLASS_ASSIGNED và tạo enrollment vào lớp. "
                    "Nếu lớp đầy thì 4xx, không insert."
                ),
                "luong": "Controller → EnrollmentRequestService.assignToClassroom → cập nhật request + enrollment",
                "ket_qua": "200/201; status CLASS_ASSIGNED; đúng 1 enrollment; full class → 4xx.",
                "neu_hoi": "Nối ENROLL (học viên xin) với CLASS (có lớp) — đúng nghiệp vụ trung tâm.",
            },
            {
                "id": "IT_ASSIGN_04",
                "title": "Từ chối enrollment request",
                "muc_dich": "STAFF reject request; status đổi, không xếp lớp.",
                "noi": (
                    "IT_ASSIGN_04: STAFF reject request. Service đổi status rejected/tương đương, "
                    "không tạo classroom enrollment."
                ),
                "luong": "Controller → Service reject → Repository update",
                "ket_qua": "Request bị từ chối; không gán lớp.",
                "neu_hoi": "Luồng alternate của assign.",
            },
        ],
    },
    {
        "sheet": "IT_ASNTEACH",
        "feature": "Phân công giáo viên vào lớp",
        "open": (
            "Sheet IT_ASNTEACH — STAFF gán TEACHER vào classroom. "
            "Case âm: role khác STAFF bị từ chối."
        ),
        "demo": ["IT_ASNTEACH_01", "IT_ASNTEACH_02"],
        "cases": [
            {
                "id": "IT_ASNTEACH_01",
                "title": "Gán giáo viên vào lớp",
                "muc_dich": "STAFF assign teacher thành công.",
                "noi": (
                    "IT_ASNTEACH_01: STAFF gọi POST "
                    "/api/staff/classrooms/{id}/teachers/{teacherId}/assign. "
                    "Service lưu quan hệ teacher–classroom. Mong đợi 200/201, DB có phân công."
                ),
                "luong": "Staff Controller → Service assign teacher → Repository",
                "ket_qua": "Assign thành công; teacher gắn với class.",
                "neu_hoi": "Sau bước này teacher mới thấy lớp ở IT_SCHEDULE.",
            },
            {
                "id": "IT_ASNTEACH_02",
                "title": "Non-staff không được assign",
                "muc_dich": "LEARNER/TEACHER không tự assign được.",
                "noi": (
                    "IT_ASNTEACH_02: gọi cùng API bằng role không phải STAFF. "
                    "Bị 403, không tạo phân công."
                ),
                "luong": "Security/authorization chặn hoặc service reject",
                "ket_qua": "403; DB không đổi.",
                "neu_hoi": "Negative case chứng minh phân quyền.",
            },
        ],
    },
    {
        "sheet": "IT_SCHEDULE",
        "feature": "Xem lịch giảng dạy (View Teaching Schedule)",
        "open": (
            "Sheet IT_SCHEDULE — giáo viên xem lớp được phân công và các buổi học (sessions). "
            "Tách riêng với điểm danh và bài tập để đúng từng UC."
        ),
        "demo": ["IT_SCHEDULE_01", "IT_SCHEDULE_02"],
        "cases": [
            {
                "id": "IT_SCHEDULE_01",
                "title": "GV xem lớp được assign",
                "muc_dich": "Teacher chỉ thấy lớp của mình.",
                "noi": (
                    "IT_SCHEDULE_01: đăng nhập TEACHER, GET danh sách lớp assigned. "
                    "Service trả lớp đã phân công cho giáo viên đó. 200 và đúng ownership."
                ),
                "luong": "Teacher Controller → Service list assigned classrooms → Repository",
                "ket_qua": "200; chỉ lớp của teacher.",
                "neu_hoi": "Phụ thuộc IT_ASNTEACH đã gán teacher.",
            },
            {
                "id": "IT_SCHEDULE_02",
                "title": "GV tải danh sách buổi học",
                "muc_dich": "Xem sessions của một classroom.",
                "noi": (
                    "IT_SCHEDULE_02: teacher chọn lớp, GET sessions. "
                    "Service load buổi học theo classroom. 200 có danh sách session."
                ),
                "luong": "Teacher Controller → Service load sessions → Repository",
                "ket_qua": "200; sessions thuộc đúng lớp.",
                "neu_hoi": "Sessions là đầu vào cho điểm danh (IT_ATTEND).",
            },
        ],
    },
    {
        "sheet": "IT_ATTEND",
        "feature": "Quản lý điểm danh lớp",
        "open": (
            "Sheet IT_ATTEND — giáo viên mở bảng điểm danh và lưu điểm danh theo buổi."
        ),
        "demo": ["IT_ATTEND_02", "IT_ATTEND_01"],
        "cases": [
            {
                "id": "IT_ATTEND_02",
                "title": "Tải bảng điểm danh",
                "muc_dich": "Load attendance sheet của session.",
                "noi": (
                    "IT_ATTEND_02: teacher GET attendance sheet theo session. "
                    "Service trả danh sách học viên + trạng thái điểm danh. 200."
                ),
                "luong": "Teacher Controller → Attendance Service → Repository",
                "ket_qua": "200; có sheet điểm danh.",
                "neu_hoi": "Đọc trước, rồi mới save ở case kia.",
            },
            {
                "id": "IT_ATTEND_01",
                "title": "Lưu điểm danh buổi học",
                "muc_dich": "Persist attendance xuống DB.",
                "noi": (
                    "IT_ATTEND_01: teacher gửi payload điểm danh (present/absent...) cho session. "
                    "Service lưu repository. 200 và DB khớp request."
                ),
                "luong": "Controller → Service save attendance → Repository",
                "ket_qua": "2xx; attendance persisted.",
                "neu_hoi": "Đây là ghi nhận chuyên cần — nghiệp vụ lớp học quan trọng.",
            },
        ],
    },
    {
        "sheet": "IT_MNGHW",
        "feature": "Giáo viên quản lý bài tập (Manage Homework)",
        "open": (
            "Sheet IT_MNGHW — phía giáo viên: tạo bài tập và chấm bài. "
            "Khác IT_HOMEWORK (phía học viên nộp bài)."
        ),
        "demo": ["IT_MNGHW_01", "IT_MNGHW_02"],
        "cases": [
            {
                "id": "IT_MNGHW_01",
                "title": "GV tạo bài tập",
                "muc_dich": "Teacher tạo homework cho lớp.",
                "noi": (
                    "IT_MNGHW_01: TEACHER tạo homework qua API, service lưu homework repository. "
                    "200/201, bài tập tồn tại trong DB gắn đúng lớp."
                ),
                "luong": "Teacher Homework Controller → Service create → Repository",
                "ket_qua": "Tạo homework thành công.",
                "neu_hoi": "Sau đó learner mới nộp được ở IT_HOMEWORK.",
            },
            {
                "id": "IT_MNGHW_02",
                "title": "GV chấm bài nộp",
                "muc_dich": "Grade submission và lưu điểm/nhận xét.",
                "noi": (
                    "IT_MNGHW_02: có submission của học viên, teacher gọi API chấm điểm. "
                    "Service cập nhật điểm. 200 và DB có grade."
                ),
                "luong": "Controller → Service grade → Repository update submission",
                "ket_qua": "Điểm được lưu đúng.",
                "neu_hoi": "Khép vòng giao bài → nộp → chấm.",
            },
        ],
    },
    {
        "sheet": "IT_HOMEWORK",
        "feature": "Học viên nộp bài tập (Submit Homework)",
        "open": (
            "Sheet IT_HOMEWORK — học viên nộp bài, xem bài của mình, và bị từ chối nếu quá hạn."
        ),
        "demo": ["IT_HOMEWORK_01", "IT_HOMEWORK_03"],
        "cases": [
            {
                "id": "IT_HOMEWORK_01",
                "title": "Học viên nộp bài",
                "muc_dich": "Submit homework hợp lệ.",
                "noi": (
                    "IT_HOMEWORK_01: LEARNER nộp bài qua API. "
                    "StudentClassroomController gọi ClassroomHomeworkService.submit, lưu submission. "
                    "200/201, có bản ghi submission."
                ),
                "luong": "Student Controller → ClassroomHomeworkService.submit → Repository",
                "ket_qua": "Nộp thành công; submission tồn tại.",
                "neu_hoi": "Cần homework đang mở và learner thuộc lớp.",
            },
            {
                "id": "IT_HOMEWORK_02",
                "title": "Xem bài nộp của mình",
                "muc_dich": "Learner xem lại submission của chính mình.",
                "noi": (
                    "IT_HOMEWORK_02: GET submission của mình. 200 và đúng bài đã nộp, không lấy bài người khác."
                ),
                "luong": "Controller → Service get my submission → Repository",
                "ket_qua": "200; đúng ownership.",
                "neu_hoi": "Lại là data scoping.",
            },
            {
                "id": "IT_HOMEWORK_03",
                "title": "Nộp sau hạn / bài đã đóng bị từ chối",
                "muc_dich": "Negative rule nghiệp vụ.",
                "noi": (
                    "IT_HOMEWORK_03: cố nộp khi quá deadline hoặc homework closed. "
                    "Service từ chối 4xx, không tạo submission mới."
                ),
                "luong": "Controller → Service validate deadline → reject",
                "ket_qua": "4xx; không submit.",
                "neu_hoi": "Chứng minh rule hạn nộp được enforce ở backend.",
            },
        ],
    },
    {
        "sheet": "IT_QUIZ",
        "feature": "Làm bài Quiz",
        "open": (
            "Sheet IT_QUIZ — GV tạo/sửa/xóa quiz; học viên làm và nộp quiz."
        ),
        "demo": ["IT_QUIZ_01", "IT_QUIZ_03", "IT_QUIZ_04"],
        "cases": [
            {
                "id": "IT_QUIZ_01",
                "title": "GV tạo quiz",
                "muc_dich": "Teacher tạo quiz practice.",
                "noi": (
                    "IT_QUIZ_01: TEACHER tạo quiz qua API, service lưu quiz + câu hỏi. 200/201."
                ),
                "luong": "Teacher Quiz Controller → Service → Repository",
                "ket_qua": "Quiz được tạo trong DB.",
                "neu_hoi": "Đầu vào cho learner take quiz.",
            },
            {
                "id": "IT_QUIZ_02",
                "title": "GV cập nhật / xóa quiz",
                "muc_dich": "Update hoặc delete phản ánh DB.",
                "noi": (
                    "IT_QUIZ_02: teacher update hoặc delete quiz. 200/204 và DB đổi đúng."
                ),
                "luong": "Controller → Service update/delete → Repository",
                "ket_qua": "DB reflects update/delete.",
                "neu_hoi": "CRUD phía giáo viên.",
            },
            {
                "id": "IT_QUIZ_03",
                "title": "Học viên bắt đầu / làm quiz",
                "muc_dich": "Learner start attempt.",
                "noi": (
                    "IT_QUIZ_03: LEARNER start/take quiz. Service tạo attempt, trả câu hỏi. 200."
                ),
                "luong": "Student Quiz Controller → Service start → Repository",
                "ket_qua": "Bắt đầu làm bài thành công.",
                "neu_hoi": "Tách start và submit để rõ từng bước.",
            },
            {
                "id": "IT_QUIZ_04",
                "title": "Học viên nộp quiz",
                "muc_dich": "Submit answers và lưu kết quả.",
                "noi": (
                    "IT_QUIZ_04: learner nộp câu trả lời. Service chấm/lưu kết quả. 200, có score hoặc submission persisted."
                ),
                "luong": "Controller → Service submit → Repository",
                "ket_qua": "Nộp thành công; kết quả được lưu.",
                "neu_hoi": "Khép vòng tạo quiz → làm → nộp.",
            },
        ],
    },
    {
        "sheet": "IT_ONLINE",
        "feature": "Quản lý khóa học online (CM)",
        "open": (
            "Sheet IT_ONLINE — Content Manager quản lý khóa online: list, create, update; "
            "role khác CM không được quản lý."
        ),
        "demo": ["IT_ONLINE_02", "IT_ONLINE_04"],
        "cases": [
            {
                "id": "IT_ONLINE_01",
                "title": "CM xem danh sách khóa online",
                "muc_dich": "CM list courses quản trị.",
                "noi": (
                    "IT_ONLINE_01: đăng nhập CM, GET list online courses phía quản trị. 200."
                ),
                "luong": "CM Controller → OnlineCourseService → Repository",
                "ket_qua": "200 + danh sách quản trị.",
                "neu_hoi": "Khác catalog public (IT_COURSE): đây là màn quản lý.",
            },
            {
                "id": "IT_ONLINE_02",
                "title": "CM tạo khóa online",
                "muc_dich": "Create course lưu DB.",
                "noi": (
                    "IT_ONLINE_02: CM POST tạo khóa. Service save OnlineCourseRepository. 200/201 có course mới."
                ),
                "luong": "Controller → Service create → Repository.save",
                "ket_qua": "Course được tạo.",
                "neu_hoi": "Sau publish mới xuất hiện ở IT_COURSE public.",
            },
            {
                "id": "IT_ONLINE_03",
                "title": "CM cập nhật khóa",
                "muc_dich": "Update field khóa học.",
                "noi": (
                    "IT_ONLINE_03: CM update title/price/status... DB khớp payload mới."
                ),
                "luong": "Controller → Service update → Repository",
                "ket_qua": "200; DB updated.",
                "neu_hoi": "CRUD update.",
            },
            {
                "id": "IT_ONLINE_04",
                "title": "Non-CM không quản lý được",
                "muc_dich": "Phân quyền CM.",
                "noi": (
                    "IT_ONLINE_04: LEARNER/role khác gọi API quản trị course → 403, không tạo/sửa được."
                ),
                "luong": "Authorization chặn non-CM",
                "ket_qua": "403; không đổi DB.",
                "neu_hoi": "Negative quyền — cô rất hay hỏi.",
            },
        ],
    },
    {
        "sheet": "IT_ADMIN",
        "feature": "Quản lý tài khoản người dùng",
        "open": (
            "Sheet IT_ADMIN — ADMIN xem danh sách user và khóa/mở khóa tài khoản."
        ),
        "demo": ["IT_ADMIN_01", "IT_ADMIN_02"],
        "cases": [
            {
                "id": "IT_ADMIN_01",
                "title": "ADMIN xem danh sách user",
                "muc_dich": "Admin list users; LEARNER/không token bị chặn.",
                "noi": (
                    "IT_ADMIN_01: ADMIN GET /api/admin/users. 200 có danh sách. "
                    "Gọi bằng LEARNER hoặc không cần xác thực thì 401/403."
                ),
                "luong": "AdminUserController → AdminUserService → User repository",
                "ket_qua": "ADMIN 200; khác role/no auth bị chặn.",
                "neu_hoi": "API quản trị luôn phải khóa role.",
            },
            {
                "id": "IT_ADMIN_02",
                "title": "Khóa / mở khóa tài khoản",
                "muc_dich": "Lock/unlock user phản ánh DB.",
                "noi": (
                    "IT_ADMIN_02: ADMIN gọi lock hoặc unlock user. "
                    "Service cập nhật trạng thái tài khoản. 200 và user bị khóa/mở đúng."
                ),
                "luong": "Controller → Service lock/unlock → Repository",
                "ket_qua": "Trạng thái user đổi đúng.",
                "neu_hoi": "Nghiệp vụ an toàn tài khoản.",
            },
        ],
    },
    {
        "sheet": "IT_BROADCAST",
        "feature": "Thông báo hệ thống (Admin Broadcast)",
        "open": (
            "Sheet IT_BROADCAST — ADMIN tạo/xem/sửa/hủy broadcast hệ thống. "
            "Khác IT_NOTIF (hộp thư cá nhân của learner). "
            "Lưu ý body cần sendInApp và sendEmail."
        ),
        "demo": ["IT_BROADCAST_01", "IT_BROADCAST_04"],
        "cases": [
            {
                "id": "IT_BROADCAST_01",
                "title": "ADMIN tạo broadcast",
                "muc_dich": "Create broadcast hợp lệ; thiếu quyền bị chặn.",
                "noi": (
                    "IT_BROADCAST_01: ADMIN POST /api/admin/broadcasts với title, message, "
                    "sendInApp, sendEmail. AdminBroadcastController gọi AdminBroadcastService.create, "
                    "lưu repository. 200 có id. Không token hoặc LEARNER: 401/403."
                ),
                "luong": "AdminBroadcastController → AdminBroadcastService → AdminBroadcastRepository",
                "ket_qua": "ADMIN tạo được; role khác không tạo được.",
                "neu_hoi": "Nhấn mạnh khác inbox thông báo cá nhân (IT_NOTIF).",
            },
            {
                "id": "IT_BROADCAST_02",
                "title": "ADMIN xem danh sách broadcast",
                "muc_dich": "List/page broadcasts.",
                "noi": (
                    "IT_BROADCAST_02: GET /api/admin/broadcasts?page&size. 200, có phân trang, "
                    "không cần xác thực thì bị chặn."
                ),
                "luong": "Controller → Service list → Repository",
                "ket_qua": "200 page data; no auth → 401/403.",
                "neu_hoi": "Read path của quản trị broadcast.",
            },
            {
                "id": "IT_BROADCAST_03",
                "title": "Cập nhật broadcast nháp/lên lịch",
                "muc_dich": "Update draft/scheduled.",
                "noi": (
                    "IT_BROADCAST_03: PUT cập nhật title/message của broadcast chưa gửi. "
                    "200 và DB khớp payload mới."
                ),
                "luong": "Controller → Service update → Repository",
                "ket_qua": "Update thành công trên draft/scheduled.",
                "neu_hoi": "Thường chỉ sửa được trước khi gửi.",
            },
            {
                "id": "IT_BROADCAST_04",
                "title": "Hủy broadcast",
                "muc_dich": "Cancel theo đúng điều kiện nghiệp vụ.",
                "noi": (
                    "IT_BROADCAST_04: ADMIN hủy broadcast (thường sau khi đã schedule). "
                    "Service đổi status canceled. 200 và không còn gửi tiếp."
                ),
                "luong": "Controller → Service cancel → Repository status update",
                "ket_qua": "Hủy thành công đúng rule.",
                "neu_hoi": "Nếu cô hỏi fail: cancel chỉ hợp lệ ở trạng thái cho phép.",
            },
        ],
    },
]


FAQ = [
    (
        "Integration Test khác Unit Test chỗ nào?",
        "Unit Test mock phụ thuộc, test từng class. "
        "Integration Test em chạy @SpringBootTest + MockMvc: request thật qua Controller, "
        "vào Service, xuống Repository/DB — kiểm tra các tầng gắn với nhau.",
    ),
    (
        "Vì sao dùng MockMvc không gọi UI?",
        "Báo cáo 5.2 là Integration Test backend API. "
        "MockMvc giả lập HTTP vào Spring MVC, không cần mở browser, vẫn đi đủ filter, controller, service.",
    ),
    (
        "Passed nghĩa là gì?",
        "API chạy đúng expected: HTTP status, dữ liệu JSON/DB, và phân quyền (role/ownership) như mô tả test case.",
    ),
    (
        "N/A nghĩa là gì? (ví dụ GMEET)",
        "Không phải fail. Thiếu môi trường bên ngoài (Google Meet provider). "
        "Em tách N/A để trung thực, không đánh Passed giả.",
    ),
    (
        "Controller – Service – Repository là gì?",
        "Controller nhận HTTP; Service chứa nghiệp vụ; Repository truy cập DB. "
        "Mỗi test case em đều nêu rõ request đi qua 3 tầng này.",
    ),
    (
        "Em có test hết hệ thống không?",
        "Em cover các Use Case chính theo SRS, chia sheet theo feature. "
        "Hôm nay em đi sâu các luồng: catalog → ghi danh/thanh toán → lớp học → học tập → quản trị.",
    ),
]


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(64)
    section.right_margin = Pt(64)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("HƯỚNG DẪN THUYẾT TRÌNH INTEGRATION TEST")
    set_run_font(r, size=18, bold=True)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("EnglishLab — SEP490_G23 — Kịch bản nói từng Test Case (không gồm AUTH)")
    set_run_font(r2, size=12, bold=True)

    add_p(
        doc,
        "Mục tiêu: nói rõ – đúng – không lan man. Mỗi case chỉ cần 30–45 giây theo khung cố định dưới đây.",
        italic=True,
    )

    add_heading(doc, "1. Cách nói để cô thấy bạn hiểu (bắt buộc thuộc)", level=1)
    add_p(doc, "Mỗi test case nói đúng 5 ý theo thứ tự — đừng đảo, đừng thêm chuyện ngoài lề:")
    for i, x in enumerate(
        [
            "Tên ID + đang test chức năng gì (1 câu).",
            "Gọi API nào, role nào, có cần đăng nhập không.",
            "Luồng: Controller nào → Service nào → Repository/DB làm gì.",
            "Kết quả mong đợi: HTTP + dữ liệu/DB + (nếu có) quyền.",
            "Ý nghĩa nghiệp vụ: case này chứng minh điều gì cho hệ thống.",
        ],
        1,
    ):
        add_bullet(doc, f"{i}. {x}")

    add_heading(doc, "Câu mở đầu cả buổi (đọc thuộc)", level=2)
    add_say(
        doc,
        "Em trình bày Integration Test backend EnglishLab. Em dùng SpringBootTest và MockMvc: "
        "gửi HTTP vào Controller, đi tiếp Service rồi Repository/DB. "
        "Excel chia theo từng Feature/Use Case. Hôm nay em không đi Auth, "
        "em tập trung luồng nghiệp vụ chính: xem khóa học, ghi danh và thanh toán, "
        "quản lý lớp, giảng dạy và đánh giá, rồi quản trị hệ thống.",
    )

    add_heading(doc, "Câu mở đầu mỗi sheet", level=2)
    add_say(
        doc,
        "Em chuyển sang sheet [TÊN SHEET] — chức năng [FEATURE]. "
        "Sheet này có [N] test case. Em xin trình bày các case chính sau.",
    )

    add_heading(doc, "Những câu KHÔNG nói (dễ bị bảo nói linh tinh)", level=2)
    for x in [
        "Không nói: 'Em test xem nó chạy được không' — quá chung.",
        "Không đọc máy móc cả đoạn Procedure tiếng Anh dài.",
        "Không nói tên class sai hoặc bịa flow không có trong Excel.",
        "Không đổ thừa 'tại AI viết' / 'em không nhớ'.",
        "Không nói Passed nếu chưa giải thích expected.",
    ]:
        add_bullet(doc, x)

    add_heading(doc, "2. Lộ trình thuyết trình gợi ý (khoảng 10–15 phút)", level=1)
    add_p(doc, "Nếu ít thời gian: chỉ demo case đánh dấu ★ trong từng sheet.")
    steps = [
        "COURSE (catalog) → ENROLL (ghi danh) → CHECKOUT (thanh toán)",
        "CLASS → ASSIGN → ASNTEACH (vận hành lớp)",
        "SCHEDULE → ATTEND → MNGHW/HOMEWORK → QUIZ (dạy & học)",
        "ONLINE → ADMIN → BROADCAST (nội dung & quản trị)",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"Phần {i}: {s}")

    add_heading(doc, "3. Kịch bản từng sheet / từng test case", level=1)

    for sh in SHEETS:
        add_heading(doc, f"{sh['sheet']} — {sh['feature']}", level=2)
        add_p(doc, "Mở sheet:", bold=True)
        add_say(doc, sh["open"])
        add_p(
            doc,
            "Case nên demo nếu thiếu thời gian: " + ", ".join("★ " + x for x in sh["demo"]),
            bold=True,
        )

        for c in sh["cases"]:
            star = " ★" if c["id"] in sh["demo"] else ""
            add_heading(doc, f"{c['id']}{star} — {c['title']}", level=3)
            add_p(doc, f"Mục đích: {c['muc_dich']}")
            add_p(doc, f"Luồng kỹ thuật: {c['luong']}", italic=True)
            add_p(doc, f"Kết quả chốt: {c['ket_qua']}", bold=True)
            add_say(doc, c["noi"])
            add_tip(doc, c["neu_hoi"])

    add_heading(doc, "4. Câu chuyển scene giữa các phần (để mạch lạc)", level=1)
    add_say(
        doc,
        "Sau khi học viên xem khóa và ghi danh/thanh toán, hệ thống cần lớp học để học. "
        "Em chuyển sang phần STAFF quản lý lớp và xếp học viên, giáo viên.",
    )
    add_say(
        doc,
        "Khi lớp đã có giáo viên và lịch, em sang phần giảng dạy: điểm danh, bài tập, quiz.",
    )
    add_say(
        doc,
        "Cuối cùng em trình bày phía quản trị: quản lý user và gửi thông báo hệ thống.",
    )

    add_heading(doc, "5. Câu kết buổi", level=1)
    add_say(
        doc,
        "Tóm lại, bộ Integration Test của em bám Use Case SRS, mỗi sheet một feature. "
        "Em kiểm tra không chỉ HTTP 200 mà còn nghiệp vụ: filter đúng, phân quyền đúng role, "
        "và dữ liệu đúng theo user. Các case chính em vừa trình bày đều Passed trên môi trường test. "
        "Em xin cảm ơn cô và sẵn sàng trả lời câu hỏi.",
    )

    add_heading(doc, "6. Hỏi–đáp nhanh (học thuộc ý, không học vẹt chữ)", level=1)
    for q, a in FAQ:
        add_p(doc, f"Q: {q}", bold=True)
        add_say(doc, a)

    add_heading(doc, "7. Checklist 1 phút trước khi lên", level=1)
    for x in [
        "Mở sẵn Excel TIENG_VIET hoặc COMPLETED, đúng sheet sắp nói.",
        "Nhớ 3 tầng: Controller – Service – Repository.",
        "Với mỗi case: API + role + expected.",
        "Biết phân biệt Passed / N/A.",
        "Biết nối story: COURSE → ENROLL/CHECKOUT → CLASS/ASSIGN → học tập → ADMIN.",
        "Nếu quên chi tiết: nói đúng khung 5 ý, không bịa thêm.",
    ]:
        add_bullet(doc, x)

    add_heading(doc, "8. Bản rút gọn 5 phút (nếu cô bảo nói nhanh)", level=1)
    add_p(doc, "Chỉ nói các case ★ sau, mỗi case ~40 giây:")
    short = [
        "IT_COURSE_01 — catalog public, chỉ PUBLISHED",
        "IT_ENROLL_02 — my-enrollments không lộ data user khác",
        "IT_CHECKOUT_01 + IT_CHECKOUT_05 — tạo đơn và chặn không token",
        "IT_CLASS_02 — tạo classroom proposal",
        "IT_ASSIGN_03 — xếp học viên vào lớp",
        "IT_HOMEWORK_01 hoặc IT_QUIZ_04 — nộp bài / nộp quiz",
        "IT_ONLINE_04 — non-CM bị chặn",
        "IT_ADMIN_02 hoặc IT_BROADCAST_01 — quản trị",
    ]
    for x in short:
        add_bullet(doc, x)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    try:
        OUT_DL.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUT_DL)
    except Exception:
        pass
    print(f"OUT: {OUT}")
    if OUT_DL.exists():
        print(f"OUT: {OUT_DL}")


if __name__ == "__main__":
    build()
