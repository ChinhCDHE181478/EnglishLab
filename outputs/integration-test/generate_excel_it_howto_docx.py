# -*- coding: utf-8 -*-
"""Word: Hướng dẫn làm Integration Test với file Excel Report 5.2."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Huong_dan_lam_Integration_Test_voi_file_Excel.docx"

EXCEL_MAIN = "SEP490_G23_Report5.2_Integration Test_COMPLETED_HONEST_FORMATTED.xlsx"
EXCEL_ALT = "SEP490_G23_Report5.2_Integration Test_COMPLETED_HONEST.xlsx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    font(run, size=11)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    font(run, size=11)


def tip(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("💡 " + text)
    font(run, size=11, color=RGBColor(0x1F, 0x4E, 0x79))


def warn(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("⚠ " + text)
    font(run, size=11, color=RGBColor(0xC0, 0x39, 0x2B))


def say(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("🎤 Nói với cô: " + text)
    font(run, size=11, color=RGBColor(0x27, 0x6E, 0x49))


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=9, name="Consolas")
    para.paragraph_format.left_indent = Cm(0.4)
    para.paragraph_format.space_after = Pt(4)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hd)
        font(r, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            font(r, size=9)
    doc.add_paragraph()


SHEETS = [
    ("IT - Auth", "Đăng ký / OTP / Login JWT", "10"),
    ("IT - User", "Hồ sơ user", "5"),
    ("IT - Notif", "Thông báo", "5"),
    ("IT - Commerce", "Giỏ / wishlist", "4"),
    ("IT - Payment", "PayOS / orders", "5"),
    ("IT - Course", "Khóa online", "6"),
    ("IT - Discuss", "Thảo luận", "5"),
    ("IT - Content", "CM khóa học", "4"),
    ("IT - Package", "Package", "3"),
    ("IT - Curriculum", "Curriculum / bank", "5"),
    ("IT - EnrollReq", "Form tư vấn", "5"),
    ("IT - Classroom", "TM lớp học", "8"),
    ("IT - LearnerCls", "HV trong lớp", "6"),
    ("IT - Teacher", "Giáo viên", "6"),
    ("IT - Quiz", "Quiz lớp", "4"),
    ("IT - Assess", "Placement / mock", "6"),
    ("IT - Support", "Support ticket", "4"),
    ("IT - Admin", "Admin", "4"),
    ("IT - Lark", "Lark", "3"),
    ("IT - Infra", "Campus / room", "3"),
    ("IT - Report", "Dashboard / revenue", "2"),
    ("IT - Proposal", "Đề xuất lớp", "3"),
    ("IT - Dispute", "Khiếu nại điểm danh", "3"),
    ("IT - Notes", "Learning notes", "2"),
]


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "HƯỚNG DẪN LÀM INTEGRATION TEST\n"
        "VỚI FILE EXCEL REPORT 5.2\n"
        "DỰ ÁN ENGLISHLAB (SEP490_G23)"
    )
    font(r, size=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Đọc Excel → chạy test → ghi Round → giải thích với cô")
    font(r, size=12)

    p(doc, "Người thực hiện: phongdx", bold=True)
    p(doc, f"File Excel chính (nên dùng): {EXCEL_MAIN}")
    p(doc, f"File dự phòng: {EXCEL_ALT}")
    p(doc, "Vị trí thường gặp: thư mục Downloads\\intergration test  hoặc  outputs\\integration-test")

    # ===== 1 =====
    h(doc, "1. Tài liệu này dùng để làm gì?", 1)
    p(
        doc,
        "Giúp bạn (kể cả mới dùng máy) biết cách mở đúng file Excel Integration Test, "
        "đọc từng cột test case, chạy kiểm thử, rồi ghi kết quả vào đúng ô Round 1–3.",
    )
    say(
        doc,
        "Em làm Integration Test theo đúng file Excel Report 5.2: mỗi dòng là một test case "
        "tích hợp Controller–Service–Repository; em thực thi rồi ghi Passed/Failed/N/A.",
    )

    # ===== 2 =====
    h(doc, "2. File Excel gồm những sheet nào?", 1)
    table(
        doc,
        ["Sheet", "Dùng để làm gì"],
        [
            ["Cover", "Trang bìa / thông tin báo cáo"],
            ["Test Cases", "Danh sách module → bấm link sang sheet IT"],
            ["Test Statistics", "Thống kê Passed/Failed/N/A (công thức)"],
            ["IT - Auth … IT - Notes", "24 sheet chi tiết — nơi bạn làm và ghi kết quả"],
        ],
    )
    tip(doc, "Làm việc chính ở các sheet bắt đầu bằng IT - … Không sửa lung tung cột Procedure/Expected.")

    h(doc, "2.1. Bảng 24 module (sheet IT)", 2)
    table(
        doc,
        ["Sheet", "Nội dung ngắn", "Số case (khoảng)"],
        [[a, b, c] for a, b, c in SHEETS],
    )
    p(doc, "Tổng khoảng 111 test case.", bold=True)

    # ===== 3 =====
    h(doc, "3. Mở file Excel đúng cách (từng click)", 1)
    numbered(doc, "Mở thư mục có file Excel (Downloads\\intergration test hoặc outputs\\integration-test).")
    numbered(doc, f"Chọn file: {EXCEL_MAIN}")
    numbered(doc, "Mở bằng Microsoft Excel (khuyến nghị) hoặc Excel online.")
    numbered(doc, "Nếu Excel hỏi Enable Editing / Enable Content → bấm Enable để công thức Statistics chạy.")
    numbered(doc, "Nhìn thanh sheet phía dưới cửa sổ Excel → kéo sang phải để thấy IT - Auth, IT - User…")
    warn(doc, "Ưu tiên bản HONEST / HONEST_FORMATTED. Không dùng bản “111 Passed” ảo nếu cô hỏi kết quả trung thực.")

    # ===== 4 =====
    h(doc, "4. Cấu trúc 1 sheet IT (ví dụ IT - Auth)", 1)
    p(doc, "Phần đầu sheet (dòng trên):", bold=True)
    bullet(doc, "Feature: Controller chính (ví dụ AuthController)")
    bullet(doc, "Test requirement: yêu cầu tích hợp cần kiểm")
    bullet(doc, "Testing Round: bảng tóm tắt Round 1/2/3 (Passed/Failed/Pending/N/A)")

    p(doc, "Bảng test case (thường từ dòng có tiêu đề cột):", bold=True)
    table(
        doc,
        ["Tên cột trên Excel", "Ý nghĩa", "Bạn làm gì với cột này"],
        [
            ["Test Case ID", "Mã case (IT_AUTH_01…)", "Đối chiếu khi chạy / thuyết trình"],
            ["Test Case Description", "Mục tiêu tích hợp", "Đọc để hiểu đang test gì"],
            ["Test Case Procedure", "Các bước + hàm Contoller→Service→Repo", "Làm theo từng bước khi chạy"],
            ["Expected Results", "Kết quả đúng", "So với thực tế sau khi chạy"],
            ["Pre-conditions", "Điều kiện trước", "Thiếu thì có thể ghi N/A"],
            ["Round 1 / 2 / 3", "Kết quả lần chạy", "Ghi Passed / Failed / N/A"],
            ["Test date", "Ngày chạy", "Ghi ngày (vd 2026-08-04)"],
            ["Tester", "Người chạy", "Ghi tên bạn (phongdx)"],
        ],
    )
    tip(doc, "Dòng chữ nhóm như “Register & verify” / “Login & security” chỉ là tiêu đề nhóm — không phải test case, khỏi ghi Round.")

    # ===== 5 =====
    h(doc, "5. Quy trình làm 1 test case với Excel (in ra dán bàn)", 1)
    numbered(doc, "Mở đúng sheet (ví dụ IT - Auth).")
    numbered(doc, "Tìm dòng có Test Case ID cần làm (ví dụ IT_AUTH_05).")
    numbered(doc, "Đọc Pre-conditions — đã đủ chưa? (DB, user, đã enroll…)")
    numbered(doc, "Đọc Description — đang tích hợp những gì?")
    numbered(doc, "Đọc Procedure — làm lần lượt từng bước (gọi API / chạy code test).")
    numbered(doc, "Đọc Expected Results — nhớ Status, dữ liệu, hành vi cần có.")
    numbered(doc, "Chạy thật (Postman hoặc JUnit) theo Procedure.")
    numbered(doc, "So sánh Actual với Expected.")
    numbered(doc, "Ghi vào Excel:")
    bullet(doc, "Cột Round 1 (hoặc 2/3): Passed hoặc Failed hoặc N/A")
    bullet(doc, "Cột Test date cạnh Round đó: ngày hôm nay")
    bullet(doc, "Cột Tester: tên bạn")
    bullet(doc, "Nếu cần giải thích: có thể ghi thêm vào Note/ô trống gần đó (HTTP 200, thiếu OTP…)")
    numbered(doc, "Sang case tiếp theo trên cùng sheet.")

    h(doc, "5.1. Ví dụ cụ thể: IT_AUTH_05 trên sheet IT - Auth", 2)
    p(doc, "1) Mở sheet IT - Auth → tìm IT_AUTH_05.", bold=True)
    p(doc, "2) Pre-condition: có LEARNER đã verify (dùng 0386852628z@gmail.com / Password123!).")
    p(doc, "3) Procedure tóm tắt trên Excel: POST /api/auth/login → AuthController/AuthService → JWT → GET /api/user/me.")
    p(doc, "4) Trên máy (Postman):")
    bullet(doc, "POST http://localhost:8080/api/auth/login với email/password Learner")
    code(doc, '{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!"\n}')
    bullet(doc, "Copy accessToken → GET http://localhost:8080/api/user/me Header Authorization: Bearer <token>")
    p(doc, "5) Expected: login 200 có token; /me 200 đúng email.")
    p(doc, "6) Ghi Excel Round 1 = Passed, Test date = hôm nay, Tester = phongdx.")

    # ===== 6 =====
    h(doc, "6. Cách ghi Passed / Failed / N/A cho đúng ý cô", 1)
    table(
        doc,
        ["Ghi vào Round", "Khi nào", "Ví dụ"],
        [
            ["Passed", "Actual khớp Expected (kể cả negative đúng)", "Sai mật khẩu → 401 đúng expected → Passed"],
            ["Failed", "Sai expected, lỗi 500, API hỏng", "Lẽ ra 200 mà ra 500"],
            ["N/A", "Thiếu precondition nên chưa kết luận đủ", "Chưa enroll khóa; không lấy được OTP; waitlist < 2"],
        ],
    )
    warn(doc, "Negative test bị hệ thống từ chối đúng → Passed, không ghi Failed.")
    say(
        doc,
        "N/A là thiếu điều kiện môi trường, không phải giấu Fail. Failed mới là hệ thống sai expected.",
    )

    # ===== 7 =====
    h(doc, "7. Chuẩn bị môi trường trước khi chạy theo Excel", 1)
    numbered(doc, "PostgreSQL chạy + database englishlab có data demo.")
    numbered(doc, "Backend Spring Boot chạy cổng 8080.")
    numbered(doc, "Có Postman (hoặc chạy JUnit AuthIT) để thực thi Procedure.")
    numbered(doc, "Tài khoản demo (password Password123!):")
    bullet(doc, "LEARNER: 0386852628z@gmail.com")
    bullet(doc, "TEACHER: classroom.teacher1@englishlab.vn")
    bullet(doc, "TM: training.manager@englishlab.vn")
    bullet(doc, "STAFF / MANAGER / CM / ADMIN: staff@…, classroom.manager@…, content.manager@…, classroom.admin@englishlab.vn")

    tip(doc, "OTP verify/reset: xem bảng auth_tokens; field API là code.")

    # ===== 8 =====
    h(doc, "8. Thứ tự nên làm các sheet Excel", 1)
    numbered(doc, "IT - Auth → IT - User → IT - Notif (nền tảng)")
    numbered(doc, "IT - Classroom → IT - LearnerCls → IT - Teacher → IT - Quiz → IT - Dispute → IT - Notes")
    numbered(doc, "IT - Commerce → IT - Payment")
    numbered(doc, "IT - Course → IT - Discuss → IT - Assess")
    numbered(doc, "Các sheet còn lại: EnrollReq, Content, Package, Curriculum, Infra, Report, Proposal, Support, Admin, Lark")

    # ===== 9 =====
    h(doc, "9. Sheet Test Cases & Test Statistics", 1)
    h(doc, "9.1. Test Cases", 2)
    p(
        doc,
        "Là mục lục: mỗi dòng có Function Name / Sheet Name. Thường có hyperlink — bấm Sheet Name để nhảy sang IT - …",
    )
    h(doc, "9.2. Test Statistics", 2)
    p(
        doc,
        "Sheet thống kê số Passed/Failed/N/A theo module. Mở bằng Excel để công thức tự đếm. "
        "Sau khi bạn sửa Round trên sheet IT, Statistics sẽ cập nhật (có thể cần tính lại công thức).",
    )
    tip(doc, "Nếu số liệu lạ: kiểm tra đã ghi đúng chữ Passed/Failed/N/A (đúng chính tả, đúng ô Round).")

    # ===== 10 =====
    h(doc, "10. Round 1 / Round 2 / Round 3 dùng thế nào?", 1)
    bullet(doc, "Round 1: lần chạy đầu.")
    bullet(doc, "Round 2: chạy lại sau khi sửa lỗi / bổ sung precondition.")
    bullet(doc, "Round 3: lần xác nhận cuối (ổn định).")
    p(
        doc,
        "Nếu case đã Passed từ Round 1 và không đổi code: có thể giữ Passed các round sau "
        "(hoặc chạy lại cho chắc trước khi nộp).",
    )

    # ===== 11 =====
    h(doc, "11. Liên hệ Excel với tool chạy (không nhầm vai trò)", 1)
    table(
        doc,
        ["Cái gì", "Vai trò"],
        [
            ["File Excel", "Bộ TEST CASE Integration (sản phẩm chính)"],
            ["Postman / JUnit", "Tool thực thi Procedure để lấy Actual"],
            ["Cột Procedure trên Excel", "Mô tả tích hợp hàm — đúng ý “inte = hàm với nhau”"],
        ],
    )
    say(
        doc,
        "Em không lấy Postman thay cho test case. Em dùng Excel làm bộ case tích hợp hàm; "
        "Postman/JUnit chỉ để chạy và ghi Round.",
    )

    # ===== 12 =====
    h(doc, "12. Checklist trước khi nộp / gặp cô", 1)
    bullet(doc, f"Đúng file Excel bản HONEST ({EXCEL_MAIN} hoặc HONEST.xlsx).")
    bullet(doc, "Biết mở sheet IT - Auth và chỉ đúng dòng IT_AUTH_05 / IT_AUTH_06.")
    bullet(doc, "Giải thích được 1 cột Procedure theo Controller→Service→Repo.")
    bullet(doc, "Round đã có Tester + Test date.")
    bullet(doc, "Phân biệt Passed negative và N/A.")
    bullet(doc, "Statistics mở bằng Excel không bị lỗi công thức.")

    # ===== 13 =====
    h(doc, "13. Kịch bản thuyết trình ngắn với cô (kèm Excel)", 1)
    say(
        doc,
        "Dạ em trình bày Integration Test trên file Excel Report 5.2. Mỗi sheet IT là một module; "
        "mỗi dòng IT_* là một test case tích hợp hàm. Em đọc Pre-condition và Procedure, thực thi, "
        "đối chiếu Expected rồi ghi Round.",
    )
    say(
        doc,
        "Ví dụ em mở IT - Auth, case IT_AUTH_05: login qua AuthController/AuthService cấp JWT, "
        "rồi /me qua Security Filter — Expected 200; em ghi Passed.",
    )
    say(
        doc,
        "Case negative IT_AUTH_06 sai mật khẩu bị từ chối đúng expected cũng Passed. "
        "Case thiếu OTP/enroll em ghi N/A và Note lý do.",
    )

    h(doc, "13.1. Cô hỏi – đáp", 2)
    table(
        doc,
        ["Cô hỏi", "Em trả lời"],
        [
            ["Inte là gì?", "Tích hợp các hàm/tầng với nhau; trên Excel cột Procedure mô tả luồng đó."],
            ["Postman liên quan gì?", "Chỉ tool chạy API lấy actual; test case nằm trên Excel."],
            ["Em viết test case ở đâu?", "Trên các sheet IT - … trong file Excel Report 5.2."],
            ["Làm sao biết đủ module?", "Sheet Test Cases / 24 sheet IT; khoảng 111 case."],
            ["Kết quả lấy ở đâu?", "Cột Round 1–3; tổng hợp ở Test Statistics."],
        ],
    )

    # ===== 14 =====
    h(doc, "14. Lỗi thường gặp khi làm với Excel", 1)
    table(
        doc,
        ["Hiện tượng", "Cách xử lý"],
        [
            ["Không thấy sheet IT", "Kéo thanh sheet dưới cùng sang phải"],
            ["Ghi Round nhưng Statistics không đổi", "Enable Content; mở bằng Excel; kiểm tra chính tả Passed/Failed/N/A"],
            ["Nhầm ghi vào dòng nhóm (Register & verify)", "Chỉ ghi vào dòng có mã IT_xxx"],
            ["Không hiểu Procedure tiếng Anh", "Đọc tên hàm (AuthController, AuthService…) + path /api/..."],
            ["Backend không chạy", "Chưa chạy test được — bật Spring Boot rồi làm tiếp"],
            ["Quên mật khẩu demo", "Password123!"],
        ],
    )

    # ===== 15 =====
    h(doc, "15. Tài liệu kèm (tuỳ chọn)", 1)
    bullet(doc, "Bo_TEST_CASE_tung_module_BUOC_CHI_TIET_cho_nguoi_moi.docx — bước bấm máy chi tiết từng case")
    bullet(doc, "Huong_dan_INTEGRATION_TEST_tich_hop_ham_cho_nguoi_moi.docx — giải thích đúng ý cô (IT ≠ Postman)")
    p(doc, "Khi gặp cô: ưu tiên mở Excel + giải thích theo Word này.", bold=True)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    PROJ.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out, PROJ / OUT_NAME)
    # also copy next to formatted excel if exists in downloads
    print(out)
    return out


if __name__ == "__main__":
    build()
