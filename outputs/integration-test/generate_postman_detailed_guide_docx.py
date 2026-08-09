# -*- coding: utf-8 -*-
"""Generate detailed Vietnamese Postman test guide (Word) — từng module Integration Test."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Huong_dan_POSTMAN_chi_tiet_tung_module_EnglishLab.docx"


def set_run_font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_vn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True)
    return h


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    set_run_font(run, size=11)
    return para


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    set_run_font(run, size=11)
    return para


def code_block(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=9, name="Consolas")
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Cm(0.5)
    return para


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
    doc.add_paragraph()


def case_block(doc, case_id, title, method, url, headers, body, steps, expected, grade, note=""):
    add_heading_vn(doc, f"{case_id} – {title}", 3)
    p(doc, f"Method / URL: {method} {url}", bold=True)
    if headers:
        p(doc, "Headers:")
        code_block(doc, headers)
    if body:
        p(doc, "Body (raw → JSON):")
        code_block(doc, body)
    p(doc, "Các bước trong Postman:", bold=True)
    for i, s in enumerate(steps, 1):
        numbered(doc, s) if False else bullet(doc, f"{i}. {s}")
    p(doc, f"Kết quả mong đợi: {expected}")
    p(doc, f"Cách chấm Excel: {grade}")
    if note:
        p(doc, f"Lưu ý: {note}")


# ---------------------------------------------------------------------------
# MODULE DATA — detailed Postman steps
# ---------------------------------------------------------------------------

MODULES = [
    {
        "title": "MODULE 1 – Authentication (Sheet: IT - Auth)",
        "sheet": "IT - Auth",
        "role": "Public (không token), trừ case /me",
        "account": "Tạo email mới; hoặc Learner 0386852628z@gmail.com / Password123!",
        "goal": "Đăng ký, OTP verify, login JWT, quên/đặt lại mật khẩu.",
        "prep": [
            "Không cần Authorization cho hầu hết case Auth.",
            "OTP: xem PostgreSQL bảng auth_tokens (cột token/code, type).",
            "Field OTP trên API là code (không phải otp).",
        ],
        "cases": [
            dict(
                id="IT_AUTH_01",
                title="Đăng ký tài khoản mới",
                method="POST",
                url="{{baseUrl}}/api/auth/register",
                headers="Content-Type: application/json",
                body='{\n  "email": "it.reg.demo01@englishlab-it.test",\n  "password": "Password123!",\n  "fullName": "IT Register User"\n}',
                steps=[
                    "New Request → đặt tên AUTH_01_register.",
                    "Chọn POST, URL như trên.",
                    "Body → raw → JSON, dán body (đổi email nếu đã tồn tại).",
                    "Send → xem Status.",
                ],
                expected="HTTP 200 hoặc 201; body thông báo thành công.",
                grade="Passed nếu 2xx. Failed nếu 500.",
            ),
            dict(
                id="IT_AUTH_02",
                title="Đăng ký trùng email (negative)",
                method="POST",
                url="{{baseUrl}}/api/auth/register",
                headers="Content-Type: application/json",
                body='{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!",\n  "fullName": "Dup"\n}',
                steps=["Gửi POST với email Learner đã tồn tại.", "Quan sát Status phải là lỗi client."],
                expected="HTTP 400 hoặc 409.",
                grade="Passed nếu bị từ chối. Failed nếu vẫn tạo được (2xx).",
            ),
            dict(
                id="IT_AUTH_03",
                title="Xác thực email bằng OTP",
                method="POST",
                url="{{baseUrl}}/api/auth/verify-email",
                headers="Content-Type: application/json",
                body='{\n  "email": "it.reg.demo01@englishlab-it.test",\n  "code": "<OTP_TU_DB>"\n}',
                steps=[
                    "Chạy AUTH_01 với email mới.",
                    "Mở DB: SELECT * FROM auth_tokens WHERE type LIKE '%EMAIL%' ORDER BY id DESC LIMIT 5;",
                    "Copy mã OTP/token → thay vào field code.",
                    "Send verify-email.",
                    "Thử login email đó để xác nhận.",
                ],
                expected="HTTP 200; sau đó login được.",
                grade="Passed nếu verify OK. N/A nếu không đọc được OTP/mail.",
            ),
            dict(
                id="IT_AUTH_04",
                title="OTP sai (negative)",
                method="POST",
                url="{{baseUrl}}/api/auth/verify-email",
                headers="Content-Type: application/json",
                body='{\n  "email": "0386852628z@gmail.com",\n  "code": "000000"\n}',
                steps=["Gửi code sai cố định 000000.", "Xem Status."],
                expected="HTTP 4xx (thường 400). Nếu user đã verify sẵn có thể 200/khác — ghi Note thực tế.",
                grade="Passed nếu bị từ chối đúng. Failed nếu 500.",
                note="Learner demo thường đã verify → có thể hành vi khác unverified user.",
            ),
            dict(
                id="IT_AUTH_05",
                title="Login lấy JWT + gọi /me",
                method="POST rồi GET",
                url="{{baseUrl}}/api/auth/login  →  {{baseUrl}}/api/user/me",
                headers="Login: Content-Type application/json\n/me: Authorization Bearer {{token}}",
                body='{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!"\n}',
                steps=[
                    "POST login với body trên.",
                    "Copy accessToken từ Response.",
                    "Hoặc tab Tests của Login: pm.environment.set('token', pm.response.json().accessToken);",
                    "Tạo GET {{baseUrl}}/api/user/me, Header Authorization = Bearer {{token}}.",
                    "Send /me.",
                ],
                expected="Login 200 có accessToken; /me 200, email đúng Learner.",
                grade="Passed nếu cả 2 bước OK.",
            ),
            dict(
                id="IT_AUTH_06",
                title="Sai mật khẩu (negative)",
                method="POST",
                url="{{baseUrl}}/api/auth/login",
                headers="Content-Type: application/json",
                body='{\n  "email": "0386852628z@gmail.com",\n  "password": "WrongPass999!"\n}',
                steps=["Send login với password sai."],
                expected="HTTP 401 hoặc 400; không có accessToken hợp lệ.",
                grade="Passed nếu không cấp token.",
            ),
            dict(
                id="IT_AUTH_07",
                title="Gọi /me không token (negative)",
                method="GET",
                url="{{baseUrl}}/api/user/me",
                headers="(không gắn Authorization)",
                body="",
                steps=["Tạo GET /me, xóa hết Header Authorization.", "Send."],
                expected="HTTP 401 hoặc 403.",
                grade="Passed nếu bị chặn.",
            ),
            dict(
                id="IT_AUTH_08",
                title="Quên mật khẩu",
                method="POST",
                url="{{baseUrl}}/api/auth/forgot-password",
                headers="Content-Type: application/json",
                body='{\n  "email": "0386852628z@gmail.com"\n}',
                steps=[
                    "Send forgot-password.",
                    "Nếu báo rate-limit / chờ vài giây → đợi rồi gọi lại.",
                ],
                expected="HTTP 200 (hoặc thông báo generic).",
                grade="Passed nếu không 5xx. Failed nếu 500.",
            ),
            dict(
                id="IT_AUTH_09",
                title="Đặt lại mật khẩu bằng OTP",
                method="POST",
                url="{{baseUrl}}/api/auth/reset-password",
                headers="Content-Type: application/json",
                body='{\n  "email": "0386852628z@gmail.com",\n  "code": "<OTP_PASSWORD_RESET>",\n  "newPassword": "Password123!"\n}',
                steps=[
                    "Chạy AUTH_08 trước.",
                    "DB: SELECT * FROM auth_tokens WHERE type LIKE '%PASSWORD%' ORDER BY id DESC;",
                    "Đặt newPassword = Password123! để không phá demo.",
                    "Send reset-password → login lại để xác nhận.",
                ],
                expected="HTTP 200; login mật khẩu mới được.",
                grade="Passed nếu reset + login OK. N/A nếu không lấy được OTP.",
            ),
            dict(
                id="IT_AUTH_10",
                title="Reset OTP sai (negative)",
                method="POST",
                url="{{baseUrl}}/api/auth/reset-password",
                headers="Content-Type: application/json",
                body='{\n  "email": "0386852628z@gmail.com",\n  "code": "000000",\n  "newPassword": "Password123!"\n}',
                steps=["Send với code sai."],
                expected="HTTP 400.",
                grade="Passed nếu bị từ chối.",
            ),
        ],
    },
    {
        "title": "MODULE 2 – Account Profile (Sheet: IT - User)",
        "sheet": "IT - User",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com / Password123!",
        "goal": "Xem/sửa hồ sơ, đổi mật khẩu (negative), avatar.",
        "prep": ["Login lấy {{token}} trước mọi request.", "Header: Authorization: Bearer {{token}}"],
        "cases": [
            dict(
                id="IT_USER_01",
                title="GET hồ sơ",
                method="GET",
                url="{{baseUrl}}/api/user/me",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["Send GET /me.", "Kiểm tra email, fullName, role trong JSON."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_USER_02",
                title="Cập nhật hồ sơ",
                method="PUT",
                url="{{baseUrl}}/api/user/me",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "fullName": "Learner IT Update",\n  "phoneNumber": "0901234567",\n  "targetExam": "TOEIC",\n  "targetScore": 700,\n  "studyGoal": "Integration test"\n}',
                steps=["PUT body trên.", "GET /me lại để xác nhận field đã đổi."],
                expected="HTTP 200.",
                grade="Passed nếu 200 và dữ liệu đổi.",
            ),
            dict(
                id="IT_USER_03",
                title="Đổi mật khẩu sai current (negative)",
                method="PUT",
                url="{{baseUrl}}/api/user/me/password",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "currentPassword": "WrongCurrent!",\n  "newPassword": "Password123!"\n}',
                steps=["Send với currentPassword sai."],
                expected="HTTP 400.",
                grade="Passed nếu bị từ chối.",
            ),
            dict(
                id="IT_USER_04",
                title="Upload avatar",
                method="POST",
                url="{{baseUrl}}/api/user/me/avatar",
                headers="Authorization: Bearer {{token}}\n(Body form-data, không set Content-Type thủ công)",
                body="form-data: key = file (type File) → chọn ảnh PNG/JPG < 1MB",
                steps=[
                    "Body → form-data.",
                    "Key: file, Type: File, Value: chọn ảnh thật.",
                    "Send.",
                ],
                expected="HTTP 200/201.",
                grade="Passed nếu upload OK. N/A nếu API từ chối định dạng ảnh.",
            ),
            dict(
                id="IT_USER_05",
                title="Sửa hồ sơ không token (negative)",
                method="PUT",
                url="{{baseUrl}}/api/user/me",
                headers="(không Authorization)",
                body='{"fullName":"x"}',
                steps=["Bỏ Bearer, Send PUT."],
                expected="HTTP 401/403.",
                grade="Passed nếu bị chặn.",
            ),
        ],
    },
    {
        "title": "MODULE 3 – Notifications (Sheet: IT - Notif)",
        "sheet": "IT - Notif",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com / Password123!",
        "goal": "Preference + list/đọc thông báo.",
        "prep": ["Cần {{token}} Learner."],
        "cases": [
            dict(
                id="IT_NOTIF_01",
                title="Lấy preference",
                method="GET",
                url="{{baseUrl}}/api/user/me/notification-preferences",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["Send GET."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_NOTIF_02",
                title="Tắt/bật in-app",
                method="PUT",
                url="{{baseUrl}}/api/user/me/notification-preferences",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "inAppEnabled": false,\n  "emailEnabled": true,\n  "larkEnabled": false\n}',
                steps=["PUT tắt inApp.", "PUT lại bật inAppEnabled=true để trả demo về trạng thái tốt."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_NOTIF_03",
                title="Body thiếu (negative)",
                method="PUT",
                url="{{baseUrl}}/api/user/me/notification-preferences",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body="{}",
                steps=["Send body rỗng {}."],
                expected="HTTP 400.",
                grade="Passed nếu validation lỗi.",
            ),
            dict(
                id="IT_NOTIF_04/05",
                title="List + unread count",
                method="GET",
                url="{{baseUrl}}/api/student/notifications\n{{baseUrl}}/api/student/notifications/unread-count",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=[
                    "GET list notifications.",
                    "GET unread-count.",
                    "Nếu có id: PATCH {{baseUrl}}/api/student/notifications/{id}/read",
                ],
                expected="HTTP 200.",
                grade="Passed nếu list/count 200.",
            ),
        ],
    },
    {
        "title": "MODULE 4 – Cart & Wishlist (Sheet: IT - Commerce)",
        "sheet": "IT - Commerce",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com / Password123!",
        "goal": "Giỏ hàng / wishlist.",
        "prep": ["Lấy courseId từ GET {{baseUrl}}/api/online-courses (không cần token)."],
        "cases": [
            dict(
                id="IT_COMMERCE_01",
                title="Thêm vào giỏ",
                method="POST",
                url="{{baseUrl}}/api/student/commerce/cart/{{courseId}}",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=[
                    "GET /api/online-courses → copy id khóa đầu tiên → set env courseId.",
                    "DELETE {{baseUrl}}/api/student/commerce/cart (xóa sạch).",
                    "POST /api/student/commerce/cart/{{courseId}}.",
                    "GET /api/student/commerce/cart → thấy khóa.",
                ],
                expected="Add 200; GET thấy course.",
                grade="Passed nếu thêm được.",
            ),
            dict(
                id="IT_COMMERCE_02",
                title="Wishlist → cart",
                method="POST",
                url="{{baseUrl}}/api/student/commerce/wishlist/{{courseId}}\n rồi POST .../wishlist/{{courseId}}/move-to-cart (hoặc endpoint move tương ứng)",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["POST wishlist.", "POST move-to-cart (đúng path API hiện có).", "GET cart kiểm tra."],
                expected="HTTP 200 hoặc báo đã có trong giỏ.",
                grade="Passed nếu chuyển được. N/A nếu khóa đã trong giỏ.",
            ),
            dict(
                id="IT_COMMERCE_03",
                title="Xóa giỏ",
                method="DELETE",
                url="{{baseUrl}}/api/student/commerce/cart",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["DELETE cart.", "GET cart → rỗng."],
                expected="HTTP 200/204.",
                grade="Passed nếu xóa được.",
            ),
            dict(
                id="IT_COMMERCE_04",
                title="Thêm lại vào giỏ",
                method="POST",
                url="{{baseUrl}}/api/student/commerce/cart/{{courseId}}",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["Lặp bước COMMERCE_01."],
                expected="HTTP 200.",
                grade="Passed nếu OK.",
            ),
        ],
    },
    {
        "title": "MODULE 5 – PayOS & Orders (Sheet: IT - Payment)",
        "sheet": "IT - Payment",
        "role": "LEARNER + MANAGER",
        "account": "Learner + classroom.manager@englishlab.vn / Password123!",
        "goal": "Quote, link PayOS, webhook negative, orders manager.",
        "prep": ["Cần courseId public.", "Không thanh toán thật trên môi trường demo nếu không cần."],
        "cases": [
            dict(
                id="IT_PAYMENT_01",
                title="Tạo PayOS link",
                method="POST",
                url="{{baseUrl}}/api/student/payments/payos/link",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "courseIds": [1]\n}',
                steps=["Thay 1 bằng courseId thật.", "Send → tìm checkoutUrl / orderCode trong Response."],
                expected="HTTP 200 có checkoutUrl hoặc orderCode.",
                grade="Passed nếu tạo link. Failed nếu 500.",
            ),
            dict(
                id="IT_PAYMENT_02",
                title="Quote giá",
                method="POST",
                url="{{baseUrl}}/api/student/payments/quote",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "courseIds": [1]\n}',
                steps=["Send quote.", "Kiểm tra totalAmount / breakdown."],
                expected="HTTP 200 có totalAmount.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_PAYMENT_03",
                title="Webhook thiếu chữ ký (negative)",
                method="POST",
                url="{{baseUrl}}/api/payos/webhook",
                headers="Content-Type: application/json",
                body="{}",
                steps=["Gửi body giả, không chữ ký PayOS."],
                expected="HTTP 400 (từ chối).",
                grade="Passed nếu bị từ chối. Failed nếu 404.",
            ),
            dict(
                id="IT_PAYMENT_04/05",
                title="Manager xem orders",
                method="GET",
                url="{{baseUrl}}/api/manager/payments/orders",
                headers="Authorization: Bearer {{managerToken}}",
                body="",
                steps=["Login MANAGER → lưu managerToken.", "GET orders."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 6 – Online Learning (Sheet: IT - Course)",
        "sheet": "IT - Course",
        "role": "Public + LEARNER",
        "account": "Public không token / Learner có token",
        "goal": "Catalog công khai; content/progress cần enroll.",
        "prep": ["Lấy slug/id từ list public."],
        "cases": [
            dict(
                id="IT_COURSE_01",
                title="Danh sách khóa public",
                method="GET",
                url="{{baseUrl}}/api/online-courses",
                headers="(không token)",
                body="",
                steps=["Send GET không Authorization."],
                expected="HTTP 200 + mảng khóa.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_COURSE_02",
                title="Chi tiết khóa",
                method="GET",
                url="{{baseUrl}}/api/online-courses/{{slugOrId}}",
                headers="(không token)",
                body="",
                steps=["Thay slugOrId từ COURSE_01.", "Send."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_COURSE_03/06",
                title="Xem content learner",
                method="GET",
                url="{{baseUrl}}/api/student/online-courses/{{courseId}}/content",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["Send với token Learner."],
                expected="200 nếu đã enroll; 400 nếu chưa.",
                grade="Passed nếu đã enroll + 200. N/A nếu 400 chưa enroll (đúng business).",
            ),
            dict(
                id="IT_COURSE_04",
                title="Cập nhật progress",
                method="PATCH",
                url="{{baseUrl}}/api/student/online-courses/{{courseId}}/lessons/{{lessonId}}/progress",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "completed": true\n}',
                steps=["Lấy lessonId từ content.", "PATCH progress."],
                expected="HTTP 200 nếu đủ quyền.",
                grade="Passed/N/A tùy enroll.",
            ),
            dict(
                id="IT_COURSE_05",
                title="Rating khóa",
                method="POST",
                url="{{baseUrl}}/api/student/online-courses/{{courseId}}/rating",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "score": 5,\n  "comment": "ok"\n}',
                steps=["POST rating."],
                expected="HTTP 200/201.",
                grade="Passed nếu gửi được. N/A nếu thiếu điều kiện.",
            ),
        ],
    },
    {
        "title": "MODULE 7 – Course Discussion (Sheet: IT - Discuss)",
        "sheet": "IT - Discuss",
        "role": "LEARNER + CONTENT_MANAGER",
        "account": "Learner + content.manager@englishlab.vn",
        "goal": "Thảo luận / report / moderation.",
        "prep": ["Một số case cần đã enroll khóa."],
        "cases": [
            dict(
                id="IT_DISCUSS_01",
                title="Tạo thảo luận",
                method="POST",
                url="{{baseUrl}}/api/student/online-courses/{{courseId}}/discussions",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "title": "IT discuss title",\n  "content": "Noi dung thao luan integration test"\n}',
                steps=["POST với token Learner đã enroll."],
                expected="200 nếu đã enroll.",
                grade="Passed nếu tạo được. N/A nếu 400 chưa enroll.",
            ),
            dict(
                id="IT_DISCUSS_02",
                title="List thảo luận",
                method="GET",
                url="{{baseUrl}}/api/online-courses/{{courseId}}/discussions",
                headers="(có thể không token)",
                body="",
                steps=["GET list."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_DISCUSS_03/04",
                title="Report thread",
                method="POST",
                url="{{baseUrl}}/api/student/online-courses/discussions/{{threadId}}/reports",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "reason": "Spam test IT"\n}',
                steps=["Lấy threadId từ list.", "POST report."],
                expected="HTTP 200.",
                grade="Passed nếu report được. N/A nếu không có thread.",
            ),
            dict(
                id="IT_DISCUSS_05",
                title="CM moderation",
                method="GET",
                url="{{baseUrl}}/api/content-manager/discussion-reports",
                headers="Authorization: Bearer {{cmToken}}",
                body="",
                steps=["Login CM → cmToken.", "GET reports."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 8 – CM Online Courses (Sheet: IT - Content)",
        "sheet": "IT - Content",
        "role": "CONTENT_MANAGER",
        "account": "content.manager@englishlab.vn / Password123!",
        "goal": "CM quản lý khóa online.",
        "prep": ["Login CM → {{cmToken}}"],
        "cases": [
            dict(
                id="IT_CONTENT_01..04",
                title="List khóa CM",
                method="GET",
                url="{{baseUrl}}/api/content-manager/online-courses",
                headers="Authorization: Bearer {{cmToken}}",
                body="",
                steps=["GET list.", "Ghi nhận có dữ liệu / rỗng vẫn 200 là OK."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
                note="Create/publish sâu: làm thêm POST/PUT theo Swagger nếu cô yêu cầu.",
            ),
        ],
    },
    {
        "title": "MODULE 9 – Packages & Bundles (Sheet: IT - Package)",
        "sheet": "IT - Package",
        "role": "CONTENT_MANAGER",
        "account": "content.manager@englishlab.vn",
        "goal": "List package/bundle.",
        "prep": ["Token CM."],
        "cases": [
            dict(
                id="IT_PACKAGE_01..03",
                title="List packages",
                method="GET",
                url="{{baseUrl}}/api/content-manager/packages",
                headers="Authorization: Bearer {{cmToken}}",
                body="",
                steps=["GET packages."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 10 – Curriculum & Banks (Sheet: IT - Curriculum)",
        "sheet": "IT - Curriculum",
        "role": "CONTENT_MANAGER",
        "account": "content.manager@englishlab.vn",
        "goal": "Chương trình, ngân hàng, rubric.",
        "prep": ["Token CM."],
        "cases": [
            dict(
                id="IT_CURRICULUM_01/05",
                title="Curriculum programs",
                method="GET",
                url="{{baseUrl}}/api/content-manager/curriculum-programs",
                headers="Authorization: Bearer {{cmToken}}",
                body="",
                steps=["GET programs."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_CURRICULUM_02",
                title="Exercise / Assessment bank",
                method="GET",
                url="{{baseUrl}}/api/content-manager/exercise-bank\n(hoặc /api/content-manager/assessment-bank)",
                headers="Authorization: Bearer {{cmToken}}",
                body="",
                steps=["Thử cả 2 path nếu một path 404."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_CURRICULUM_03",
                title="Learning paths",
                method="GET",
                url="{{baseUrl}}/api/content-manager/learning-paths",
                headers="Authorization: Bearer {{cmToken}}",
                body="",
                steps=["GET learning-paths."],
                expected="HTTP 200 hoặc N/A nếu chưa có API.",
                grade="Ghi đúng thực tế.",
            ),
            dict(
                id="IT_CURRICULUM_04",
                title="Rubrics",
                method="GET",
                url="{{baseUrl}}/api/content-manager/rubrics",
                headers="Authorization: Bearer {{cmToken}}",
                body="",
                steps=["GET rubrics."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 11 – Enrollment Requests (Sheet: IT - EnrollReq)",
        "sheet": "IT - EnrollReq",
        "role": "LEARNER + STAFF",
        "account": "Learner + staff@englishlab.vn",
        "goal": "HV gửi form tư vấn; Staff xem.",
        "prep": ["GET /api/course-offerings lấy courseOfferingId."],
        "cases": [
            dict(
                id="IT_ENROLLREQ_01/04",
                title="HV tạo request",
                method="POST",
                url="{{baseUrl}}/api/student/course-enrollment-requests",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "courseOfferingId": 1,\n  "contactName": "HV IT",\n  "contactEmail": "0386852628z@gmail.com",\n  "contactPhone": "0901111222",\n  "consultationTrack": "TOEIC"\n}',
                steps=[
                    "GET course-offerings → id.",
                    "POST enrollment request.",
                    "GET {{baseUrl}}/api/student/course-enrollment-requests/my",
                ],
                expected="200 nếu tạo mới được.",
                grade="Passed nếu tạo + /my OK. N/A nếu đã có form đang xử lý.",
            ),
            dict(
                id="IT_ENROLLREQ_02/03/05",
                title="Staff list",
                method="GET",
                url="{{baseUrl}}/api/staff/enrollment-requests",
                headers="Authorization: Bearer {{staffToken}}",
                body="",
                steps=["Login STAFF.", "GET list."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 12 – TM Classroom Ops (Sheet: IT - Classroom)",
        "sheet": "IT - Classroom",
        "role": "TRAINING_MANAGER (+ public)",
        "account": "training.manager@englishlab.vn / Password123!",
        "goal": "Offerings, registrations, waitlist.",
        "prep": ["Login TM → {{tmToken}}"],
        "cases": [
            dict(
                id="IT_CLASS_01",
                title="Public offerings",
                method="GET",
                url="{{baseUrl}}/api/classroom-offerings",
                headers="(không token)",
                body="",
                steps=["GET public."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_CLASS_02/08",
                title="TM list lớp",
                method="GET",
                url="{{baseUrl}}/api/training-manager/classrooms",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["GET list.", "Copy 1 id lớp → env classroomId."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_CLASS_03",
                title="TM chi tiết lớp",
                method="GET",
                url="{{baseUrl}}/api/training-manager/classrooms/{{classroomId}}",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["GET detail."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_CLASS_04/06",
                title="Registrations",
                method="GET",
                url="{{baseUrl}}/api/training-manager/classrooms/registrations?classroomOfferingId={{classroomId}}",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["GET registrations."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_CLASS_05",
                title="Reorder waitlist",
                method="PUT",
                url="{{baseUrl}}/api/training-manager/classrooms/{{classroomId}}/waitlist/order",
                headers="Authorization: Bearer {{tmToken}}\nContent-Type: application/json",
                body='{\n  "enrollmentIds": [101, 102]\n}',
                steps=[
                    "GET registrations?status=WAITLIST&classroomOfferingId=...",
                    "Nếu ≥ 2 HV: đổi thứ tự enrollmentIds rồi PUT.",
                    "Nếu < 2 → ghi N/A.",
                ],
                expected="HTTP 200 nếu đủ data.",
                grade="Passed nếu reorder OK. N/A nếu <2 HV waitlist hoặc endpoint chưa có trên bản code hiện tại.",
                note="Nếu PUT trả 404: ghi N/A + Note endpoint chưa deploy — không Failed hệ thống nếu chưa có quyền sửa code.",
            ),
            dict(
                id="IT_CLASS_07",
                title="Xem lớp trước khi gán GV",
                method="GET",
                url="{{baseUrl}}/api/training-manager/classrooms/{{classroomId}}",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["GET detail lớp bất kỳ."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 13 – Learner Classroom (Sheet: IT - LearnerCls)",
        "sheet": "IT - LearnerCls",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com (nên đã được gán lớp)",
        "goal": "Lớp của tôi, session, homework, materials, gradebook.",
        "prep": ["Token Learner.", "Nếu my-classrooms = [] → case sau ghi N/A."],
        "cases": [
            dict(
                id="IT_LEARNERCLS_01",
                title="My classrooms",
                method="GET",
                url="{{baseUrl}}/api/student/classrooms/my-classrooms",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET my-classrooms.", "Copy id nếu có."],
                expected="HTTP 200 (có thể []).",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_LEARNERCLS_02",
                title="Sessions",
                method="GET",
                url="{{baseUrl}}/api/student/classrooms/{{classroomId}}/sessions",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET sessions."],
                expected="HTTP 200.",
                grade="Passed nếu 200. N/A nếu không có lớp.",
            ),
            dict(
                id="IT_LEARNERCLS_03/05",
                title="Homework",
                method="GET",
                url="{{baseUrl}}/api/student/classrooms/{{classroomId}}/homework",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET homework."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_LEARNERCLS_04",
                title="Materials",
                method="GET",
                url="{{baseUrl}}/api/student/classrooms/{{classroomId}}/materials",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET materials."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_LEARNERCLS_06",
                title="Gradebook của tôi",
                method="GET",
                url="{{baseUrl}}/api/student/classrooms/{{classroomId}}/gradebook/me",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET gradebook/me."],
                expected="HTTP 200 hoặc 204.",
                grade="Passed nếu 200/204.",
            ),
        ],
    },
    {
        "title": "MODULE 14 – Teacher Operations (Sheet: IT - Teacher)",
        "sheet": "IT - Teacher",
        "role": "TEACHER",
        "account": "classroom.teacher1@englishlab.vn / Password123!",
        "goal": "Lớp assigned, homework, điểm danh, gradebook, change request.",
        "prep": ["Login Teacher → {{teacherToken}}"],
        "cases": [
            dict(
                id="IT_TEACH_01/06",
                title="Lớp assigned",
                method="GET",
                url="{{baseUrl}}/api/teacher/classrooms/assigned",
                headers="Authorization: Bearer {{teacherToken}}",
                body="",
                steps=["GET assigned.", "Copy classroomId."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_TEACH_02",
                title="Homework",
                method="GET",
                url="{{baseUrl}}/api/teacher/classrooms/{{classroomId}}/homework",
                headers="Authorization: Bearer {{teacherToken}}",
                body="",
                steps=["GET homework."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_TEACH_03",
                title="Điểm danh theo session",
                method="GET",
                url="{{baseUrl}}/api/teacher/classrooms/sessions/{{sessionId}}/attendance",
                headers="Authorization: Bearer {{teacherToken}}",
                body="",
                steps=[
                    "GET {{baseUrl}}/api/teacher/classrooms/{{classroomId}}/sessions",
                    "Copy sessionId → GET attendance.",
                ],
                expected="HTTP 200.",
                grade="Passed nếu 200. N/A nếu chưa có session.",
            ),
            dict(
                id="IT_TEACH_04",
                title="Gradebook lớp",
                method="GET",
                url="{{baseUrl}}/api/teacher/classrooms/{{classroomId}}/gradebook",
                headers="Authorization: Bearer {{teacherToken}}",
                body="",
                steps=["GET gradebook."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_TEACH_05",
                title="Change requests của tôi",
                method="GET",
                url="{{baseUrl}}/api/teacher/classrooms/requests/mine",
                headers="Authorization: Bearer {{teacherToken}}",
                body="",
                steps=["GET requests/mine."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 15 – Classroom Quiz (Sheet: IT - Quiz)",
        "sheet": "IT - Quiz",
        "role": "TEACHER + LEARNER",
        "account": "Teacher + Learner",
        "goal": "List quiz; tránh xóa data demo.",
        "prep": ["Cần offeringId từ lớp teacher."],
        "cases": [
            dict(
                id="IT_QUIZ_01/02",
                title="Teacher list quiz",
                method="GET",
                url="{{baseUrl}}/api/teacher/classrooms/{{offeringId}}/quizzes",
                headers="Authorization: Bearer {{teacherToken}}",
                body="",
                steps=["GET quizzes."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_QUIZ_03",
                title="Learner list quiz",
                method="GET",
                url="{{baseUrl}}/api/student/classrooms/quizzes",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET learner quizzes."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_QUIZ_04",
                title="Xóa quiz (destructive)",
                method="DELETE",
                url="{{baseUrl}}/api/teacher/quizzes/{{quizId}}",
                headers="Authorization: Bearer {{teacherToken}}",
                body="",
                steps=["Chỉ xóa quiz tự tạo để test.", "Không xóa quiz demo."],
                expected="HTTP 204 nếu xóa quiz test.",
                grade="N/A nếu không muốn xóa data demo.",
            ),
        ],
    },
    {
        "title": "MODULE 16 – Assessment & Placement (Sheet: IT - Assess)",
        "sheet": "IT - Assess",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com",
        "goal": "Placement, mock test; assessment khóa cần enroll.",
        "prep": ["Token Learner."],
        "cases": [
            dict(
                id="IT_ASSESS_01",
                title="Placement hiện tại",
                method="GET",
                url="{{baseUrl}}/api/student/placement-tests/current",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET current placement."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_ASSESS_02",
                title="Submit placement thiếu đáp án (negative)",
                method="POST",
                url="{{baseUrl}}/api/student/placement-tests/current/submit",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "answers": []\n}',
                steps=["POST answers rỗng."],
                expected="HTTP 400 hoặc 200 tùy trạng thái (ghi Note).",
                grade="Passed nếu hành vi hợp lý. Failed nếu 500.",
            ),
            dict(
                id="IT_ASSESS_03/05",
                title="Assessments theo khóa",
                method="GET",
                url="{{baseUrl}}/api/student/online-courses/{{courseId}}/assessments",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET assessments."],
                expected="200 nếu enroll.",
                grade="Passed nếu 200. N/A nếu 400 chưa enroll.",
            ),
            dict(
                id="IT_ASSESS_04/06",
                title="Mock tests",
                method="GET",
                url="{{baseUrl}}/api/student/mock-tests",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET mock-tests."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 17 – Support Tickets (Sheet: IT - Support)",
        "sheet": "IT - Support",
        "role": "LEARNER + MANAGER",
        "account": "Learner + classroom.manager@englishlab.vn",
        "goal": "Tạo/list ticket; manager xem.",
        "prep": ["subject ≥ 5 ký tự; message ≥ 10 ký tự."],
        "cases": [
            dict(
                id="IT_SUPPORT_01",
                title="Tạo ticket",
                method="POST",
                url="{{baseUrl}}/api/student/support-tickets",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body='{\n  "subject": "IT Support ticket",\n  "category": "TECHNICAL",\n  "message": "Mo ta chi tiet loi integration test"\n}',
                steps=["POST ticket.", "Nhớ id nếu có."],
                expected="HTTP 200/201.",
                grade="Passed nếu tạo được.",
            ),
            dict(
                id="IT_SUPPORT_02",
                title="List của tôi",
                method="GET",
                url="{{baseUrl}}/api/student/support-tickets",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET list."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_SUPPORT_03",
                title="Manager list",
                method="GET",
                url="{{baseUrl}}/api/manager/support-tickets",
                headers="Authorization: Bearer {{managerToken}}",
                body="",
                steps=["Login MANAGER.", "GET tickets."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_SUPPORT_04",
                title="Body rỗng (negative)",
                method="POST",
                url="{{baseUrl}}/api/student/support-tickets",
                headers="Authorization: Bearer {{token}}\nContent-Type: application/json",
                body="{}",
                steps=["POST {}."],
                expected="HTTP 400.",
                grade="Passed nếu validation lỗi.",
            ),
        ],
    },
    {
        "title": "MODULE 18 – Administration (Sheet: IT - Admin)",
        "sheet": "IT - Admin",
        "role": "ADMIN",
        "account": "classroom.admin@englishlab.vn / Password123!",
        "goal": "Users, audit, config.",
        "prep": ["Login ADMIN → {{adminToken}}"],
        "cases": [
            dict(
                id="IT_ADMIN_01/02",
                title="List users",
                method="GET",
                url="{{baseUrl}}/api/admin/users",
                headers="Authorization: Bearer {{adminToken}}",
                body="",
                steps=["GET users."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_ADMIN_03",
                title="Audit logs",
                method="GET",
                url="{{baseUrl}}/api/admin/audit-logs",
                headers="Authorization: Bearer {{adminToken}}",
                body="",
                steps=["GET audit-logs."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_ADMIN_04",
                title="System config",
                method="GET",
                url="{{baseUrl}}/api/admin/system/config",
                headers="Authorization: Bearer {{adminToken}}",
                body="",
                steps=["GET config."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 19 – Lark Meetings (Sheet: IT - Lark)",
        "sheet": "IT - Lark",
        "role": "Public webhook + TM",
        "account": "TM cho sync",
        "goal": "Webhook challenge; sync recording.",
        "prep": ["Lark phụ thuộc cấu hình môi trường → nhiều case có thể N/A."],
        "cases": [
            dict(
                id="IT_LARK_01/02",
                title="Webhook challenge",
                method="POST",
                url="{{baseUrl}}/api/lark/events",
                headers="Content-Type: application/json",
                body='{\n  "type": "url_verification",\n  "challenge": "it-challenge-123"\n}',
                steps=["POST body challenge.", "Xem có echo challenge không."],
                expected="200 nếu format đúng; có thể 400 nếu thiếu chữ ký.",
                grade="Passed nếu 200. N/A nếu 400 do cấu hình Lark.",
            ),
            dict(
                id="IT_LARK_03",
                title="Sync Lark session",
                method="POST",
                url="{{baseUrl}}/api/training-manager/recordings/sessions/{{sessionId}}/sync-lark",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["Dùng sessionId thật từ TM.", "POST sync."],
                expected="200 nếu session hợp lệ.",
                grade="N/A nếu session không tồn tại / chưa cấu hình Lark.",
            ),
        ],
    },
    {
        "title": "MODULE 20 – Infrastructure (Sheet: IT - Infra)",
        "sheet": "IT - Infra",
        "role": "TRAINING_MANAGER",
        "account": "training.manager@englishlab.vn",
        "goal": "Campus, room, session template.",
        "prep": ["{{tmToken}}"],
        "cases": [
            dict(
                id="IT_INFRA_01",
                title="Campuses",
                method="GET",
                url="{{baseUrl}}/api/training-manager/infrastructure/campuses",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["GET campuses."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_INFRA_02",
                title="Rooms",
                method="GET",
                url="{{baseUrl}}/api/training-manager/infrastructure/rooms",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["GET rooms."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_INFRA_03",
                title="Session templates",
                method="GET",
                url="{{baseUrl}}/api/training-manager/infrastructure/session-templates",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["GET session-templates."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 21 – Reports & Revenue (Sheet: IT - Report)",
        "sheet": "IT - Report",
        "role": "TM + CM",
        "account": "training.manager@englishlab.vn + content.manager@englishlab.vn",
        "goal": "Dashboard / doanh thu.",
        "prep": ["Hai token riêng: tmToken, cmToken."],
        "cases": [
            dict(
                id="IT_REPORT_01",
                title="TM dashboard",
                method="GET",
                url="{{baseUrl}}/api/training-manager/dashboard",
                headers="Authorization: Bearer {{tmToken}}",
                body="",
                steps=["GET dashboard."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_REPORT_02",
                title="Revenue analytics",
                method="GET",
                url="{{baseUrl}}/api/content-manager/revenue/analytics",
                headers="Authorization: Bearer {{cmToken}}",
                body="",
                steps=["GET revenue analytics."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 22 – Classroom Proposals (Sheet: IT - Proposal)",
        "sheet": "IT - Proposal",
        "role": "STAFF",
        "account": "staff@englishlab.vn",
        "goal": "Đề xuất mở lớp.",
        "prep": ["{{staffToken}}"],
        "cases": [
            dict(
                id="IT_PROPOSAL_01..03",
                title="List proposals",
                method="GET",
                url="{{baseUrl}}/api/staff/classroom-proposals",
                headers="Authorization: Bearer {{staffToken}}",
                body="",
                steps=["GET proposals."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 23 – Attendance Disputes (Sheet: IT - Dispute)",
        "sheet": "IT - Dispute",
        "role": "LEARNER + TEACHER",
        "account": "Learner + Teacher",
        "goal": "Khiếu nại điểm danh.",
        "prep": ["Hai token."],
        "cases": [
            dict(
                id="IT_DISPUTE_01",
                title="HV xem disputes",
                method="GET",
                url="{{baseUrl}}/api/student/attendance/disputes",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET disputes (có thể [])."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
            dict(
                id="IT_DISPUTE_02/03",
                title="GV pending disputes",
                method="GET",
                url="{{baseUrl}}/api/teacher/attendance-disputes/pending",
                headers="Authorization: Bearer {{teacherToken}}",
                body="",
                steps=["GET pending."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
    {
        "title": "MODULE 24 – Learning Notes (Sheet: IT - Notes)",
        "sheet": "IT - Notes",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com",
        "goal": "Ghi chú bài học.",
        "prep": ["{{token}}"],
        "cases": [
            dict(
                id="IT_NOTES_01/02",
                title="List notes",
                method="GET",
                url="{{baseUrl}}/api/student/learning/notes",
                headers="Authorization: Bearer {{token}}",
                body="",
                steps=["GET notes."],
                expected="HTTP 200.",
                grade="Passed nếu 200.",
            ),
        ],
    },
]


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "HƯỚNG DẪN TEST POSTMAN CHI TIẾT\n"
        "TỪNG MODULE INTEGRATION TEST\n"
        "DỰ ÁN ENGLISHLAB (SEP490_G23)"
    )
    set_run_font(r, size=18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Tài liệu tự thực hiện bằng Postman → ghi kết quả vào Excel Integration Test")
    set_run_font(r, size=12)

    p(doc, "Người thực hiện / Tester: phongdx", bold=True)
    p(
        doc,
        "Môi trường: Backend http://localhost:8080 · PostgreSQL đã seed · "
        "File Excel: SEP490_G23_Report5.2_Integration Test_COMPLETED_HONEST_FORMATTED.xlsx",
    )

    add_heading_vn(doc, "PHẦN A – Mục đích", 1)
    p(
        doc,
        "Tài liệu hướng dẫn chi tiết cách dùng Postman để tự test 24 module Integration Test. "
        "Mỗi case có Method, URL, Headers, Body mẫu, từng bước bấm trong Postman, kết quả mong đợi và cách ghi Excel.",
    )

    add_heading_vn(doc, "PHẦN B – Chuẩn bị hệ thống", 1)
    numbered(doc, "Bật PostgreSQL (database englishlab) và seed demo.")
    numbered(doc, "Chạy backend Spring Boot cổng 8080.")
    numbered(doc, "Mở Swagger (tuỳ chọn): http://localhost:8080/swagger-ui.html để đối chiếu path.")
    numbered(doc, "Cài Postman (Desktop).")
    numbered(doc, "Mở file Excel Integration Test bản HONEST để ghi Round 1–3.")

    add_heading_vn(doc, "PHẦN C – Tạo Environment Postman (làm 1 lần)", 1)
    numbered(doc, "Postman → Environments → + → đặt tên EnglishLab-Local.")
    numbered(doc, "Thêm biến (TYPE = default):")
    add_table(
        doc,
        ["Biến", "Initial value", "Công dụng"],
        [
            ["baseUrl", "http://localhost:8080", "Prefix mọi API"],
            ["token", "(để trống)", "JWT Learner"],
            ["teacherToken", "(để trống)", "JWT Teacher"],
            ["tmToken", "(để trống)", "JWT Training Manager"],
            ["staffToken", "(để trống)", "JWT Staff"],
            ["managerToken", "(để trống)", "JWT Manager"],
            ["cmToken", "(để trống)", "JWT Content Manager"],
            ["adminToken", "(để trống)", "JWT Admin"],
            ["courseId", "(điền sau)", "Id khóa online"],
            ["classroomId", "(điền sau)", "Id lớp"],
            ["sessionId", "(điền sau)", "Id buổi học"],
        ],
    )
    numbered(doc, "Góc trên-phải Postman chọn Environment = EnglishLab-Local.")
    numbered(doc, "URL viết dạng {{baseUrl}}/api/...")

    add_heading_vn(doc, "PHẦN D – Request Login chuẩn (copy dùng lại)", 1)
    p(doc, "Tạo folder Auth → request LOGIN_LEARNER:", bold=True)
    p(doc, "POST {{baseUrl}}/api/auth/login")
    p(doc, "Body raw JSON:")
    code_block(
        doc,
        '{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!"\n}',
    )
    p(doc, "Tab Tests (JavaScript) — tự lưu token:", bold=True)
    code_block(
        doc,
        'if (pm.response.code === 200) {\n'
        '  const json = pm.response.json();\n'
        '  pm.environment.set("token", json.accessToken);\n'
        '}',
    )
    p(doc, "Tạo thêm các request Login khác, đổi email + đổi dòng set:", bold=True)
    bullet(doc, 'Teacher → pm.environment.set("teacherToken", json.accessToken);')
    bullet(doc, 'TM → pm.environment.set("tmToken", json.accessToken);')
    bullet(doc, 'Staff → pm.environment.set("staffToken", json.accessToken);')
    bullet(doc, 'Manager → pm.environment.set("managerToken", json.accessToken);')
    bullet(doc, 'CM → pm.environment.set("cmToken", json.accessToken);')
    bullet(doc, 'Admin → pm.environment.set("adminToken", json.accessToken);')

    add_heading_vn(doc, "PHẦN E – Gắn Bearer token cho request sau", 1)
    numbered(doc, "Mở request cần bảo vệ → tab Authorization.")
    numbered(doc, "Type = Bearer Token.")
    numbered(doc, "Token = {{token}} (hoặc {{teacherToken}}, {{tmToken}}, …).")
    numbered(doc, "Hoặc tab Headers: Key=Authorization, Value=Bearer {{token}}.")

    add_heading_vn(doc, "PHẦN F – Quy trình 1 case + cách ghi Excel", 1)
    numbered(doc, "Mở đúng sheet module trong Excel (cột Test Case ID = IT_xxx).")
    numbered(doc, "Đọc Pre-conditions → login đúng role.")
    numbered(doc, "Làm theo từng bước trong tài liệu này.")
    numbered(doc, "So Status + body với Expected.")
    numbered(doc, "Ghi Excel: Round = Passed/Failed/N/A · Test date · Tester=phongdx · Note=HTTP xxx hoặc lý do N/A.")

    add_table(
        doc,
        ["Kết quả", "Khi nào ghi"],
        [
            ["Passed", "Đúng expected: happy-path OK hoặc negative bị từ chối đúng (4xx)."],
            ["Failed", "Sai expected, lỗi 500, path 404 không tồn tại khi đáng lẽ có."],
            ["N/A", "Thiếu precondition (chưa enroll, thiếu OTP, waitlist < 2, không xóa data demo, Lark chưa cấu hình…)."],
        ],
    )

    add_heading_vn(doc, "PHẦN G – Tài khoản demo (Password123!)", 1)
    add_table(
        doc,
        ["Role", "Email"],
        [
            ["LEARNER", "0386852628z@gmail.com"],
            ["TEACHER", "classroom.teacher1@englishlab.vn"],
            ["TRAINING_MANAGER", "training.manager@englishlab.vn"],
            ["STAFF", "staff@englishlab.vn"],
            ["MANAGER", "classroom.manager@englishlab.vn"],
            ["CONTENT_MANAGER", "content.manager@englishlab.vn"],
            ["ADMIN", "classroom.admin@englishlab.vn"],
        ],
    )

    add_heading_vn(doc, "PHẦN H – OTP từ database (verify / reset)", 1)
    p(doc, "Khi không đọc được email, dùng SQL:")
    code_block(
        doc,
        "SELECT id, user_id, type, token, code, expires_at, created_at\n"
        "FROM auth_tokens\n"
        "ORDER BY id DESC\n"
        "LIMIT 20;",
    )
    bullet(doc, "API verify-email / reset-password dùng field code (không dùng otp).")
    bullet(doc, "Sau reset nên đặt lại Password123! để không phá tài khoản demo.")

    add_heading_vn(doc, "PHẦN I – Hướng dẫn chi tiết TỪNG MODULE", 1)
    p(
        doc,
        "Có 24 module. Nên làm theo thứ tự: Auth → User → Notif → Classroom → Teacher → Commerce → Payment → các module còn lại.",
    )

    for m in MODULES:
        add_heading_vn(doc, m["title"], 2)
        p(doc, f"Sheet Excel: {m['sheet']}", bold=True)
        p(doc, f"Vai trò: {m['role']}")
        p(doc, f"Tài khoản: {m['account']}")
        p(doc, f"Mục tiêu: {m['goal']}")
        p(doc, "Chuẩn bị trước module:", bold=True)
        for item in m["prep"]:
            bullet(doc, item)
        for c in m["cases"]:
            case_block(
                doc,
                c["id"],
                c["title"],
                c["method"],
                c["url"],
                c["headers"],
                c["body"],
                c["steps"],
                c["expected"],
                c["grade"],
                c.get("note", ""),
            )

    add_heading_vn(doc, "PHẦN J – Cấu trúc Collection Postman gợi ý", 1)
    bullet(doc, "00_Auth_Login (login theo từng role + Tests script)")
    bullet(doc, "01_Auth … 24_Notes — mỗi folder = 1 module; mỗi request đặt tên đúng mã IT_xxx")
    bullet(doc, "Sau mỗi Send: chụp Status (tuỳ chọn) và ghi Excel ngay.")

    add_heading_vn(doc, "PHẦN K – Checklist trước khi nộp / thuyết trình", 1)
    bullet(doc, "Đã tự chạy Postman và ghi Round cho các module cô yêu cầu demo.")
    bullet(doc, "Tester = tên thật.")
    bullet(doc, "Note có HTTP status hoặc lý do N/A.")
    bullet(doc, "Biết giải thích 1 case Passed và 1 case N/A.")
    bullet(doc, "Không thanh toán PayOS thật / không xóa data demo quan trọng.")

    add_heading_vn(doc, "PHẦN L – Câu trả lời ngắn nếu cô hỏi", 1)
    p(
        doc,
        "Em tự test bằng Postman theo đúng mã IT trên từng sheet Excel: tạo Environment baseUrl, "
        "login lấy JWT Bearer, gửi request đúng Method/Body, đối chiếu Expected Results rồi ghi "
        "Passed / Failed / N/A vào Round 1–3. N/A là thiếu precondition trên môi trường demo, không phải giấu Fail.",
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUT_DIR / OUT_NAME
    doc.save(out_path)
    PROJ.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out_path, PROJ / OUT_NAME)
    print(f"DOCX {out_path}")
    print(f"COPY {PROJ / OUT_NAME}")
    return out_path


if __name__ == "__main__":
    build()
