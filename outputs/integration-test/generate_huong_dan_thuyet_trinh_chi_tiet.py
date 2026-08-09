# -*- coding: utf-8 -*-
"""
Word huong dan thuyet trinh CHI TIET theo tung cot Excel:
Test Case ID -> Description -> Pre-conditions -> Procedure (tung buoc) -> Expected Results.
Bo AUTH. Lay noi dung tu uc_modules.py + dich de doc.
"""
from __future__ import annotations

import hashlib
import json
import re
import time
from pathlib import Path

from deep_translator import GoogleTranslator
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))
from uc_modules import MODULES  # noqa: E402

OUT = Path(__file__).resolve().parent / "Huong_dan_THUYET_TRINH_CHI_TIET_tung_cot_Excel.docx"
OUT_DL = Path(
    r"C:\Users\phong\Downloads\intergration test"
    r"\Huong_dan_THUYET_TRINH_CHI_TIET_tung_cot_Excel.docx"
)
CACHE = Path(__file__).resolve().parent / "_vi_speak_cache.json"

PRIORITY = [
    "COURSE",
    "ENROLL",
    "CHECKOUT",
    "CLASS",
    "ASSIGN",
    "ASNTEACH",
    "SCHEDULE",
    "ATTEND",
    "MNGHW",
    "HOMEWORK",
    "QUIZ",
    "ONLINE",
    "ADMIN",
    "BROADCAST",
]

FEATURE_VI = {
    "View public courses": "Xem khóa học công khai",
    "Enroll in Course": "Đăng ký khóa học",
    "Checkout": "Thanh toán (Checkout)",
    "Manage Classrooms": "Quản lý lớp học",
    "Assign Learner to Classroom": "Xếp học viên vào lớp",
    "Assign Teacher to Classroom": "Phân công giáo viên vào lớp",
    "View Teaching Schedule": "Xem lịch giảng dạy",
    "Manage Class Attendance": "Quản lý điểm danh lớp",
    "Manage Homework": "Quản lý bài tập (giáo viên)",
    "Submit Homework": "Nộp bài tập",
    "Take Quiz": "Làm bài quiz",
    "Manage Online Courses": "Quản lý khóa học online",
    "Manage User Accounts": "Quản lý tài khoản người dùng",
    "Manage System Notifications": "Quản lý thông báo hệ thống",
}

translator = GoogleTranslator(source="en", target="vi")
_cache: dict[str, str] = {}


def load_cache():
    global _cache
    if CACHE.exists():
        _cache = json.loads(CACHE.read_text(encoding="utf-8"))


def save_cache():
    CACHE.write_text(json.dumps(_cache, ensure_ascii=False, indent=0), encoding="utf-8")


def protect(text: str):
    mapping = {}
    idx = 0
    pat = re.compile(
        r"("
        r"/api/[A-Za-z0-9_\-/{}.?=*&]+"
        r"|HTTP\s*\d{3}(?:\s*OK)?"
        r"|\bIT_[A-Z0-9_]+\b"
        r"|\b(?:PUBLISHED|DRAFT|SUBMITTED|WAITING_FOR_CLASS|CLASS_ASSIGNED|"
        r"LEARNER|STAFF|ADMIN|TEACHER|CM|JWT|OTP|MockMvc|PayOS|Google Meet|"
        r"Content-Type|application/json|Bearer)\b"
        r"|\b[A-Za-z][A-Za-z0-9]*(?:\.[A-Za-z][A-Za-z0-9]*)+\(\)"
        r"|\b[A-Z][A-Za-z0-9]+(?:Controller|Service|ServiceImpl|Repository|Request|Response)\b"
        r"|\b[A-Za-z][A-Za-z0-9]*\(\)"
        r"|\{[A-Za-z0-9_]+\}"
        r")"
    )

    def repl(m):
        nonlocal idx
        key = f"[[#{idx}]]"
        mapping[key] = m.group(0)
        idx += 1
        return key

    return pat.sub(repl, text), mapping


def restore(text: str, mapping: dict[str, str]) -> str:
    out = text
    for key, val in sorted(mapping.items(), key=lambda x: -len(x[0])):
        if key in out:
            out = out.replace(key, val)
            continue
        n = re.search(r"(\d+)", key).group(1)
        out, c = re.subn(rf"\[\[\s*#\s*{n}\s*\]\]", val, out, count=1)
        if not c:
            out, _ = re.subn(rf"#\s*{n}", val, out, count=1)
    return out


def polish(text: str) -> str:
    reps = [
        (r"Bộ điều khiển-Dịch vụ-Kho lưu trữ", "Controller–Service–Repository"),
        (r"Bộ điều khiển", "Controller"),
        (r"đại biểu cho", "gọi tiếp"),
        (r"Đại biểu cho", "Gọi tiếp"),
        (r"ủy quyền cho", "gọi tiếp"),
        (r"tiêu đề Ủy quyền", "header Authorization"),
        (r"Người mang ủy quyền", "Authorization Bearer"),
        (r"mã thông báo Người mang", "token Authorization Bearer"),
        (r"qua MockMvc mà không được phép", "qua MockMvc (không cần xác thực)"),
        (r"mà không có sự cho phép", " (không cần xác thực)"),
        (r"không có sự cho phép", "không cần xác thực"),
        (r"không có header Authorization", "không cần xác thực"),
        (r"without Authorization header", "không cần xác thực"),
        (r"without Authorization", "không cần xác thực"),
        (r"Trạng thái HTTP là", "HTTP status là"),
        (r"văn bản thuần túy", "plain text"),
        (r"được băm", "được hash"),
        (r"dạng băm", "dạng hash"),
        (r"tải trọng", "payload"),
        (r"^CNTT để ", "IT cho "),
        (r"\bCNTT\b", "IT"),
        (r"thông qua Controller", "qua Controller"),
    ]
    out = text
    for pat, repl in reps:
        out = re.sub(pat, repl, out)
    out = out.replace(" .", ".").replace(" ,", ",")
    out = re.sub(r"[ \t]{2,}", " ", out)
    return out


def vi(text: str) -> str:
    if not text or not text.strip():
        return text
    if text.strip() in ("None.", "None"):
        return "Không."
    key = hashlib.sha1(text.encode("utf-8")).hexdigest()
    if key in _cache:
        return _cache[key]
    protected, mapping = protect(text)
    attempt = 0
    while True:
        try:
            translated = translator.translate(protected)
            break
        except Exception:
            attempt += 1
            if attempt >= 5:
                translated = protected
                break
            time.sleep(1.2 * attempt)
    out = polish(restore(translated, mapping))
    _cache[key] = out
    return out


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def p(doc, text, bold=False, italic=False, size=11, indent=0, space_after=4):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=size, bold=bold)
    run.italic = italic
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    if indent:
        para.paragraph_format.left_indent = Pt(indent)
    return para


def say(doc, text):
    para = doc.add_paragraph()
    tag = para.add_run("NÓI: ")
    set_font(tag, bold=True, color=RGBColor(0x0B, 0x5F, 0xA5))
    body = para.add_run(text)
    set_font(body)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Pt(14)
    return para


def label(doc, text, color=RGBColor(0x1F, 0x4E, 0x79)):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=12, bold=True, color=color)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(3)
    return para


def box_en(doc, title, content):
    p(doc, title, bold=True, size=10)
    for line in (content or "").splitlines() or ["(trống)"]:
        p(doc, line, italic=True, size=10, indent=10, space_after=1)


def explain_desc(desc_vi: str, case_id: str) -> str:
    return (
        f"Cột Description nói ngắn gọn mục tiêu của {case_id}. "
        f"Ý nghĩa: {desc_vi.rstrip('.')} "
        f"Em nhấn mạnh đây là Integration Test qua Controller, không phải test UI."
    )


def explain_pre(pre_vi: str) -> str:
    if not pre_vi or pre_vi.strip() in ("Không.", "Không"):
        return (
            "Cột Pre-conditions: không cần dữ liệu đặc biệt ngoài môi trường test đã chạy. "
            "Em nói: tiền điều kiện là hệ thống SpringBootTest sẵn sàng."
        )
    return (
        "Cột Pre-conditions là điều kiện phải có TRƯỚC khi chạy test. "
        f"Cụ thể: {pre_vi} "
        "Nếu thiếu tiền điều kiện mà test fail thì không kết luận code sai."
    )


def explain_step(step_en: str, step_vi: str, idx: int) -> str:
    s = step_en.lower()
    if s.startswith("login") or "login as" in s:
        return (
            f"Bước {idx}: đăng nhập lấy JWT. "
            f"{step_vi} "
            "Token này gắn vào header Authorization các bước sau."
        )
    if "without authorization" in s or "không cần xác thực" in step_vi.lower():
        return (
            f"Bước {idx}: gọi API không cần xác thực (public hoặc case âm thiếu token). "
            f"{step_vi}"
        )
    if s.startswith("call get") or s.startswith("call post") or s.startswith("call put") or s.startswith("call patch") or s.startswith("call delete"):
        return (
            f"Bước {idx}: gửi HTTP thật bằng MockMvc. "
            f"{step_vi} "
            "MockMvc đi qua filter/security và vào Controller."
        )
    if "delegates" in s or "gọi tiếp" in step_vi.lower() or "controller" in s:
        return (
            f"Bước {idx}: chứng minh luồng Controller → Service. "
            f"{step_vi} "
            "Em nói rõ tên Controller và Service để cô thấy em hiểu kiến trúc."
        )
    if "repository" in s or "db" in s or "seed" in s or "query" in s or "save" in s or "persist" in s:
        return (
            f"Bước {idx}: tầng dữ liệu / Repository / DB. "
            f"{step_vi} "
            "Đây là điểm khác Unit Test: có đụng DB hoặc repository thật."
        )
    if "repeat" in s or "retry" in s or "lặp" in step_vi.lower() or "thử lại" in step_vi.lower():
        return (
            f"Bước {idx}: case âm / lặp lại với điều kiện khác (sai quyền, thiếu token...). "
            f"{step_vi}"
        )
    return f"Bước {idx}: {step_vi}"


def explain_expected(exp_vi: str) -> str:
    return (
        "Cột Expected Results là tiêu chí Pass/Fail. "
        f"Em chốt: {exp_vi.replace(chr(10), ' ')} "
        "Đủ các ý này thì đánh Passed."
    )


def build():
    load_cache()
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Pt(50)
        sec.bottom_margin = Pt(50)
        sec.left_margin = Pt(56)
        sec.right_margin = Pt(56)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("HƯỚNG DẪN THUYẾT TRÌNH CHI TIẾT THEO TỪNG CỘT EXCEL")
    set_font(r, size=16, bold=True)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = st.add_run(
        "EnglishLab — SEP490_G23 — Integration Test\n"
        "Thứ tự nói: Test Case ID → Description → Pre-conditions → Procedure → Expected Results\n"
        "(Không gồm AUTH)"
    )
    set_font(r2, size=11, bold=True)

    p(
        doc,
        "Cách dùng: với mỗi test case, nói lần lượt 5 cột như bên dưới. "
        "Phần chữ nghiêng = nội dung trên Excel (đã dịch dễ đọc). "
        "Phần NÓI = câu nói to khi thuyết trình.",
        italic=True,
        space_after=10,
    )

    # general method
    h = doc.add_heading("0. Khung nói chung cho MỌI test case (thuộc lòng)", level=1)
    for run in h.runs:
        set_font(run, size=14, bold=True)

    p(doc, "1) Test Case ID — Em đang nói case nào.", bold=True)
    say(doc, "Em trình bày test case [ID], thuộc sheet [SHEET], chức năng [FEATURE].")

    p(doc, "2) Description — Case này kiểm tra gì (1–2 câu).", bold=True)
    say(doc, "Mô tả: [đọc ý Description]. Mục tiêu là chứng minh [ý nghiệp vụ] qua luồng Controller–Service–Repository.")

    p(doc, "3) Pre-conditions — Trước khi chạy cần gì.", bold=True)
    say(doc, "Tiền điều kiện: [đọc Pre-conditions]. Nếu thiếu thì không kết luận code sai.")

    p(doc, "4) Procedure — Làm từng bước, nhấn mạnh API + Controller + Service + DB.", bold=True)
    say(
        doc,
        "Các bước: một, [bước 1]; hai, [bước 2]; ba, [bước 3]… "
        "Request đi MockMvc vào Controller, gọi tiếp Service, Service dùng Repository/DB.",
    )

    p(doc, "5) Expected Results — Thế nào là đạt.", bold=True)
    say(doc, "Kết quả mong đợi: HTTP …; dữ liệu JSON/DB …; phân quyền … . Đủ thì Passed.")

    # modules
    for code in PRIORITY:
        mod = next((m for m in MODULES if m["code"] == code), None)
        if not mod:
            continue
        feat = FEATURE_VI.get(mod["function"], mod["function"])
        h = doc.add_heading(f"{mod['sheet']} — {feat}", level=1)
        for run in h.runs:
            set_font(run, size=14, bold=True)

        p(doc, f"Function Name (Excel): {mod['function']}", italic=True)
        p(doc, f"Thành phần chính: {mod.get('components', '')}", italic=True)
        req_vi = vi(mod.get("requirement", ""))
        p(doc, f"Yêu cầu kiểm thử: {req_vi}")
        say(
            doc,
            f"Em chuyển sang sheet {mod['sheet']}, chức năng {feat}. "
            f"Sheet này kiểm tra {req_vi} Em xin đi từng test case theo đúng cột Excel.",
        )

        for g in mod["groups"]:
            group_name = g["name"]
            group_vi = FEATURE_VI.get(group_name, vi(group_name) if re.search(r"[A-Za-z]", group_name) else group_name)
            h2 = doc.add_heading(f"Nhóm UC: {group_vi}", level=2)
            for run in h2.runs:
                set_font(run, size=12, bold=True)

            for c in g["cases"]:
                cid = c["id"]
                desc_en = c["desc"]
                proc_en = c["proc"]
                exp_en = c["exp"]
                pre_en = c["pre"]

                desc_vi = vi(desc_en)
                pre_vi = vi(pre_en)
                exp_vi = vi(exp_en)
                steps_en = [x.strip() for x in proc_en.splitlines() if x.strip()]
                steps_vi = [vi(x) for x in steps_en]

                h3 = doc.add_heading(cid, level=3)
                for run in h3.runs:
                    set_font(run, size=12, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

                # --- ID ---
                label(doc, "① TEST CASE ID")
                box_en(doc, "Trên Excel:", cid)
                p(
                    doc,
                    "Giải thích: Đây là mã định danh duy nhất. Tiền tố sheet + số thứ tự. "
                    "Khi cô hỏi 'case nào?' em luôn nêu đúng ID này.",
                )
                say(
                    doc,
                    f"Em trình bày test case {cid}, thuộc sheet {mod['sheet']}, "
                    f"nhóm {group_vi}, chức năng {feat}.",
                )

                # --- Description ---
                label(doc, "② DESCRIPTION (Mô tả Test Case)")
                box_en(doc, "Trên Excel (gốc EN):", desc_en)
                box_en(doc, "Diễn giải dễ đọc:", desc_vi)
                p(doc, explain_desc(desc_vi, cid))
                say(
                    doc,
                    f"Về mô tả: {desc_vi} "
                    f"Case này thuộc Integration Test: em không test từng hàm lẻ mà test "
                    f"Controller nối Service nối Repository.",
                )

                # --- Pre ---
                label(doc, "③ PRE-CONDITIONS (Tiền điều kiện)")
                box_en(doc, "Trên Excel (gốc EN):", pre_en)
                box_en(doc, "Diễn giải dễ đọc:", pre_vi)
                p(doc, explain_pre(pre_vi))
                say(doc, explain_pre(pre_vi))

                # --- Procedure ---
                label(doc, "④ PROCEDURE (Các bước thực hiện)")
                box_en(doc, "Trên Excel (gốc EN):", proc_en)
                p(
                    doc,
                    "Giải thích: Procedure là kịch bản chạy test. Em nói lần lượt từng bước, "
                    "đừng đọc như robot — mỗi bước nói rõ đang làm gì và tầng nào xử lý.",
                )
                for i, (en_line, vi_line) in enumerate(zip(steps_en, steps_vi), 1):
                    p(doc, f"Bước {i} (Excel): {en_line}", italic=True, indent=8, size=10)
                    p(doc, f"→ Dễ đọc: {vi_line}", indent=8, size=10)
                    say(doc, explain_step(en_line, vi_line, i))

                say(
                    doc,
                    "Tóm lại procedure: request đi MockMvc → Controller nhận → gọi tiếp Service → "
                    "Service dùng Repository/DB → em đối chiếu kết quả.",
                )

                # --- Expected ---
                label(doc, "⑤ EXPECTED RESULTS (Kết quả mong đợi)")
                box_en(doc, "Trên Excel (gốc EN):", exp_en)
                box_en(doc, "Diễn giải dễ đọc:", exp_vi)
                p(doc, explain_expected(exp_vi))
                # per expected line
                for i, line in enumerate([x for x in exp_en.splitlines() if x.strip()], 1):
                    line_vi = vi(line)
                    p(doc, f"Ý {i}: {line_vi}", indent=8)
                say(doc, explain_expected(exp_vi))
                say(
                    doc,
                    f"Chốt {cid}: nếu đúng các ý Expected trên thì case Passed. "
                    f"Em xin chuyển case tiếp theo." if c != g["cases"][-1] or True else "",
                )

                # tip
                tip = doc.add_paragraph()
                tg = tip.add_run("NẾU CÔ HỎI THÊM: ")
                set_font(tg, bold=True, color=RGBColor(0xC0, 0x56, 0x00))
                tb = tip.add_run(
                    "Em trả lời đúng khung: (1) đang test chức năng gì, (2) API + role, "
                    "(3) Controller–Service–Repository, (4) expected HTTP/DB/quyền. "
                    "Không nói lan man ngoài Excel."
                )
                set_font(tb)
                tip.paragraph_format.space_after = Pt(12)

                save_cache()

    # closing
    h = doc.add_heading("Phần kết và hỏi đáp ngắn", level=1)
    for run in h.runs:
        set_font(run, size=14, bold=True)
    say(
        doc,
        "Tóm lại: mỗi test case em đều trình bày đủ 5 cột Excel — ID, mô tả, tiền điều kiện, "
        "các bước Controller–Service–Repository, và kết quả mong đợi. "
        "Em xin cảm ơn cô và sẵn sàng trả lời câu hỏi.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    try:
        OUT_DL.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUT_DL)
    except Exception as e:
        print("DL copy skip:", e)
    save_cache()
    print("OUT:", OUT)
    if OUT_DL.exists():
        print("OUT:", OUT_DL)


if __name__ == "__main__":
    build()
