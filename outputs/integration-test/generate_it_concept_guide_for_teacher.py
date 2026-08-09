# -*- coding: utf-8 -*-
"""
Word: Integration Test đúng nghĩa (tích hợp hàm với nhau)
+ hướng dẫn người mới + cách giải thích với cô
(Postman chỉ là tool thực thi, không phải bản chất IT)
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Huong_dan_INTEGRATION_TEST_tich_hop_ham_cho_nguoi_moi.docx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    font(run, size=11)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    font(run, size=11)


def tip(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("💡 " + text)
    font(run, size=11, color=RGBColor(0x1F, 0x4E, 0x79))


def warn(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("⚠ " + text)
    font(run, size=11, color=RGBColor(0xC0, 0x39, 0x2B))


def say(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("🎤 Nói với cô: " + text)
    font(run, size=11, color=RGBColor(0x27, 0x6E, 0x49))


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=9, name="Consolas")
    para.paragraph_format.left_indent = Cm(0.4)
    para.paragraph_format.space_after = Pt(6)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hd)
        font(r, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            font(r, size=9)
    doc.add_paragraph()


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "HƯỚNG DẪN INTEGRATION TEST\n"
        "TÍCH HỢP HÀM VỚI NHAU (CHO NGƯỜI MỚI)\n"
        "+ CÁCH GIẢI THÍCH ĐÚNG Ý CÔ GIÁO\n"
        "EnglishLab – SEP490_G23"
    )
    font(r, size=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Test case là trung tâm · Tool (Postman/JUnit) chỉ để chạy")
    font(r, size=12)

    p(doc, "Sinh viên: phongdx", bold=True)
    p(doc, "File Excel đối chiếu: SEP490_G23_Report5.2_Integration Test (các sheet IT - …)")

    # ========== A ==========
    h(doc, "A. Cô nói gì? Em phải hiểu thế nào?", 1)
    p(doc, "Cô nhắn ý chính:", bold=True)
    bullet(doc, "Integration Test = tích hợp các hàm/thành phần với nhau.")
    bullet(doc, "Postman chỉ là tool để test API.")
    bullet(doc, "Postman không phải bản chất của Integration Test → “không liên quan” theo nghĩa: không được coi Postman = IT.")
    bullet(doc, "Cần viết test case bình thường (đúng cấu trúc test case tích hợp).")

    warn(
        doc,
        "Đừng giải thích với cô kiểu “em làm IT bằng Postman”. Hãy nói: “Em viết test case tích hợp giữa Controller–Service–Repository; "
        "Postman/JUnit chỉ là công cụ thực thi case đó.”",
    )

    say(
        doc,
        "Dạ em hiểu Integration Test là kiểm tra các hàm/tầng tích hợp với nhau. "
        "Postman chỉ là tool gọi API. Phần chính em nộp là bộ test case IT trên Excel "
        "(Procedure mô tả luồng Controller → Service → DB).",
    )

    table(
        doc,
        ["Cái này", "Là gì?", "Có phải Integration Test?"],
        [
            ["Test case IT (Excel)", "Mô tả: gọi gì, hàm nào chạy, DB đổi gì, kỳ vọng gì", "CÓ – đây là sản phẩm IT"],
            ["Postman", "Tool bấm Send để gọi API", "KHÔNG – chỉ là dụng cụ"],
            ["JUnit + MockMvc", "Tool chạy test bằng code", "KHÔNG – cũng chỉ là dụng cụ"],
            ["Unit test 1 hàm mock hết", "Test 1 hàm tách rời", "KHÔNG phải Integration"],
        ],
    )

    # ========== B ==========
    h(doc, "B. Integration Test là gì? (giải thích siêu dễ)", 1)
    p(
        doc,
        "Trong project Spring Boot của em, một thao tác người dùng (ví dụ Đăng nhập) không chỉ chạy 1 hàm. "
        "Nó đi qua nhiều tầng:",
    )
    code(
        doc,
        "Request HTTP\n"
        "   → Controller  (nhận request, gọi service)\n"
        "      → Service  (xử lý nghiệp vụ, gọi repository)\n"
        "         → Repository / Database  (đọc/ghi dữ liệu)\n"
        "   ← Response JSON + HTTP Status",
    )
    p(
        doc,
        "Integration Test kiểm tra cả chuỗi này có khớp nhau: Controller gọi đúng Service, Service gọi đúng Repository, "
        "DB đúng, response đúng. Không chỉ test 1 hàm đơn lẻ.",
    )

    h(doc, "B.1. So với Unit Test", 2)
    table(
        doc,
        ["", "Unit Test", "Integration Test"],
        [
            ["Phạm vi", "1 hàm / 1 class", "Nhiều hàm/tầng cùng chạy"],
            ["Phụ thuộc", "Thường mock (giả)", "Dùng thật hoặc gần thật (DB, security…)"],
            ["Câu hỏi", "Hàm này đúng không?", "Chúng nó nối nhau đúng không?"],
            ["Ví dụ", "AuthService.login() với UserRepository giả", "Login thật: AuthController → AuthService → users → JWT → /me"],
        ],
    )

    tip(doc, "Nhớ 1 câu: Unit = tách ra test. Integration = ráp vào test.")

    # ========== C ==========
    h(doc, "C. “Viết test case bình thường” nghĩa là viết những cột gì?", 1)
    p(
        doc,
        "Một test case Integration “bình thường” (đúng chuẩn báo cáo) thường có các phần sau. "
        "Trên Excel của nhóm đã có dạng này:",
    )
    table(
        doc,
        ["Thành phần test case", "Ý nghĩa cho người mới", "Ví dụ IT_AUTH_05"],
        [
            ["Test Case ID", "Mã định danh", "IT_AUTH_05"],
            ["Description / Mục tiêu", "Định kiểm tra tích hợp gì", "Login trả JWT và /me đọc được user"],
            ["Components / Hàm liên quan", "Những class/hàm tham gia", "AuthController, AuthService, UserRepository, JwtFilter, UserController"],
            ["Pre-condition", "Điều kiện trước khi chạy", "Có LEARNER đã verify, DB chạy"],
            ["Procedure / Các bước", "Luồng tích hợp từng bước", "1. POST /login 2. AuthService xác thực 3. cấp JWT 4. GET /me qua filter"],
            ["Expected Result", "Kỳ vọng HTTP + DB + dữ liệu", "200 + accessToken; /me đúng email"],
            ["Actual / Round", "Kết quả khi chạy thật", "Passed / Failed / N/A"],
        ],
    )

    p(doc, "Procedure của IT phải nói được “hàm nào gọi hàm nào”, không chỉ “bấm Postman”.", bold=True)
    p(doc, "Ví dụ Procedure viết ĐÚNG kiểu tích hợp hàm:")
    code(
        doc,
        "1. Gọi POST /api/auth/login (đi vào AuthController.login).\n"
        "2. AuthController ủy quyền AuthService.login.\n"
        "3. AuthService xác thực qua UserRepository (đọc bảng users).\n"
        "4. AuthService phát hành JWT.\n"
        "5. Gọi GET /api/user/me kèm Bearer token.\n"
        "6. JwtAuthenticationFilter xác thực token trước khi vào UserController.\n"
        "7. UserController → UserService → UserRepository trả hồ sơ.",
    )
    p(doc, "Ví dụ Procedure viết SAI (chỉ như tool, cô không thích):")
    code(
        doc,
        "1. Mở Postman.\n"
        "2. Bấm Send login.\n"
        "3. Copy token.\n"
        "4. Bấm Send /me.",
    )
    warn(doc, "Postman vẫn dùng được để CHẠY case — nhưng nội dung test case phải viết theo kiểu tích hợp hàm ở trên.")

    # ========== D ==========
    h(doc, "D. Học cách ĐỌC 1 test case trên Excel (từng dòng)", 1)
    p(doc, "Mở sheet IT - Auth → tìm IT_AUTH_05. Đọc theo thứ tự:", bold=True)
    numbered(doc, "Đọc Description: đang kiểm tra tích hợp gì?")
    numbered(doc, "Đọc Components: liệt kê class nào? (Controller, Service, Repo, Filter…)")
    numbered(doc, "Đọc Pre-condition: thiếu gì thì phải ghi N/A, không Forced Passed.")
    numbered(doc, "Đọc Procedure: khoanh vào các tên hàm (AuthController.login, AuthService.login…).")
    numbered(doc, "Đọc Expected: Status? Có token? DB/user đúng?")
    numbered(doc, "Khi chạy thật (bằng tool): ghi Round + Note.")

    say(
        doc,
        "Em đọc test case theo đúng cấu trúc: ID – mục tiêu tích hợp – các hàm tham gia – precondition – "
        "procedure luồng hàm – expected – actual.",
    )

    # ========== E ==========
    h(doc, "E. Học cách VIẾT / GIẢI THÍCH 1 test case (làm theo mẫu)", 1)
    p(doc, "Công thức 7 bước viết Procedure tích hợp (áp dụng mọi module):", bold=True)
    numbered(doc, "Bước vào hệ thống bằng gì? (HTTP method + path) → tên Controller method.")
    numbered(doc, "Controller gọi Service method nào?")
    numbered(doc, "Service gọi Repository / component nào?")
    numbered(doc, "Dữ liệu DB thay đổi / đọc gì? (insert/update/select)")
    numbered(doc, "Có lớp Security/Filter/Mail… xen vào không?")
    numbered(doc, "Response trả về gì? (HTTP + field JSON)")
    numbered(doc, "Nếu là negative: điểm nào từ chối? (không gọi save / không cấp token…)")

    h(doc, "E.1. Mẫu hoàn chỉnh – IT_AUTH_01 (Register)", 2)
    table(
        doc,
        ["Mục", "Nội dung"],
        [
            ["ID", "IT_AUTH_01"],
            ["Mục tiêu tích hợp", "Đăng ký: AuthController ↔ AuthService ↔ UserRepository ↔ AuthTokenRepository"],
            ["Pre-condition", "DB sẵn sàng; email chưa tồn tại"],
            ["Procedure", "POST /api/auth/register → AuthController.register → AuthService.register → UserRepository.save + AuthTokenRepository lưu OTP"],
            ["Expected", "HTTP 2xx; có row users (password hash); có row auth_tokens verification"],
            ["Negative liên quan", "IT_AUTH_02: email trùng → Service từ chối, không insert user thứ 2"],
        ],
    )
    say(
        doc,
        "IT_AUTH_01 kiểm tra tích hợp đăng ký: Controller nhận request, Service tạo user và OTP qua Repository, "
        "Expected là HTTP thành công kèm dữ liệu DB tương ứng.",
    )

    h(doc, "E.2. Mẫu negative – IT_AUTH_06 (sai mật khẩu)", 2)
    p(
        doc,
        "Negative cũng là test case bình thường. Mục tiêu: chứng minh các hàm tích hợp vẫn đúng khi từ chối.",
    )
    bullet(doc, "Controller vẫn nhận request login.")
    bullet(doc, "Service xác thực thất bại (password không khớp hash trong DB).")
    bullet(doc, "Không phát hành JWT usable.")
    bullet(doc, "Expected: 401/400 — đây là Passed nếu đúng expected.")
    tip(doc, "Nhấn với cô: negative pass = hệ thống tích hợp đúng nhánh lỗi.")

    # ========== F ==========
    h(doc, "F. Tool dùng để CHẠY test case (không nhầm với viết test case)", 1)
    p(doc, "Sau khi có test case, em cần thực thi để lấy Actual Result. Có 2 tool phổ biến:", bold=True)
    table(
        doc,
        ["Tool", "Dùng khi nào", "Nói với cô thế nào"],
        [
            ["Postman", "Gọi API nhanh, xem Status/JSON, phù hợp demo", "Tool thực thi API để lấy actual result của test case IT"],
            ["JUnit + MockMvc / @SpringBootTest", "Chạy bằng code, lặp lại được, gần CI", "Tool tự động hóa cùng bộ test case IT"],
        ],
    )
    p(
        doc,
        "Cô nói “Postman không liên quan” = đừng lấy Postman thay cho việc viết/hiểu test case tích hợp. "
        "Em vẫn có thể dùng Postman để chạy, miễn là Excel/test case mô tả đúng tích hợp hàm.",
    )

    h(doc, "F.1. Khi chạy bằng Postman, em phải “đối chiếu” test case thế nào?", 2)
    numbered(doc, "Mở đúng case trên Excel (ví dụ IT_AUTH_05).")
    numbered(doc, "Nhìn Procedure: bước 1 là login → em Send POST /api/auth/login.")
    numbered(doc, "Nhìn Expected: có accessToken? → kiểm tra JSON.")
    numbered(doc, "Procedure bước /me → Send GET /api/user/me + Bearer.")
    numbered(doc, "So Expected → ghi Passed/Failed/N/A.")
    tip(doc, "Trong đầu luôn hỏi: bước này đang đi qua hàm nào? Không chỉ hỏi: Status bao nhiêu?")

    h(doc, "F.2. Khi chạy bằng code (AuthIT.java) thì sao?", 2)
    p(
        doc,
        "File AuthIT.java cũng chỉ là cách thực thi. @DisplayName(\"IT_AUTH_05 …\") map sang Excel. "
        "MockMvc.perform(...) = gọi HTTP như Postman, nhưng bằng code; Spring vẫn chạy Controller→Service→DB.",
    )
    say(
        doc,
        "Code test của em map mã IT_* trên @DisplayName với Excel; mỗi method là một test case tích hợp được tự động hóa.",
    )

    # ========== G ==========
    h(doc, "G. Bản đồ module → hàm chính (để giải thích nhanh)", 1)
    p(doc, "Khi cô hỏi “module này tích hợp gì?”, trả lời theo bảng:", bold=True)
    table(
        doc,
        ["Sheet Excel", "Tích hợp chính (nói ngắn)"],
        [
            ["IT - Auth", "AuthController ↔ AuthService ↔ UserRepository / AuthTokenRepository ↔ JWT"],
            ["IT - User", "UserController ↔ UserService ↔ UserRepository (+ avatar storage)"],
            ["IT - Commerce", "Commerce API ↔ Service giỏ/wishlist ↔ persistence"],
            ["IT - Payment", "PaymentController ↔ PaymentService ↔ PayOS/order persistence"],
            ["IT - Course", "OnlineCourseController ↔ CourseService ↔ content/progress"],
            ["IT - Classroom", "TM ClassroomController ↔ OfferingService ↔ Enrollment/Waitlist"],
            ["IT - Teacher", "TeacherController ↔ classroom services ↔ attendance/gradebook"],
            ["IT - Support", "SupportTicketController ↔ SupportTicketService ↔ DB"],
            ["IT - Admin", "AdminController ↔ admin services ↔ users/audit/config"],
        ],
    )
    tip(doc, "Không cần thuộc hết. Thuộc Auth + Classroom + 1 module em demo là đủ thuyết trình.")

    # ========== H ==========
    h(doc, "H. Cách chấm Passed / Failed / N/A (trung thực)", 1)
    table(
        doc,
        ["Kết quả", "Khi nào", "Giải thích với cô"],
        [
            ["Passed", "Actual khớp Expected (kể cả negative đúng)", "Luồng tích hợp đúng nhánh đó"],
            ["Failed", "Sai expected, 500, sai dữ liệu", "Phát hiện lỗi tích hợp/hệ thống"],
            ["N/A", "Thiếu precondition (chưa enroll, thiếu OTP, thiếu data…)", "Chưa đủ điều kiện kết luận; không phải giấu Fail"],
        ],
    )

    # ========== I ==========
    h(doc, "I. Bài tập tự học 60 phút (làm được là hiểu)", 1)
    numbered(doc, "Mở Excel sheet IT - Auth.")
    numbered(doc, "Với IT_AUTH_01, IT_AUTH_02, IT_AUTH_05: viết ra giấy 3 cột: Controller method | Service method | Repository/DB.")
    numbered(doc, "Chạy thực tế 3 case đó (Postman hoặc AuthIT) và ghi Round.")
    numbered(doc, "Tập nói to 1 phút theo mục J bên dưới.")
    numbered(doc, "Chọn thêm 1 case Classroom: nêu được OfferingService / Enrollment tham gia.")

    # ========== J ==========
    h(doc, "J. Kịch bản giải thích với cô (học thuộc ý)", 1)

    h(doc, "J.1. Nếu cô nói: Postman không liên quan", 2)
    say(
        doc,
        "Dạ em hiểu ạ. Postman chỉ là tool test API. Integration Test của em là bộ test case kiểm tra "
        "các hàm/tầng tích hợp Controller–Service–Repository. Em dùng Postman hoặc JUnit chỉ để lấy Actual Result.",
    )

    h(doc, "J.2. Mở đầu thuyết trình", 2)
    say(
        doc,
        "Em trình bày Integration Test theo hướng tích hợp hàm. Mỗi test case trên Excel mô tả "
        "precondition, procedure theo luồng hàm, expected về HTTP và dữ liệu. Sau đó em thực thi để ghi kết quả.",
    )

    h(doc, "J.3. Demo 1 case (Auth login)", 2)
    say(
        doc,
        "Ví dụ IT_AUTH_05: AuthController.login gọi AuthService.login, Service đọc UserRepository, "
        "phát JWT; sau đó UserController.getCurrentUser đi qua JwtFilter. Expected: login 200 có token, /me đúng user. "
        "Em chạy và ghi Passed.",
    )

    h(doc, "J.4. Demo negative", 2)
    say(
        doc,
        "IT_AUTH_06: cùng luồng login nhưng password sai — Service từ chối, không cấp token. "
        "Expected 4xx → Passed. Như vậy nhánh lỗi cũng được tích hợp đúng.",
    )

    h(doc, "J.5. Câu hỏi – đáp nhanh", 2)
    table(
        doc,
        ["Cô hỏi", "Em trả lời"],
        [
            ["Inte là gì?", "Tích hợp các hàm/thành phần với nhau, không test 1 hàm tách rời."],
            ["Postman là gì?", "Tool gọi API để thực thi test case, không phải định nghĩa IT."],
            ["Em viết test case ở đâu?", "Trên Excel các sheet IT - …; mỗi dòng là 1 case đủ Procedure/Expected."],
            ["Làm sao biết hàm nào chạy?", "Procedure ghi Controller → Service → Repository; đối chiếu source nếu cần."],
            ["Khác Unit Test?", "Unit mock phụ thuộc; IT để các tầng thật nối nhau."],
            ["Vì sao có N/A?", "Thiếu precondition nên chưa kết luận đủ; ghi trung thực."],
        ],
    )

    h(doc, "J.6. Kết thúc", 2)
    say(
        doc,
        "Tóm lại: sản phẩm Integration Test của em là bộ test case tích hợp hàm trên Excel; "
        "tool chỉ hỗ trợ chạy. Em có thể demo thêm case cô chỉ định.",
    )

    # ========== K ==========
    h(doc, "K. Nên dùng file Word / Excel nào?", 1)
    table(
        doc,
        ["File", "Dùng để"],
        [
            ["Word NÀY (tích hợp hàm)", "Học đúng khái niệm + giải thích với cô (ưu tiên đọc file này)"],
            ["Excel Integration Test (HONEST)", "Bộ test case chính thức + ghi Round"],
            ["Word Postman chi tiết", "Chỉ khi cần hướng dẫn bấm tool để lấy Actual — không dùng thay lời giải thích IT"],
            ["Word/Code AuthIT.java", "Tuỳ chọn: tự động hóa cùng mã IT_*"],
        ],
    )
    warn(
        doc,
        "Khi gặp cô: cầm Excel (test case) + giải thích theo Word này. Đừng mở đầu bằng “em test Postman”.",
    )

    # ========== L ==========
    h(doc, "L. Checklist trước khi gặp cô", 1)
    bullet(doc, "Thuộc câu: IT = tích hợp hàm; Postman = tool.")
    bullet(doc, "Giải thích được 1 case Auth theo Controller→Service→Repo.")
    bullet(doc, "Giải thích được 1 case negative vẫn Passed.")
    bullet(doc, "Mở được Excel đúng dòng IT_AUTH_05 / IT_AUTH_06.")
    bullet(doc, "Biết N/A khác Failed.")
    bullet(doc, "Không nói “IT của em là Postman”.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    PROJ.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out, PROJ / OUT_NAME)
    print(out)
    return out


if __name__ == "__main__":
    build()
