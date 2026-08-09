# -*- coding: utf-8 -*-
"""
Word: Dịch Test Case Procedure (Excel) → hướng dẫn chi tiết từng bước trên máy.
Bám đúng thứ tự Procedure trên Excel (Controller→Service→Repo).
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from full_modules import MODULES  # noqa: E402

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Huong_dan_dich_Procedure_Excel_chi_tiet_tung_buoc.docx"
EXCEL = "SEP490_G23_Report5.2_Integration Test_COMPLETED_HONEST_FORMATTED.xlsx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(3)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    font(run, size=10)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    font(run, size=10)


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=9, name="Consolas")
    para.paragraph_format.left_indent = Cm(0.35)


# ---------- helpers: split procedure lines ----------
def split_proc(proc: str) -> list[str]:
    lines = []
    for raw in proc.split("\n"):
        s = raw.strip()
        if not s:
            continue
        s = re.sub(r"^\d+\.\s*", "", s)
        lines.append(s)
    return lines


def translate_step(en: str) -> str:
    """Heuristic EN→VI for common IT procedure phrases; keep class/method names."""
    t = en
    reps = [
        (r"Call POST (.+?) via MockMvc with (.+)\.", r"Gọi POST \1 bằng MockMvc/Postman với \2."),
        (r"Call GET (.+?) via MockMvc with (.+)\.", r"Gọi GET \1 bằng MockMvc/Postman với \2."),
        (r"Call PUT (.+?) via MockMvc with (.+)\.", r"Gọi PUT \1 bằng MockMvc/Postman với \2."),
        (r"Call PATCH (.+?) via MockMvc with (.+)\.", r"Gọi PATCH \1 bằng MockMvc/Postman với \2."),
        (r"Call DELETE (.+?) via MockMvc(.*)\.", r"Gọi DELETE \1 bằng MockMvc/Postman\2."),
        (r"Call POST (.+?) via MockMvc\.", r"Gọi POST \1 bằng MockMvc/Postman."),
        (r"Call GET (.+?) via MockMvc\.", r"Gọi GET \1 bằng MockMvc/Postman."),
        (r"via MockMvc", "bằng MockMvc/Postman"),
        (r"delegates to", "ủy quyền / gọi tiếp"),
        (r"Authenticate as", "Đăng nhập với vai trò"),
        (r"obtain a JWT", "lấy JWT (accessToken)"),
        (r"Bearer token", "token Bearer"),
        (r"Authorization Bearer", "Header Authorization Bearer"),
        (r"Seed an?", "Chuẩn bị sẵn / seed"),
        (r"Query the", "Truy vấn bảng"),
        (r"Reload", "Tải lại"),
        (r"Compare", "So sánh"),
        (r"Observe that", "Quan sát rằng"),
        (r"must not", "không được"),
        (r"must ", "phải "),
        (r"with a unique email, password and fullName", "với email chưa dùng, password và fullName"),
        (r"with the same email", "với cùng email"),
        (r"with an incorrect password", "với mật khẩu sai"),
        (r"with a wrong OTP", "với OTP sai"),
        (r"with email and OTP", "với email và OTP"),
        (r"without Authorization header", "không gắn header Authorization"),
        (r"tables for the new email", "theo email vừa đăng ký"),
        (r"creates the account through", "tạo tài khoản qua"),
        (r"stores a verification OTP through", "lưu OTP xác thực qua"),
        (r"validates the token via", "kiểm tra token qua"),
        (r"updates the user via", "cập nhật user qua"),
        (r"checks UserRepository before insert", "kiểm tra UserRepository trước khi insert"),
        (r"Count users rows for that email", "Đếm số dòng users của email đó"),
        (r"Database is available", "Database đang chạy"),
        (r"mail sender is stubbed", "mail được stub (không cần gửi mail thật)"),
    ]
    for a, b in reps:
        t = re.sub(a, b, t, flags=re.IGNORECASE)
    # if still mostly English, wrap with note
    return t


def howto_for_step(en: str, case_id: str, step_idx: int) -> list[str]:
    """Concrete machine steps mapped from a procedure line."""
    low = en.lower()
    tips: list[str] = []

    # HTTP calls
    m = re.search(r"(POST|GET|PUT|PATCH|DELETE)\s+(/api/[^\s]+)", en, re.I)
    if m:
        method, path = m.group(1).upper(), m.group(2).rstrip(".,)")
        tips.append(f"Mở Postman → New Request → Method chọn {method}.")
        tips.append(f"URL gõ: http://localhost:8080{path}  (hoặc {{baseUrl}}{path})")
        if "authorization" in low or "bearer" in low or "jwt" in low:
            tips.append("Tab Authorization → Bearer Token → dán accessToken (hoặc {{token}}).")
        if "without authorization" in low or "without a token" in low or "no access token is required" in low:
            tips.append("Không gắn Authorization (No Auth).")
        if method in ("POST", "PUT", "PATCH") and "register" in path:
            tips.append("Body → raw → JSON, ví dụ:")
            tips.append(
                '{"email":"it.reg.demo01@englishlab-it.test","password":"Password123!","fullName":"IT Register User"}'
            )
        elif method in ("POST", "PUT", "PATCH") and "login" in path:
            tips.append("Body → raw → JSON:")
            tips.append('{"email":"0386852628z@gmail.com","password":"Password123!"}')
        elif method in ("POST", "PUT", "PATCH") and "verify-email" in path:
            tips.append('Body JSON: {"email":"<email>","code":"<OTP>"} — field tên code.')
        elif method in ("POST", "PUT", "PATCH") and ("reset-password" in path or "forgot" in path):
            tips.append("Body JSON theo Swagger; OTP field = code.")
        elif method in ("POST", "PUT", "PATCH"):
            tips.append("Body → raw → JSON theo đúng request của API (xem Swagger nếu cần).")
        tips.append("Bấm Send → ghi lại Status (200/201/4xx…).")

    if "authcontroller" in low and "authservice" in low:
        tips.append(
            "Ý nghĩa tích hợp: request vào AuthController trước, Controller không tự lưu DB mà gọi AuthService."
        )
        tips.append("Bạn không bấm vào class Java — chỉ cần biết luồng đi đúng tầng này khi API chạy.")

    if "userrepository" in low or "users" in low and ("save" in low or "query" in low or "insert" in low or "row" in low):
        tips.append("Kiểm tra DB: mở pgAdmin/DBeaver → database englishlab → bảng users.")
        tips.append("Tìm email vừa dùng. Password phải là chuỗi hash (không phải Password123!明文).")

    if "authtokenrepository" in low or "auth_tokens" in low:
        tips.append("Kiểm tra DB bảng auth_tokens: SELECT * FROM auth_tokens ORDER BY id DESC LIMIT 10;")
        tips.append("Phải có dòng OTP gắn user vừa đăng ký (verification).")

    if "count users rows" in low or "exactly one users row" in low:
        tips.append("Đếm số user cùng email: chỉ được 1 dòng → đăng ký trùng bị chặn đúng.")

    if "email_verified" in low:
        tips.append("Xem cột email_verified (hoặc tương đương) trên bảng users.")

    if "wrong otp" in low or "invalid otp" in low or "incorrect password" in low:
        tips.append("Đây là bước negative: cố tình nhập sai để hệ thống từ chối.")

    if "jwt" in low or "accesstoken" in low.replace(" ", ""):
        tips.append("Trong Response JSON tìm accessToken — copy chuỗi dài.")

    if "get /api/user/me" in low or "/api/user/me" in low:
        tips.append("Sau khi có token: GET /api/user/me + Bearer → kiểm tra email khớp.")

    if "seed" in low:
        tips.append(
            "Seed = dữ liệu đã có sẵn (demo). Dùng tài khoản Password123! đã seed, không cần tự insert nếu data demo đủ."
        )

    if not tips:
        tips.append("Làm đúng ý bước này trên hệ thống đang chạy (API/DB) như mô tả Procedure.")
        tips.append(f"Giữ tên class/method gốc để đối chiếu Excel: « {en[:120]} »")

    tips.append("Sau bước này: tick tạm vào checklist rồi sang bước Procedure tiếp theo.")
    return tips


def expand_expected(exp: str) -> list[str]:
    out = []
    for line in exp.split("\n"):
        s = line.strip()
        if not s:
            continue
        out.append(translate_step(s))
    return out


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.6)
        s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(1.7)
        s.right_margin = Cm(1.7)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "HƯỚNG DẪN LÀM THEO CỘT TEST CASE PROCEDURE (EXCEL)\n"
        "DỊCH TỪNG BƯỚC + CÁCH LÀM CHI TIẾT TRÊN MÁY\n"
        "EnglishLab – SEP490_G23"
    )
    font(r, size=14, bold=True)

    p(doc, f"Bám đúng thứ tự bước trong Excel file: {EXCEL}", bold=True)
    p(
        doc,
        "Cách đọc: mỗi Test Case ID → lấy đúng các dòng 1. 2. 3. … trong cột Test Case Procedure → "
        "dưới đây có (A) bản dịch, (B) hướng dẫn bấm máy cho từng bước đó.",
    )

    h(doc, "0. Ví dụ mẫu cô đang xem: IT_AUTH_01 (Register & verify)", 1)
    p(doc, "Trên Excel sheet IT - Auth, dòng IT_AUTH_01, cột Test Case Procedure có 4 bước. Làm lần lượt:", bold=True)

    # Hardcoded crystal-clear walkthrough matching the screenshot
    demo = [
        (
            "1. Call POST /api/auth/register via MockMvc with a unique email, password and fullName.",
            "Gọi POST /api/auth/register (MockMvc hoặc Postman) với email chưa dùng, password và fullName.",
            [
                "Bật backend localhost:8080 và PostgreSQL.",
                "Mở Postman → Method POST → URL http://localhost:8080/api/auth/register",
                "Body → raw → JSON, dán (đổi email nếu trùng):",
                '{"email":"it.reg.demo01@englishlab-it.test","password":"Password123!","fullName":"IT Register User"}',
                "Không cần Bearer token (đăng ký công khai).",
                "Bấm Send. Nhớ Status (cần 200/201).",
                "MockMvc trong code test cũng gọi cùng URL này — cùng một Procedure.",
            ],
        ),
        (
            "2. AuthController.register() delegates to AuthService.register().",
            "AuthController.register() ủy quyền / gọi tiếp AuthService.register().",
            [
                "Bạn không cần mở IDE để “bấm” hàm này khi chạy Postman — đây là bước tích hợp hàm.",
                "Ý nghĩa: request đi vào AuthController trước, Controller gọi AuthService (không tự lưu DB).",
                "Khi giải thích với cô: bước 2 chứng minh Controller và Service nối với nhau.",
                "Nếu chạy JUnit AuthIT: MockMvc cũng đi đúng qua 2 class này.",
            ],
        ),
        (
            "3. AuthService creates the account through UserRepository.save() and stores a verification OTP through AuthTokenRepository.",
            "AuthService tạo tài khoản qua UserRepository.save() và lưu OTP xác thực qua AuthTokenRepository.",
            [
                "Đây là bước Service → Repository → Database.",
                "Sau khi Send register thành công, mở tool DB (pgAdmin/DBeaver).",
                "Bảng users: tìm email vừa đăng ký — phải có 1 dòng; cột password là hash (không phải plaintext).",
                "Bảng auth_tokens: phải có dòng OTP verification gắn user đó.",
                "SQL gợi ý: SELECT * FROM auth_tokens ORDER BY id DESC LIMIT 10;",
            ],
        ),
        (
            "4. Query the users and auth_tokens tables for the new email.",
            "Truy vấn bảng users và auth_tokens theo email vừa đăng ký để xác nhận dữ liệu.",
            [
                "Tự kiểm tra lại 2 bảng như bước 3 (đây là bước Expected/verify dữ liệu).",
                "Khớp Expected Excel: HTTP 200/201; có users; có auth_tokens; không cần accessToken để gọi register.",
                "Ghi Round trên Excel: Passed nếu đủ; Failed nếu 500/không có row; N/A nếu DB không mở được.",
            ],
        ),
    ]

    for en, vi, how in demo:
        para = doc.add_paragraph()
        rr = para.add_run("Procedure gốc (Excel): ")
        font(rr, size=10, bold=True)
        rr2 = para.add_run(en)
        font(rr2, size=10)
        p(doc, "Dịch: " + vi, bold=True)
        p(doc, "Hướng dẫn chi tiết trên máy:", bold=True)
        for x in how:
            if x.startswith("{") or x.startswith("SELECT"):
                code(doc, x)
            else:
                numbered(doc, x)
        p(doc, "—" * 28)

    p(doc, "Expected Results (dịch nhanh):", bold=True)
    bullet(doc, "Response 200/201 kèm thông báo thành công.")
    bullet(doc, "Có 1 dòng users; password đã hash (không lưu plaintext).")
    bullet(doc, "Có 1 dòng auth_tokens verification gắn user.")
    bullet(doc, "Không cần accessToken để gọi API register.")

    # ---------- ALL CASES ----------
    h(doc, "1. Toàn bộ Test Case ID — dịch Procedure từng bước + hướng dẫn", 1)
    p(
        doc,
        "Dưới đây làm giống ví dụ IT_AUTH_01: với mỗi dòng trong cột Procedure trên Excel, "
        "có bản dịch và cách làm trên máy. Làm đúng thứ tự 1 → 2 → 3 → …",
    )

    total = 0
    for mi, m in enumerate(MODULES, 1):
        h(doc, f"Sheet {m['sheet']} — {m['name']}", 2)
        for g in m["groups"]:
            p(doc, f"Nhóm Excel: {g['name']}", bold=True)
            for c in g["cases"]:
                total += 1
                cid = c["id"]
                para = doc.add_paragraph()
                rr = para.add_run(cid)
                font(rr, size=12, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

                bullet(doc, f"Mở Excel → sheet {m['sheet']} → Ctrl+F → {cid}")
                bullet(doc, "Đọc Pre-conditions trước, rồi làm lần lượt các bước Procedure bên dưới.")

                steps = split_proc(c["proc"])
                for i, en in enumerate(steps, 1):
                    p(doc, f"Bước Procedure {i} (như trên Excel):", bold=True)
                    code(doc, f"{i}. {en}")
                    p(doc, "Dịch tiếng Việt: " + translate_step(en))
                    p(doc, "Làm trên máy:", bold=True)
                    for tip in howto_for_step(en, cid, i):
                        if tip.startswith("{") or tip.startswith("SELECT") or tip.startswith('{"'):
                            code(doc, tip)
                        else:
                            numbered(doc, tip)

                p(doc, "Đối chiếu Expected Results (dịch):", bold=True)
                for e in expand_expected(c["exp"]):
                    bullet(doc, e)

                p(doc, "Ghi Excel: Round = Passed/Failed/N/A · Test date · Tester", bold=True)
                p(doc, "—" * 32)

    h(doc, "2. Nhắc với cô", 1)
    p(
        doc,
        "Em làm đúng theo cột Test Case Procedure trên Excel: mỗi bước là một mắt xích tích hợp "
        "(gọi API → Controller → Service → Repository/DB). Tài liệu này dịch và hướng dẫn chi tiết từng bước đó.",
    )
    p(doc, f"Tổng số Test Case ID: {total}.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    shutil.copy2(out, PROJ / OUT_NAME)
    print(out, "cases", total)
    return out


if __name__ == "__main__":
    build()
