# -*- coding: utf-8 -*-
"""
Word: Bộ TEST CASE từng module + bước thực hiện CHI TIẾT trên máy (cho người mới).
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from full_modules import MODULES  # noqa: E402

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Bo_TEST_CASE_tung_module_BUOC_CHI_TIET_cho_nguoi_moi.docx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.1


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    font(run, size=10)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    font(run, size=10)


def field(doc, label, value):
    para = doc.add_paragraph()
    r1 = para.add_run(label + " ")
    font(r1, size=11, bold=True)
    r2 = para.add_run(str(value))
    font(r2, size=11)
    para.paragraph_format.space_after = Pt(2)


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=9, name="Consolas")
    para.paragraph_format.left_indent = Cm(0.35)
    para.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Chi tiết bước máy theo từng mã case (người mới)
# ---------------------------------------------------------------------------

COMMON_OPEN = [
    "Bật máy tính và mở Postman (icon màu cam).",
    "Góc trên bên phải Postman: chọn Environment tên EnglishLab-Local (nếu chưa có thì xem Phần chuẩn bị đầu tài liệu).",
    "Kiểm tra backend đang chạy: mở Chrome gõ http://localhost:8080/swagger-ui.html — nếu trang mở được là ổn.",
]

LOGIN_LEARNER = [
    "Trong Postman, mở (hoặc tạo) request tên LOGIN_LEARNER.",
    "Ô Method (bên trái URL): chọn POST.",
    "Ô URL gõ đúng: {{baseUrl}}/api/auth/login",
    "Bấm tab Body → chọn raw → góc phải chọn JSON (không chọn Text).",
    "Dán đoạn JSON sau vào ô lớn:",
]
LOGIN_BODY = '{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!"\n}'
AFTER_LOGIN = [
    "Bấm nút màu xanh Send.",
    "Nhìn dòng Status bên dưới: phải thấy 200 OK.",
    "Trong Body response, tìm chữ accessToken — copy chuỗi dài phía sau (hoặc dùng script đã lưu sẵn vào biến token).",
    "Sang request tiếp theo: tab Authorization → Type = Bearer Token → ô Token gõ {{token}} rồi Save.",
]


def steps_get(url, need_login=True, role="LEARNER"):
    s = []
    s.extend(COMMON_OPEN)
    if need_login:
        if role == "LEARNER":
            s.append("Nếu chưa login: làm request LOGIN_LEARNER (POST /api/auth/login) để có {{token}}.")
        else:
            s.append(f"Login đúng role {role} và lưu token vào biến Environment tương ứng.")
    s += [
        "Collections → Add request → đặt tên theo mã case.",
        "Method chọn GET.",
        f"URL gõ: {url}",
    ]
    if need_login:
        tok = {
            "LEARNER": "{{token}}",
            "TEACHER": "{{teacherToken}}",
            "TM": "{{tmToken}}",
            "STAFF": "{{staffToken}}",
            "MANAGER": "{{managerToken}}",
            "CM": "{{cmToken}}",
            "ADMIN": "{{adminToken}}",
        }.get(role, "{{token}}")
        s += [
            "Tab Authorization → Type = Bearer Token.",
            f"Ô Token gõ: {tok}",
        ]
    else:
        s.append("Không gắn Authorization (để trống).")
    s += [
        "Bấm Send.",
        "Nhìn Status (góc dưới). So với mục Expected của case.",
        "Mở Excel sheet tương ứng → ghi Round: Passed/Failed/N/A và Note = HTTP …",
    ]
    return s


def steps_post(url, body, need_login=True, role="LEARNER", extra_before=None, extra_after=None):
    s = []
    s.extend(COMMON_OPEN)
    if extra_before:
        s.extend(extra_before)
    if need_login:
        s.append(f"Đảm bảo đã login role {role} và có token trong Environment.")
    s += [
        "Tạo request mới (Add request), đặt tên theo mã IT.",
        "Method chọn POST.",
        f"URL: {url}",
        "Tab Body → raw → JSON.",
        "Xóa hết chữ trong ô body rồi dán JSON mẫu bên dưới (nhớ sửa id nếu cần).",
    ]
    if need_login:
        tok = {
            "LEARNER": "{{token}}",
            "TEACHER": "{{teacherToken}}",
            "TM": "{{tmToken}}",
            "STAFF": "{{staffToken}}",
            "MANAGER": "{{managerToken}}",
            "CM": "{{cmToken}}",
            "ADMIN": "{{adminToken}}",
        }.get(role, "{{token}}")
        s += ["Tab Authorization → Bearer Token → " + tok]
    s += ["Bấm Send.", "Đọc Status + vài dòng JSON.", "Ghi Excel Round + Note."]
    if extra_after:
        s.extend(extra_after)
    return s, body


def steps_put(url, body, role="LEARNER"):
    s, b = steps_post(url, body, need_login=True, role=role)
    # replace POST with PUT in copy
    s = [x.replace("Method chọn POST.", "Method chọn PUT.") for x in s]
    return s, b


def steps_patch(url, body, role="LEARNER"):
    s, b = steps_post(url, body, need_login=True, role=role)
    s = [x.replace("Method chọn POST.", "Method chọn PATCH.") for x in s]
    return s, b


def steps_delete(url, role="LEARNER"):
    s = steps_get(url, need_login=True, role=role)
    s = [x.replace("Method chọn GET.", "Method chọn DELETE.") for x in s]
    return s


# Per-case detailed steps: returns (steps_list, optional_body_str)
DETAILED = {}


def _reg():
    # AUTH
    DETAILED["IT_AUTH_01"] = (
        COMMON_OPEN
        + [
            "Tạo request AUTH_01_register.",
            "Method = POST. URL = {{baseUrl}}/api/auth/register",
            "Body → raw → JSON. Dán body; nếu email đã dùng thì đổi số demo01 → demo02…",
            "Không cần Authorization.",
            "Bấm Send. Status 200/201 = đạt.",
            "Ghi Excel IT_AUTH_01: Passed + Note HTTP …",
        ],
        '{\n  "email": "it.reg.demo01@englishlab-it.test",\n  "password": "Password123!",\n  "fullName": "IT Register User"\n}',
    )
    DETAILED["IT_AUTH_02"] = (
        COMMON_OPEN
        + [
            "Copy request register → đổi tên AUTH_02.",
            "Giữ POST {{baseUrl}}/api/auth/register",
            "Đổi email thành đúng email Learner đã có: 0386852628z@gmail.com",
            "Send. Phải ra 400 hoặc 409 (bị từ chối).",
            "Nếu ra 200 = Failed. Nếu 4xx = Passed (negative đúng).",
        ],
        '{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!",\n  "fullName": "Dup"\n}',
    )
    DETAILED["IT_AUTH_03"] = (
        COMMON_OPEN
        + [
            "Chạy lại AUTH_01 với email mới (chưa dùng).",
            "Mở phần mềm DB (pgAdmin hoặc DBeaver) → kết nối database englishlab.",
            "Chạy SQL: SELECT * FROM auth_tokens ORDER BY id DESC LIMIT 10;",
            "Tìm dòng type liên quan EMAIL / VERIFICATION → copy mã code/token.",
            "Postman: POST {{baseUrl}}/api/auth/verify-email",
            "Body: email đúng + code = mã vừa copy (field tên là code, không phải otp).",
            "Send → 200 thì OK. Thử login email đó.",
            "Không lấy được OTP → ghi N/A trên Excel.",
        ],
        '{\n  "email": "it.reg.demo01@englishlab-it.test",\n  "code": "123456"\n}',
    )
    DETAILED["IT_AUTH_04"] = (
        COMMON_OPEN
        + [
            "POST {{baseUrl}}/api/auth/verify-email",
            "Body dùng code cố định 000000 (cố tình sai).",
            "Send → thường 400. Ghi Passed nếu bị từ chối.",
            "Lưu ý: user demo có thể đã verify sẵn — ghi Note đúng Status bạn thấy.",
        ],
        '{\n  "email": "0386852628z@gmail.com",\n  "code": "000000"\n}',
    )
    DETAILED["IT_AUTH_05"] = (
        COMMON_OPEN
        + LOGIN_LEARNER
        + [
            "Dán body login Learner.",
            "Send → 200, có accessToken.",
            "Tạo request GET {{baseUrl}}/api/user/me",
            "Authorization → Bearer Token → dán token hoặc {{token}}",
            "Send → 200 và email đúng 0386852628z@gmail.com",
            "Cả 2 bước OK → Passed.",
        ],
        LOGIN_BODY,
    )
    DETAILED["IT_AUTH_06"] = (
        COMMON_OPEN
        + [
            "POST {{baseUrl}}/api/auth/login",
            "Body password đổi thành WrongPass999!",
            "Send → 401/400, không có accessToken dùng được → Passed.",
        ],
        '{\n  "email": "0386852628z@gmail.com",\n  "password": "WrongPass999!"\n}',
    )
    DETAILED["IT_AUTH_07"] = (
        COMMON_OPEN
        + [
            "GET {{baseUrl}}/api/user/me",
            "Tab Authorization → chọn No Auth (hoặc xóa header Authorization).",
            "Send → 401/403 → Passed.",
        ],
        "",
    )
    DETAILED["IT_AUTH_08"] = (
        COMMON_OPEN
        + [
            "POST {{baseUrl}}/api/auth/forgot-password",
            "Body chỉ có email Learner.",
            "Send. Nếu bảo đợi vài giây → đợi rồi Send lại.",
            "Status không phải 500 → có thể Passed.",
        ],
        '{\n  "email": "0386852628z@gmail.com"\n}',
    )
    DETAILED["IT_AUTH_09"] = (
        COMMON_OPEN
        + [
            "Làm AUTH_08 trước.",
            "Mở DB bảng auth_tokens, tìm type PASSWORD_RESET, copy code.",
            "POST {{baseUrl}}/api/auth/reset-password",
            "newPassword nên để Password123! để không phá tài khoản demo.",
            "Send → 200. Login lại xác nhận.",
            "Không có OTP → N/A.",
        ],
        '{\n  "email": "0386852628z@gmail.com",\n  "code": "123456",\n  "newPassword": "Password123!"\n}',
    )
    DETAILED["IT_AUTH_10"] = (
        COMMON_OPEN
        + [
            "POST {{baseUrl}}/api/auth/reset-password với code 000000.",
            "Send → 400 → Passed.",
        ],
        '{\n  "email": "0386852628z@gmail.com",\n  "code": "000000",\n  "newPassword": "Password123!"\n}',
    )

    # USER
    DETAILED["IT_USER_01"] = (steps_get("{{baseUrl}}/api/user/me"), "")
    DETAILED["IT_USER_02"] = steps_put(
        "{{baseUrl}}/api/user/me",
        '{\n  "fullName": "Learner IT Update",\n  "phoneNumber": "0901234567",\n  "targetExam": "TOEIC",\n  "targetScore": 700,\n  "studyGoal": "IT"\n}',
    )
    DETAILED["IT_USER_03"] = steps_put(
        "{{baseUrl}}/api/user/me/password",
        '{\n  "currentPassword": "WrongCurrent!",\n  "newPassword": "Password123!"\n}',
    )
    DETAILED["IT_USER_04"] = (
        COMMON_OPEN
        + [
            "Login Learner lấy {{token}}.",
            "POST {{baseUrl}}/api/user/me/avatar",
            "Tab Body → form-data (không chọn raw).",
            "Key gõ: file → đổi Type từ Text sang File → chọn 1 ảnh PNG/JPG nhỏ trên máy.",
            "Authorization Bearer {{token}}.",
            "Send → 200/201 = Passed. Lỗi định dạng ảnh → N/A + Note.",
        ],
        "",
    )
    DETAILED["IT_USER_05"] = (
        COMMON_OPEN
        + [
            "PUT {{baseUrl}}/api/user/me",
            "Body JSON nhỏ: {\"fullName\":\"x\"}",
            "Không gắn Bearer.",
            "Send → 401/403 → Passed.",
        ],
        '{"fullName":"x"}',
    )

    # NOTIF
    DETAILED["IT_NOTIF_01"] = (steps_get("{{baseUrl}}/api/user/me/notification-preferences"), "")
    DETAILED["IT_NOTIF_02"] = steps_put(
        "{{baseUrl}}/api/user/me/notification-preferences",
        '{\n  "inAppEnabled": false,\n  "emailEnabled": true,\n  "larkEnabled": false\n}',
    )
    DETAILED["IT_NOTIF_03"] = steps_put("{{baseUrl}}/api/user/me/notification-preferences", "{}")
    DETAILED["IT_NOTIF_04"] = (steps_get("{{baseUrl}}/api/student/notifications"), "")
    DETAILED["IT_NOTIF_05"] = (
        steps_get("{{baseUrl}}/api/student/notifications/unread-count")
        + [
            "Nếu list có id: tạo PATCH {{baseUrl}}/api/student/notifications/{id}/read rồi Send.",
        ],
        "",
    )

    # COMMERCE
    DETAILED["IT_COMMERCE_01"] = (
        COMMON_OPEN
        + [
            "GET {{baseUrl}}/api/online-courses (không token) → copy id khóa đầu tiên.",
            "Environment → sửa courseId = số vừa copy → Save.",
            "DELETE {{baseUrl}}/api/student/commerce/cart + Bearer {{token}} → Send (xóa sạch).",
            "POST {{baseUrl}}/api/student/commerce/cart/{{courseId}} + Bearer → Send.",
            "GET {{baseUrl}}/api/student/commerce/cart → thấy khóa → Passed.",
        ],
        "",
    )
    DETAILED["IT_COMMERCE_02"] = (
        COMMON_OPEN
        + [
            "POST {{baseUrl}}/api/student/commerce/wishlist/{{courseId}} + Bearer.",
            "Tìm endpoint move-to-cart trên Swagger (thường gần wishlist) → POST.",
            "GET cart kiểm tra. Đã có trong giỏ → có thể N/A.",
        ],
        "",
    )
    DETAILED["IT_COMMERCE_03"] = (steps_delete("{{baseUrl}}/api/student/commerce/cart"), "")
    DETAILED["IT_COMMERCE_04"] = (
        ["Lặp lại đúng các thao tác của IT_COMMERCE_01 (thêm lại vào giỏ)."],
        "",
    )

    # PAYMENT
    DETAILED["IT_PAYMENT_01"] = steps_post(
        "{{baseUrl}}/api/student/payments/payos/link",
        '{\n  "courseIds": [1]\n}',
        extra_before=["Thay số 1 bằng courseId thật lấy từ GET /api/online-courses.", "Không cần mở link PayOS trả tiền thật."],
    )
    DETAILED["IT_PAYMENT_02"] = steps_post(
        "{{baseUrl}}/api/student/payments/quote",
        '{\n  "courseIds": [1]\n}',
        extra_before=["Thay courseIds bằng id thật."],
    )
    DETAILED["IT_PAYMENT_03"] = (
        COMMON_OPEN
        + [
            "POST {{baseUrl}}/api/payos/webhook",
            "Body raw JSON: {}",
            "Không gắn chữ ký PayOS.",
            "Send → 400 = Passed (từ chối đúng). 404 = Failed.",
        ],
        "{}",
    )
    DETAILED["IT_PAYMENT_04"] = (steps_get("{{baseUrl}}/api/manager/payments/orders", role="MANAGER"), "")
    DETAILED["IT_PAYMENT_05"] = (
        steps_get("{{baseUrl}}/api/manager/payments/orders", role="MANAGER")
        + ["Nếu có id đơn: mở chi tiết theo Swagger (GET by id) nếu case yêu cầu."],
        "",
    )

    # COURSE
    DETAILED["IT_COURSE_01"] = (steps_get("{{baseUrl}}/api/online-courses", need_login=False), "")
    DETAILED["IT_COURSE_02"] = (
        COMMON_OPEN
        + [
            "Từ COURSE_01 copy slug hoặc id.",
            "GET {{baseUrl}}/api/online-courses/{{slugOrId}} (không token).",
            "Send → 200 → Passed.",
        ],
        "",
    )
    DETAILED["IT_COURSE_03"] = (steps_get("{{baseUrl}}/api/student/online-courses/{{courseId}}/content"), "")
    DETAILED["IT_COURSE_04"] = steps_patch(
        "{{baseUrl}}/api/student/online-courses/{{courseId}}/lessons/{{lessonId}}/progress",
        '{\n  "completed": true\n}',
    )
    DETAILED["IT_COURSE_05"] = steps_post(
        "{{baseUrl}}/api/student/online-courses/{{courseId}}/rating",
        '{\n  "score": 5,\n  "comment": "ok"\n}',
    )
    DETAILED["IT_COURSE_06"] = (
        steps_get("{{baseUrl}}/api/student/online-courses/{{courseId}}/content")
        + ["Nếu 400 vì chưa enroll → ghi N/A hoặc Passed tùy Expected trên Excel (thường N/A thiếu precondition / Passed nếu expected là từ chối)."],
        "",
    )

    # DISCUSS
    DETAILED["IT_DISCUSS_01"] = steps_post(
        "{{baseUrl}}/api/student/online-courses/{{courseId}}/discussions",
        '{\n  "title": "IT discuss title",\n  "content": "Noi dung thao luan integration test"\n}',
        extra_before=["Cần courseId đã enroll. Chưa enroll mà 400 → N/A."],
    )
    DETAILED["IT_DISCUSS_02"] = (steps_get("{{baseUrl}}/api/online-courses/{{courseId}}/discussions", need_login=False), "")
    DETAILED["IT_DISCUSS_03"] = steps_post(
        "{{baseUrl}}/api/student/online-courses/discussions/{{threadId}}/reports",
        '{\n  "reason": "Spam test IT"\n}',
        extra_before=["Lấy threadId từ list discussions."],
    )
    DETAILED["IT_DISCUSS_04"] = (
        ["Giống DISCUSS_03; nếu không có threadId → ghi N/A."],
        "",
    )
    DETAILED["IT_DISCUSS_05"] = (steps_get("{{baseUrl}}/api/content-manager/discussion-reports", role="CM"), "")

    # CONTENT / PACKAGE / CURRICULUM
    DETAILED["IT_CONTENT_01"] = (steps_get("{{baseUrl}}/api/content-manager/online-courses", role="CM"), "")
    DETAILED["IT_CONTENT_02"] = (steps_get("{{baseUrl}}/api/content-manager/online-courses", role="CM") + ["Xem Swagger nếu case yêu cầu POST tạo khóa — làm thêm theo form API."], "")
    DETAILED["IT_CONTENT_03"] = (steps_get("{{baseUrl}}/api/content-manager/online-courses", role="CM") + ["Publish/unpublish: tìm PUT/PATCH trên Swagger theo id khóa."], "")
    DETAILED["IT_CONTENT_04"] = (steps_get("{{baseUrl}}/api/content-manager/online-courses", role="CM"), "")
    DETAILED["IT_PACKAGE_01"] = (steps_get("{{baseUrl}}/api/content-manager/packages", role="CM"), "")
    DETAILED["IT_PACKAGE_02"] = (steps_get("{{baseUrl}}/api/content-manager/packages", role="CM"), "")
    DETAILED["IT_PACKAGE_03"] = (steps_get("{{baseUrl}}/api/content-manager/packages", role="CM"), "")
    DETAILED["IT_CURRICULUM_01"] = (steps_get("{{baseUrl}}/api/content-manager/curriculum-programs", role="CM"), "")
    DETAILED["IT_CURRICULUM_02"] = (steps_get("{{baseUrl}}/api/content-manager/exercise-bank", role="CM") + ["Nếu 404 thử /api/content-manager/assessment-bank"], "")
    DETAILED["IT_CURRICULUM_03"] = (steps_get("{{baseUrl}}/api/content-manager/learning-paths", role="CM"), "")
    DETAILED["IT_CURRICULUM_04"] = (steps_get("{{baseUrl}}/api/content-manager/rubrics", role="CM"), "")
    DETAILED["IT_CURRICULUM_05"] = (steps_get("{{baseUrl}}/api/content-manager/curriculum-programs", role="CM"), "")

    # ENROLLREQ
    DETAILED["IT_ENROLLREQ_01"] = steps_post(
        "{{baseUrl}}/api/student/course-enrollment-requests",
        '{\n  "courseOfferingId": 1,\n  "contactName": "HV IT",\n  "contactEmail": "0386852628z@gmail.com",\n  "contactPhone": "0901111222",\n  "consultationTrack": "TOEIC"\n}',
        extra_before=[
            "GET {{baseUrl}}/api/course-offerings → copy id → thay courseOfferingId.",
            "Sau khi tạo: GET {{baseUrl}}/api/student/course-enrollment-requests/my",
            "Nếu báo đã có form đang xử lý → N/A.",
        ],
    )
    DETAILED["IT_ENROLLREQ_02"] = (steps_get("{{baseUrl}}/api/staff/enrollment-requests", role="STAFF"), "")
    DETAILED["IT_ENROLLREQ_03"] = (steps_get("{{baseUrl}}/api/staff/enrollment-requests", role="STAFF") + ["Chi tiết xử lý: xem Swagger PUT/PATCH theo id."], "")
    DETAILED["IT_ENROLLREQ_04"] = (steps_get("{{baseUrl}}/api/student/course-enrollment-requests/my"), "")
    DETAILED["IT_ENROLLREQ_05"] = (
        ["Thử tạo request trùng / body thiếu field → kỳ vọng 4xx. Ghi Passed nếu validation đúng."],
        "{}",
    )

    # CLASSROOM TM
    DETAILED["IT_CLASS_01"] = (steps_get("{{baseUrl}}/api/classroom-offerings", need_login=False), "")
    DETAILED["IT_CLASS_02"] = (steps_get("{{baseUrl}}/api/training-manager/classrooms", role="TM") + ["Copy 1 id → lưu classroomId trong Environment."], "")
    DETAILED["IT_CLASS_03"] = (steps_get("{{baseUrl}}/api/training-manager/classrooms/{{classroomId}}", role="TM"), "")
    DETAILED["IT_CLASS_04"] = (
        steps_get(
            "{{baseUrl}}/api/training-manager/classrooms/registrations?classroomOfferingId={{classroomId}}",
            role="TM",
        ),
        "",
    )
    DETAILED["IT_CLASS_05"] = (
        COMMON_OPEN
        + [
            "Login TM → {{tmToken}}.",
            "GET registrations?status=WAITLIST&classroomOfferingId={{classroomId}}",
            "Đếm số HV waitlist. Nếu < 2 → ghi N/A (thiếu data), dừng.",
            "Nếu ≥ 2: copy 2 enrollment id, đảo thứ tự.",
            "PUT {{baseUrl}}/api/training-manager/classrooms/{{classroomId}}/waitlist/order",
            "Body enrollmentIds theo thứ tự mới.",
            "Nếu 404 endpoint → N/A + Note chưa có API trên bản đang chạy.",
        ],
        '{\n  "enrollmentIds": [101, 102]\n}',
    )
    DETAILED["IT_CLASS_06"] = (
        steps_get(
            "{{baseUrl}}/api/training-manager/classrooms/registrations?classroomOfferingId={{classroomId}}",
            role="TM",
        ),
        "",
    )
    DETAILED["IT_CLASS_07"] = (steps_get("{{baseUrl}}/api/training-manager/classrooms/{{classroomId}}", role="TM"), "")
    DETAILED["IT_CLASS_08"] = (steps_get("{{baseUrl}}/api/training-manager/classrooms", role="TM"), "")

    # LEARNER CLS
    DETAILED["IT_LEARNERCLS_01"] = (steps_get("{{baseUrl}}/api/student/classrooms/my-classrooms"), "")
    DETAILED["IT_LEARNERCLS_02"] = (steps_get("{{baseUrl}}/api/student/classrooms/{{classroomId}}/sessions"), "")
    DETAILED["IT_LEARNERCLS_03"] = (steps_get("{{baseUrl}}/api/student/classrooms/{{classroomId}}/homework"), "")
    DETAILED["IT_LEARNERCLS_04"] = (steps_get("{{baseUrl}}/api/student/classrooms/{{classroomId}}/materials"), "")
    DETAILED["IT_LEARNERCLS_05"] = (steps_get("{{baseUrl}}/api/student/classrooms/{{classroomId}}/homework"), "")
    DETAILED["IT_LEARNERCLS_06"] = (steps_get("{{baseUrl}}/api/student/classrooms/{{classroomId}}/gradebook/me"), "")

    # TEACHER
    DETAILED["IT_TEACH_01"] = (steps_get("{{baseUrl}}/api/teacher/classrooms/assigned", role="TEACHER") + ["Copy classroomId."], "")
    DETAILED["IT_TEACH_02"] = (steps_get("{{baseUrl}}/api/teacher/classrooms/{{classroomId}}/homework", role="TEACHER"), "")
    DETAILED["IT_TEACH_03"] = (
        COMMON_OPEN
        + [
            "Login Teacher.",
            "GET {{baseUrl}}/api/teacher/classrooms/{{classroomId}}/sessions → copy sessionId.",
            "GET {{baseUrl}}/api/teacher/classrooms/sessions/{{sessionId}}/attendance + Bearer teacherToken.",
            "Không có session → N/A.",
        ],
        "",
    )
    DETAILED["IT_TEACH_04"] = (steps_get("{{baseUrl}}/api/teacher/classrooms/{{classroomId}}/gradebook", role="TEACHER"), "")
    DETAILED["IT_TEACH_05"] = (steps_get("{{baseUrl}}/api/teacher/classrooms/requests/mine", role="TEACHER"), "")
    DETAILED["IT_TEACH_06"] = (steps_get("{{baseUrl}}/api/teacher/classrooms/assigned", role="TEACHER"), "")

    # QUIZ
    DETAILED["IT_QUIZ_01"] = (steps_get("{{baseUrl}}/api/teacher/classrooms/{{classroomId}}/quizzes", role="TEACHER"), "")
    DETAILED["IT_QUIZ_02"] = (steps_get("{{baseUrl}}/api/teacher/classrooms/{{classroomId}}/quizzes", role="TEACHER"), "")
    DETAILED["IT_QUIZ_03"] = (steps_get("{{baseUrl}}/api/student/classrooms/quizzes"), "")
    DETAILED["IT_QUIZ_04"] = (
        COMMON_OPEN
        + [
            "Chỉ xóa quiz do bạn tự tạo để test — không xóa quiz demo.",
            "DELETE {{baseUrl}}/api/teacher/quizzes/{{quizId}} + teacherToken.",
            "Không muốn xóa → ghi N/A.",
        ],
        "",
    )

    # ASSESS
    DETAILED["IT_ASSESS_01"] = (steps_get("{{baseUrl}}/api/student/placement-tests/current"), "")
    DETAILED["IT_ASSESS_02"] = steps_post(
        "{{baseUrl}}/api/student/placement-tests/current/submit",
        '{\n  "answers": []\n}',
    )
    DETAILED["IT_ASSESS_03"] = (steps_get("{{baseUrl}}/api/student/online-courses/{{courseId}}/assessments"), "")
    DETAILED["IT_ASSESS_04"] = (steps_get("{{baseUrl}}/api/student/mock-tests"), "")
    DETAILED["IT_ASSESS_05"] = (
        steps_get("{{baseUrl}}/api/student/online-courses/{{courseId}}/assessments")
        + ["Chưa enroll → thường 400 → N/A."],
        "",
    )
    DETAILED["IT_ASSESS_06"] = (steps_get("{{baseUrl}}/api/student/mock-tests"), "")

    # SUPPORT
    DETAILED["IT_SUPPORT_01"] = steps_post(
        "{{baseUrl}}/api/student/support-tickets",
        '{\n  "subject": "IT Support ticket",\n  "category": "TECHNICAL",\n  "message": "Mo ta chi tiet loi integration test"\n}',
    )
    DETAILED["IT_SUPPORT_02"] = (steps_get("{{baseUrl}}/api/student/support-tickets"), "")
    DETAILED["IT_SUPPORT_03"] = (steps_get("{{baseUrl}}/api/manager/support-tickets", role="MANAGER"), "")
    DETAILED["IT_SUPPORT_04"] = steps_post("{{baseUrl}}/api/student/support-tickets", "{}")

    # ADMIN
    DETAILED["IT_ADMIN_01"] = (steps_get("{{baseUrl}}/api/admin/users", role="ADMIN"), "")
    DETAILED["IT_ADMIN_02"] = (steps_get("{{baseUrl}}/api/admin/users", role="ADMIN"), "")
    DETAILED["IT_ADMIN_03"] = (steps_get("{{baseUrl}}/api/admin/audit-logs", role="ADMIN"), "")
    DETAILED["IT_ADMIN_04"] = (steps_get("{{baseUrl}}/api/admin/system/config", role="ADMIN"), "")

    # LARK
    DETAILED["IT_LARK_01"] = (
        COMMON_OPEN
        + [
            "POST {{baseUrl}}/api/lark/events",
            "Body challenge như mẫu.",
            "200 = Passed. 400 do thiếu cấu hình = N/A.",
        ],
        '{\n  "type": "url_verification",\n  "challenge": "it-challenge-123"\n}',
    )
    DETAILED["IT_LARK_02"] = (
        ["Gửi webhook thiếu field/chữ ký → quan sát 4xx. Ghi Passed/N/A theo Expected Excel."],
        "{}",
    )
    DETAILED["IT_LARK_03"] = (
        COMMON_OPEN
        + [
            "Login TM.",
            "Lấy sessionId thật từ lớp.",
            "POST {{baseUrl}}/api/training-manager/recordings/sessions/{{sessionId}}/sync-lark",
            "Thiếu session/Lark → N/A.",
        ],
        "",
    )

    # INFRA REPORT PROPOSAL DISPUTE NOTES
    DETAILED["IT_INFRA_01"] = (steps_get("{{baseUrl}}/api/training-manager/infrastructure/campuses", role="TM"), "")
    DETAILED["IT_INFRA_02"] = (steps_get("{{baseUrl}}/api/training-manager/infrastructure/rooms", role="TM"), "")
    DETAILED["IT_INFRA_03"] = (
        steps_get("{{baseUrl}}/api/training-manager/infrastructure/session-templates", role="TM"),
        "",
    )
    DETAILED["IT_REPORT_01"] = (steps_get("{{baseUrl}}/api/training-manager/dashboard", role="TM"), "")
    DETAILED["IT_REPORT_02"] = (steps_get("{{baseUrl}}/api/content-manager/revenue/analytics", role="CM"), "")
    DETAILED["IT_PROPOSAL_01"] = (steps_get("{{baseUrl}}/api/staff/classroom-proposals", role="STAFF"), "")
    DETAILED["IT_PROPOSAL_02"] = (steps_get("{{baseUrl}}/api/staff/classroom-proposals", role="STAFF"), "")
    DETAILED["IT_PROPOSAL_03"] = (steps_get("{{baseUrl}}/api/staff/classroom-proposals", role="STAFF"), "")
    DETAILED["IT_DISPUTE_01"] = (steps_get("{{baseUrl}}/api/student/attendance/disputes"), "")
    DETAILED["IT_DISPUTE_02"] = (steps_get("{{baseUrl}}/api/teacher/attendance-disputes/pending", role="TEACHER"), "")
    DETAILED["IT_DISPUTE_03"] = (
        steps_get("{{baseUrl}}/api/teacher/attendance-disputes/pending", role="TEACHER")
        + ["Nếu có id: xem Swagger để approve/reject — chỉ làm trên data test."],
        "",
    )
    DETAILED["IT_NOTES_01"] = (steps_get("{{baseUrl}}/api/student/learning/notes"), "")
    DETAILED["IT_NOTES_02"] = (steps_get("{{baseUrl}}/api/student/learning/notes"), "")


_reg()


VN_GOAL = {
    "IT_AUTH_01": "Đăng ký tài khoản mới (tích hợp AuthController → AuthService → DB).",
    "IT_AUTH_05": "Login JWT + /me (Auth + Security + User).",
    "IT_AUTH_06": "Login sai mật khẩu (negative).",
}


def get_steps(case_id: str, proc: str):
    if case_id in DETAILED:
        item = DETAILED[case_id]
        if isinstance(item, tuple):
            steps, body = item
            # steps_post returns (steps, body) already; steps_get returns list only wrapped as (list,"")
            if isinstance(steps, tuple):
                return steps[0], steps[1]
            return steps, body
    # fallback: biến procedure kỹ thuật thành bước máy chung
    steps = COMMON_OPEN + [
        "Đọc Procedure kỹ thuật bên dưới để biết API nào cần gọi.",
        "Trong Postman tạo request đúng Method/URL như Procedure (thường có dạng /api/...).",
        "Nếu API cần đăng nhập: Authorization → Bearer Token {{token}} (hoặc token đúng role).",
        "Nếu có Body: tab Body → raw → JSON, nhập theo Expected/Swagger.",
        "Bấm Send → xem Status.",
        "So với Expected → ghi Excel Passed/Failed/N/A.",
        "Procedure gốc (để hiểu tích hợp hàm):",
    ]
    for line in proc.split("\n"):
        if line.strip():
            steps.append("· " + line.strip())
    return steps, ""


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.7)
        s.bottom_margin = Cm(1.7)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "BỘ TEST CASE INTEGRATION TEST TỪNG MODULE\n"
        "+ HƯỚNG DẪN BƯỚC THỰC HIỆN CHI TIẾT TRÊN MÁY\n"
        "(Dành cho người mới dùng máy tính)\n"
        "EnglishLab – SEP490_G23"
    )
    font(r, size=15, bold=True)

    p(doc, "Người dùng tài liệu: sinh viên chưa quen Postman/API cũng làm theo từng dòng được.", bold=True)
    p(doc, "Mỗi case vẫn là TEST CASE tích hợp hàm (để giải thích với cô). Phần bước máy chỉ giúp lấy Actual Result.")

    h(doc, "PHẦN 0 – Chuẩn bị máy (làm 1 lần trước mọi module)", 1)
    numbered(doc, "Bật PostgreSQL và backend (cổng 8080).")
    numbered(doc, "Cài Postman Desktop → mở ứng dụng.")
    numbered(doc, "Tạo Environment tên EnglishLab-Local với biến:")
    bullet(doc, "baseUrl = http://localhost:8080")
    bullet(doc, "token, teacherToken, tmToken, staffToken, managerToken, cmToken, adminToken (để trống, login sau sẽ có)")
    bullet(doc, "courseId, classroomId (điền khi làm tới)")
    numbered(doc, "Góc phải Postman chọn đúng Environment EnglishLab-Local.")
    numbered(doc, "Tạo request LOGIN_LEARNER: POST {{baseUrl}}/api/auth/login")
    code(doc, LOGIN_BODY)
    numbered(doc, "Tab Scripts/Tests (Post-response) dán:")
    code(
        doc,
        'if (pm.response.code === 200) {\n  pm.environment.set("token", pm.response.json().accessToken);\n}',
    )
    numbered(doc, "Send login → kiểm tra Environment đã có token.")
    numbered(doc, "Tạo tương tự login các role khác, đổi email + tên biến token.")
    p(doc, "Tài khoản demo — mật khẩu Password123!:", bold=True)
    bullet(doc, "LEARNER: 0386852628z@gmail.com → {{token}}")
    bullet(doc, "TEACHER: classroom.teacher1@englishlab.vn → {{teacherToken}}")
    bullet(doc, "TM: training.manager@englishlab.vn → {{tmToken}}")
    bullet(doc, "STAFF: staff@englishlab.vn → {{staffToken}}")
    bullet(doc, "MANAGER: classroom.manager@englishlab.vn → {{managerToken}}")
    bullet(doc, "CM: content.manager@englishlab.vn → {{cmToken}}")
    bullet(doc, "ADMIN: classroom.admin@englishlab.vn → {{adminToken}}")

    h(doc, "Cách ghi Excel sau mỗi case", 2)
    bullet(doc, "Round = Passed / Failed / N/A")
    bullet(doc, "Note = HTTP 200 (hoặc lý do N/A ngắn)")
    bullet(doc, "Tester = tên bạn")

    h(doc, "PHẦN 1 – TEST CASE CHI TIẾT TỪNG MODULE", 1)

    total = 0
    for mi, m in enumerate(MODULES, 1):
        h(doc, f"MODULE {mi}. {m['name']} — Sheet Excel: {m['sheet']}", 2)
        field(doc, "Tích hợp chính:", m.get("components", ""))
        field(doc, "Ý nghĩa module:", m.get("requirement", ""))

        n = 0
        for g in m["groups"]:
            h(doc, f"Nhóm: {g['name']}", 3)
            for c in g["cases"]:
                n += 1
                total += 1
                cid = c["id"]
                para = doc.add_paragraph()
                rr = para.add_run(f"TEST CASE {n}: {cid}")
                font(rr, size=12, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

                field(doc, "1. Mục tiêu kiểm thử:", VN_GOAL.get(cid, c["desc"]))
                field(doc, "2. Thành phần tích hợp:", m.get("components", ""))
                p(doc, "3. Tiền điều kiện:", bold=True)
                for line in c["pre"].split("\n"):
                    if line.strip():
                        bullet(doc, line.strip())

                p(doc, "4. Ý nghĩa tích hợp hàm (để nói với cô):", bold=True)
                for line in c["proc"].split("\n"):
                    if line.strip():
                        bullet(doc, line.strip())

                p(doc, "5. CÁC BƯỚC THỰC HIỆN TRÊN MÁY (làm từng dòng):", bold=True)
                steps, body = get_steps(cid, c["proc"])
                for i, st in enumerate(steps, 1):
                    numbered(doc, st)
                if body:
                    p(doc, "Body JSON mẫu (copy/paste):", bold=True)
                    code(doc, body)

                p(doc, "6. Kết quả mong đợi:", bold=True)
                for line in c["exp"].split("\n"):
                    if line.strip():
                        bullet(doc, line.strip())

                field(
                    doc,
                    "7. Actual (điền khi làm xong):",
                    "☐ Passed   ☐ Failed   ☐ N/A    Note: ........................",
                )
                p(doc, "—" * 36)

        p(doc, f"(Module này: {n} test case)", bold=True)

    h(doc, "PHẦN 2 – Nhắc lại với cô", 1)
    p(
        doc,
        "Tài liệu này là bộ TEST CASE theo module. Mục 4 mô tả tích hợp hàm. "
        "Mục 5 chỉ hướng dẫn người mới bấm máy để lấy Actual. Postman là tool, không thay thế test case.",
    )
    p(doc, f"Tổng số test case trong tài liệu: {total}.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    PROJ.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out, PROJ / OUT_NAME)
    print(out, "cases", total)
    return out


if __name__ == "__main__":
    build()
