# -*- coding: utf-8 -*-
"""
Word: Integration Test ĐÚNG = Spring Boot Test + MockMvc
xác minh Controller–Service–Repository và Service–Service.
Postman chỉ bổ trợ, chưa đủ.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Huong_dan_Integration_Test_MockMvc_SpringBootTest.docx"
EXCEL = "SEP490_G23_Report5.2_Integration Test_COMPLETED_HONEST_FORMATTED.xlsx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    font(run, size=11)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    font(run, size=11)


def tip(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("💡 " + text)
    font(run, size=11, color=RGBColor(0x1F, 0x4E, 0x79))


def warn(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("⚠ " + text)
    font(run, size=11, color=RGBColor(0xC0, 0x39, 0x2B))


def say(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("🎤 Nói với cô: " + text)
    font(run, size=11, color=RGBColor(0x27, 0x6E, 0x49))


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=9, name="Consolas")
    para.paragraph_format.left_indent = Cm(0.4)
    para.paragraph_format.space_after = Pt(4)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hd)
        font(r, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            font(r, size=9)
    doc.add_paragraph()


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "HƯỚNG DẪN INTEGRATION TEST ĐÚNG CHUẨN\n"
        "SPRING BOOT TEST + MOCKMVC\n"
        "Xác minh Controller–Service–Repository & Service–Service\n"
        "EnglishLab – SEP490_G23"
    )
    font(r, size=15, bold=True)

    p(doc, "Kết luận quan trọng:", bold=True)
    warn(
        doc,
        "Chỉ test API bằng Postman là CHƯA ĐỦ để gọi là Integration Test theo yêu cầu môn. "
        "IT phải dùng Spring Boot Test / MockMvc để xác minh tương tác giữa các tầng/hàm trong ứng dụng.",
    )
    say(
        doc,
        "Em hiểu Integration Test phải xác minh tương tác Controller–Service–Repository và Service–Service "
        "bằng @SpringBootTest + MockMvc. Postman chỉ là công cụ phụ để quan sát API, không thay thế IT.",
    )

    # ===== 1 =====
    h(doc, "1. Vì sao Postman chưa đủ?", 1)
    table(
        doc,
        ["Cách làm", "Kiểm được gì?", "Có đủ IT?"],
        [
            ["Postman gọi API", "HTTP Status + JSON bên ngoài", "CHƯA ĐỦ — không chứng minh tầng Java nối nhau trong test framework"],
            ["Unit test 1 class + mock hết", "1 hàm tách rời", "KHÔNG phải Integration"],
            ["@SpringBootTest + MockMvc (+ DB thật/test)", "Request đi qua Controller→Service→Repository (và service khác)", "ĐÚNG Integration Test"],
        ],
    )
    p(
        doc,
        "Postman không load Spring context, không gắn với JUnit, không map @DisplayName = IT_AUTH_01 trên Excel "
        "theo kiểu “chạy test tự động trong project”. Cô yêu cầu IT trong codebase bằng Spring Boot Test.",
    )

    # ===== 2 =====
    h(doc, "2. Integration Test xác minh những tương tác nào?", 1)

    h(doc, "2.1. Controller → Service → Repository", 2)
    code(
        doc,
        "MockMvc.perform(POST /api/...)\n"
        "   → AuthController.register()\n"
        "      → AuthService.register()\n"
        "         → UserRepository.save()\n"
        "         → AuthTokenRepository.save()\n"
        "   ← HTTP 200/201 + dữ liệu DB",
    )
    p(
        doc,
        "Khi @SpringBootTest chạy, cả chuỗi bean thật được nối. MockMvc giả lập HTTP nhưng vẫn đi vào Controller thật "
        "→ Service thật → Repository thật (trừ khi bạn cố tình @MockBean).",
    )

    h(doc, "2.2. Service → Service (tương tác chéo)", 2)
    p(
        doc,
        "Một số luồng không chỉ 1 service: ví dụ PaymentService gọi OrderService / NotificationService; "
        "ClassroomOfferingService gọi EnrollmentService… IT phải để các service thật gọi nhau trong context, "
        "không mock mất hết phụ thuộc nếu mục tiêu là tích hợp.",
    )
    tip(
        doc,
        "Trên Excel cột Procedure đã viết kiểu “Controller delegates to Service… Repository…”. "
        "Code MockMvc chính là cách thực thi Procedure đó trong Spring.",
    )

    # ===== 3 =====
    h(doc, "3. Hai annotation bắt buộc phải hiểu", 1)
    table(
        doc,
        ["Annotation", "Ý nghĩa đơn giản"],
        [
            ["@SpringBootTest", "Khởi động gần như cả ứng dụng Spring (context) để các bean thật nối nhau"],
            ["@AutoConfigureMockMvc", "Cho phép dùng MockMvc gọi HTTP vào Controller mà không cần mở trình duyệt/Postman"],
            ["@DisplayName(\"IT_AUTH_01 …\")", "Map đúng Test Case ID trên Excel"],
            ["@Test", "Một method = một test case"],
            ["mockMvc.perform(...).andExpect(...)", "Gửi request giả lập + khẳng định Status/JSON"],
        ],
    )
    warn(doc, "Project dùng Spring Boot 4: import AutoConfigureMockMvc từ org.springframework.boot.webmvc.test.autoconfigure")

    # ===== 4 =====
    h(doc, "4. Cấu trúc code IT trong project (đã có mẫu)", 1)
    p(doc, "Thư mục đúng (src/test, không phải src/main):", bold=True)
    code(
        doc,
        "backend/src/test/java/fu/sap490/g23/backend/it/\n"
        "  ItSupport.java     ← helper login (không có @Test)\n"
        "  AuthIT.java        ← IT sheet Auth (có @Test + @DisplayName IT_AUTH_*)",
    )
    p(doc, "Khung class chuẩn:", bold=True)
    code(
        doc,
        "@SpringBootTest\n"
        "@AutoConfigureMockMvc\n"
        "public class AuthIT {\n"
        "    @Autowired MockMvc mockMvc;\n"
        "\n"
        "    @Test\n"
        "    @DisplayName(\"IT_AUTH_01 register\")\n"
        "    void itAuth01_register() throws Exception { ... }\n"
        "}",
    )

    # ===== 5 =====
    h(doc, "5. Map Excel Procedure → code MockMvc (ví dụ IT_AUTH_01)", 1)
    p(doc, f"Excel: {EXCEL} → sheet IT - Auth → IT_AUTH_01", bold=True)

    table(
        doc,
        ["Bước Procedure trên Excel", "Làm trong MockMvc / IT nghĩa là gì"],
        [
            ["1. Call POST /api/auth/register via MockMvc …", "mockMvc.perform(post(\"/api/auth/register\").content(...))"],
            ["2. AuthController.register() → AuthService.register()", "Spring tự đi qua 2 bean thật (không viết tay)"],
            ["3. UserRepository.save + AuthTokenRepository OTP", "Sau perform thành công, dữ liệu vào DB qua repository thật"],
            ["4. Query users / auth_tokens", "Có thể assert JSON Status; nâng cao: @Autowired repo/JdbcTemplate kiểm tra DB"],
        ],
    )

    p(doc, "Đoạn code tương ứng trong AuthIT (ý chính):", bold=True)
    code(
        doc,
        "@Test\n"
        "@DisplayName(\"IT_AUTH_01 register\")\n"
        "void itAuth01_register() throws Exception {\n"
        "    String email = \"it.reg.\" + UUID.randomUUID() + \"@englishlab-it.test\";\n"
        "    String body = \"\"\"\n"
        "            {\"email\":\"%s\",\"password\":\"%s\",\"fullName\":\"IT Register User\"}\n"
        "            \"\"\".formatted(email, PASSWORD);\n"
        "    mockMvc.perform(post(\"/api/auth/register\")\n"
        "                    .contentType(MediaType.APPLICATION_JSON)\n"
        "                    .content(body))\n"
        "            .andExpect(status().is2xxSuccessful());\n"
        "}",
    )
    tip(
        doc,
        "Khi method này PASS: chứng minh request đã đi qua Controller–Service–Repository trong Spring context "
        "(đúng tinh thần Procedure Excel).",
    )

    # ===== 6 =====
    h(doc, "6. Ví dụ IT có Security + nhiều tầng: IT_AUTH_05", 1)
    p(
        doc,
        "Procedure Excel: login → AuthService cấp JWT → GET /api/user/me đi qua JwtAuthenticationFilter → UserController → UserService → UserRepository.",
    )
    p(doc, "Đây vừa là Controller–Service–Repository, vừa có tương tác với Security filter (chuỗi tích hợp).", bold=True)
    code(
        doc,
        "String token = login(mockMvc, LEARNER, PASSWORD);  // AuthController→AuthService\n"
        "mockMvc.perform(get(\"/api/user/me\")\n"
        "        .header(\"Authorization\", bearer(token)))\n"
        "    .andExpect(status().isOk())\n"
        "    .andExpect(jsonPath(\"$.email\").value(LEARNER));",
    )

    # ===== 7 =====
    h(doc, "7. Hướng dẫn chạy MockMvc IT trên máy (chi tiết người mới)", 1)
    numbered(doc, "Bật PostgreSQL (DB englishlab) — @SpringBootTest cần kết nối DB.")
    numbered(doc, "Mở Terminal tại thư mục backend (có file mvnw.cmd):")
    code(doc, "cd D:\\EngLishLab\\EnglishLab\\backend")
    numbered(doc, "Chạy đúng class IT:")
    code(doc, ".\\mvnw.cmd \"-Dtest=AuthIT\" test")
    numbered(doc, "Hoặc chạy 1 method:")
    code(doc, ".\\mvnw.cmd \"-Dtest=AuthIT#itAuth01_register\" test")
    numbered(doc, "Xem kết quả: Tests run / Failures / Errors.")
    numbered(doc, "Mở Excel sheet IT - Auth → ghi Round theo @DisplayName (IT_AUTH_01…).")
    numbered(doc, "Trong Cursor/IDE: mở AuthIT.java → Run Test trên class (không chạy ItSupport.java).")

    warn(doc, "Đứng sai thư mục (ví dụ đứng trong .../it) sẽ không thấy mvnw.cmd. Phải cd về backend.")
    tip(doc, "Nếu lỗi TimeZone Asia/Saigon: project đã set user.timezone=Asia/Ho_Chi_Minh trong BackendApplication / surefire.")

    # ===== 8 =====
    h(doc, "8. Cách viết thêm IT cho module khác (cùng chuẩn MockMvc)", 1)
    numbered(doc, "Tạo class mới trong package fu.sap490.g23.backend.it (ví dụ UserIT, PaymentIT).")
    numbered(doc, "Gắn @SpringBootTest + @AutoConfigureMockMvc.")
    numbered(doc, "Mỗi case Excel = 1 method @Test + @DisplayName(\"IT_xxx …\").")
    numbered(doc, "Procedure Excel viết Call … via MockMvc → đổi thành mockMvc.perform(...).")
    numbered(doc, "Expected Excel → andExpect(status()…) / jsonPath…; negative → is4xxClientError().")
    numbered(doc, "Cần login: dùng ItSupport.login(...) rồi header Authorization.")
    numbered(doc, "Không @MockBean toàn bộ Service/Repository nếu mục tiêu là IT tích hợp thật.")

    p(doc, "Gợi ý đặt tên class theo sheet Excel:", bold=True)
    bullet(doc, "IT - Auth → AuthIT (đã có)")
    bullet(doc, "IT - User → UserIT")
    bullet(doc, "IT - Classroom → TrainingManagerClassroomIT / …")
    bullet(doc, "… đủ 24 sheet nếu cô yêu cầu full")

    # ===== 9 =====
    h(doc, "9. Vai trò Excel / MockMvc / Postman sau khi hiểu đúng", 1)
    table(
        doc,
        ["Thành phần", "Vai trò đúng"],
        [
            ["Excel (IT_*)", "Thiết kế test case + Procedure + Expected + ghi Round"],
            ["Spring Boot Test + MockMvc", "Thực thi IT — xác minh tương tác tầng/hàm"],
            ["Postman", "Phụ: debug tay, demo nhanh — KHÔNG thay MockMvc"],
            ["@DisplayName(IT_*)", "Cầu nối code ↔ Excel"],
        ],
    )

    # ===== 10 =====
    h(doc, "10. Checklist trước khi thuyết trình với cô", 1)
    bullet(doc, "Nói được: IT = tương tác Controller–Service–Repository (+ Service–Service) bằng @SpringBootTest/MockMvc.")
    bullet(doc, "Nói được: Postman chưa đủ.")
    bullet(doc, "Mở được AuthIT.java, chỉ @DisplayName IT_AUTH_01 / IT_AUTH_05 / IT_AUTH_06.")
    bullet(doc, "Chạy được .\\mvnw.cmd -Dtest=AuthIT test.")
    bullet(doc, "Excel Round khớp kết quả Maven.")
    bullet(doc, "Giải thích 1 Procedure Excel ↔ 1 đoạn mockMvc.perform.")

    # ===== 11 =====
    h(doc, "11. Kịch bản trả lời cô", 1)
    say(
        doc,
        "Integration Test của em dùng @SpringBootTest và MockMvc để request đi qua Controller thật, "
        "Service thật, Repository/DB thật, nên xác minh được tương tác giữa các tầng. "
        "Một số luồng còn gồm Service gọi Service khác trong cùng context.",
    )
    say(
        doc,
        "Postman chỉ giúp em quan sát API thủ công. Em không coi Postman là Integration Test đủ điều kiện. "
        "Bộ case nằm trên Excel; thực thi chính là các class *IT trong src/test.",
    )
    say(
        doc,
        "Ví dụ IT_AUTH_01: mockMvc.perform POST /api/auth/register — đi AuthController→AuthService→UserRepository/"
        "AuthTokenRepository. IT_AUTH_05 thêm nhánh JWT + UserController/UserService.",
    )

    h(doc, "11.1. Hỏi – đáp", 2)
    table(
        doc,
        ["Cô hỏi", "Em trả lời"],
        [
            ["Sao không dùng Postman?", "Postman không chạy Spring Test context; không đủ chứng minh tích hợp tầng trong IT."],
            ["Em xác minh tương tác thế nào?", "MockMvc gọi Controller; bean Service/Repository thật trong @SpringBootTest; assert Status/JSON/DB."],
            ["Service–Service?", "Context load nhiều service; luồng nghiệp vụ để service gọi nhau thật, không mock hết."],
            ["Map Excel thế nào?", "@DisplayName chứa IT_*; Procedure Excel viết via MockMvc khớp code."],
            ["Unit khác IT?", "Unit mock phụ thuộc; IT để các tầng nối nhau."],
        ],
    )

    # ===== 12 =====
    h(doc, "12. Việc cần làm tiếp (nếu cô bắt đủ module)", 1)
    numbered(doc, "Giữ Excel HONEST làm thiết kế case.")
    numbered(doc, "Hoàn thiện AuthIT (đã có) — chạy ổn, ghi Round.")
    numbered(doc, "Viết thêm *IT.java cho các sheet còn lại theo cùng pattern MockMvc.")
    numbered(doc, "Mỗi method map 1 Test Case ID.")
    numbered(doc, "Không thay IT bằng collection Postman.")

    tip(
        doc,
        "Nếu cần generate đủ class *IT cho 24 module trong src/test: nhờ mentor/AI viết code vào "
        "backend/src/test/java/.../it/ (không bỏ vào src/main).",
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    PROJ.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out, PROJ / OUT_NAME)
    print(out)
    return out


if __name__ == "__main__":
    build()
