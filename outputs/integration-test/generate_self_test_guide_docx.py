# -*- coding: utf-8 -*-
"""Generate Vietnamese self-test guide (Word) for Integration Test modules."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT = OUT_DIR / "Huong_dan_tu_test_Integration_Test_EnglishLab.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_vn(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)
    return h


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    set_run_font(run, size=11)
    return para


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    set_run_font(run, size=11)
    return para


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


MODULES = [
    {
        "name": "1. Authentication (IT - Auth)",
        "role": "Không cần token (trừ case dùng /me)",
        "account": "Tự đăng ký email mới; hoặc Learner 0386852628z@gmail.com / Password123!",
        "goal": "Kiểm tra đăng ký, xác thực email, đăng nhập JWT, quên/đặt lại mật khẩu.",
        "tool": "Postman",
        "cases": [
            ("IT_AUTH_01", "Đăng ký tài khoản mới",
             "POST /api/auth/register\nBody JSON: email (chưa dùng), password Password123!, fullName",
             "HTTP 200/201; tài khoản được tạo (có thể kiểm tra bằng login sau khi verify).",
             "Passed nếu 200/201. Failed nếu 500."),
            ("IT_AUTH_02", "Đăng ký trùng email (negative)",
             "POST /api/auth/register với email Learner đã tồn tại",
             "HTTP 400/409",
             "Passed nếu bị từ chối. Failed nếu vẫn tạo được."),
            ("IT_AUTH_03", "Xác thực email bằng OTP",
             "1) Register email mới\n2) Lấy OTP trong bảng auth_tokens (type EMAIL_VERIFICATION) hoặc mail\n3) POST /api/auth/verify-email Body: email, code=<OTP>",
             "HTTP 200; sau đó login được",
             "Passed nếu verify thành công. N/A nếu không đọc được OTP."),
            ("IT_AUTH_04", "OTP sai (negative)",
             "POST /api/auth/verify-email với otp/code = 000000",
             "HTTP 400",
             "Passed nếu bị từ chối."),
            ("IT_AUTH_05", "Login lấy JWT + gọi /me",
             "1) POST /api/auth/login {email, password}\n2) Copy accessToken\n3) GET /api/user/me Header Authorization: Bearer <token>",
             "Login 200 có accessToken; /me 200 có thông tin user",
             "Passed nếu cả 2 bước OK."),
            ("IT_AUTH_06", "Sai mật khẩu (negative)",
             "POST /api/auth/login với password sai",
             "HTTP 401/400",
             "Passed nếu không cấp token."),
            ("IT_AUTH_07", "Gọi /me không token (negative)",
             "GET /api/user/me (không gắn Authorization)",
             "HTTP 401/403",
             "Passed nếu bị chặn."),
            ("IT_AUTH_08", "Quên mật khẩu",
             "POST /api/auth/forgot-password {email}. Nếu báo chờ 15–60 giây thì đợi rồi gọi lại.",
             "HTTP 200",
             "Passed nếu 200. Failed nếu 500."),
            ("IT_AUTH_09", "Đặt lại mật khẩu bằng OTP",
             "1) forgot-password\n2) Lấy OTP type PASSWORD_RESET trong auth_tokens\n3) POST /api/auth/reset-password {email, code, newPassword}\nNên đặt lại đúng Password123! để không phá demo.",
             "HTTP 200; login bằng mật khẩu mới được",
             "Passed nếu reset + login OK."),
            ("IT_AUTH_10", "Reset OTP sai (negative)",
             "POST /api/auth/reset-password với code=000000",
             "HTTP 400",
             "Passed nếu bị từ chối."),
        ],
    },
    {
        "name": "2. Account Profile (IT - User)",
        "role": "LEARNER (Bearer token)",
        "account": "0386852628z@gmail.com / Password123!",
        "goal": "Xem/sửa hồ sơ, đổi mật khẩu (negative), avatar, bảo mật.",
        "tool": "Postman",
        "cases": [
            ("IT_USER_01", "GET hồ sơ", "GET /api/user/me + Bearer", "HTTP 200 + email/fullName", "Passed nếu 200."),
            ("IT_USER_02", "Cập nhật hồ sơ", "PUT /api/user/me với fullName, phoneNumber, targetExam, targetScore, studyGoal", "HTTP 200", "Passed nếu 200."),
            ("IT_USER_03", "Đổi mật khẩu sai current (negative)", "PUT /api/user/me/password {currentPassword sai, newPassword}", "HTTP 400", "Passed nếu bị từ chối."),
            ("IT_USER_04", "Upload avatar", "POST /api/user/me/avatar form-data field file = ảnh PNG/JPG thật", "HTTP 200/201", "Passed nếu upload OK. N/A nếu API báo không đọc được ảnh."),
            ("IT_USER_05", "Sửa hồ sơ không token (negative)", "PUT /api/user/me không Authorization", "HTTP 401/403", "Passed nếu bị chặn."),
        ],
    },
    {
        "name": "3. Notifications (IT - Notif)",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com / Password123!",
        "goal": "Preference thông báo + danh sách/đọc thông báo.",
        "tool": "Postman",
        "cases": [
            ("IT_NOTIF_01", "Lấy preference", "GET /api/user/me/notification-preferences", "HTTP 200", "Passed nếu 200."),
            ("IT_NOTIF_02", "Tắt/bật in-app", "PUT /api/user/me/notification-preferences {inAppEnabled, emailEnabled, larkEnabled} rồi bật lại", "HTTP 200", "Passed nếu 200."),
            ("IT_NOTIF_03", "Body thiếu (negative)", "PUT preference với {}", "HTTP 400", "Passed nếu validation lỗi."),
            ("IT_NOTIF_04/05", "List + unread", "GET /api/student/notifications và GET /api/student/notifications/unread-count; có thể PATCH .../{id}/read", "HTTP 200", "Passed nếu list/count 200."),
        ],
    },
    {
        "name": "4. Cart & Wishlist (IT - Commerce)",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com / Password123!",
        "goal": "Thêm/xóa giỏ, wishlist.",
        "tool": "Postman",
        "cases": [
            ("IT_COMMERCE_01", "Thêm vào giỏ", "1) GET /api/online-courses lấy id\n2) DELETE /api/student/commerce/cart (xóa sạch)\n3) POST /api/student/commerce/cart/{courseId}\n4) GET /api/student/commerce/cart", "Add 200 + GET thấy khóa", "Passed nếu thêm được."),
            ("IT_COMMERCE_02", "Wishlist → cart", "POST /api/student/commerce/wishlist/{courseId} rồi POST .../move-to-cart", "HTTP 200 hoặc báo đã có trong giỏ", "Passed nếu chuyển được. N/A nếu khóa đã trong giỏ."),
            ("IT_COMMERCE_03", "Xóa giỏ", "DELETE /api/student/commerce/cart", "HTTP 200/204", "Passed nếu xóa được."),
            ("IT_COMMERCE_04", "Thêm lại", "Lặp COMMERCE_01", "HTTP 200", "Passed nếu OK."),
        ],
    },
    {
        "name": "5. PayOS & Orders (IT - Payment)",
        "role": "LEARNER + MANAGER",
        "account": "Learner + classroom.manager@englishlab.vn / Password123!",
        "goal": "Quote, tạo link PayOS, webhook (negative), xem orders.",
        "tool": "Postman",
        "cases": [
            ("IT_PAYMENT_01", "Tạo PayOS link", "POST /api/student/payments/payos/link Body: {\"courseIds\":[<id khóa public>]}", "HTTP 200 có checkoutUrl/orderCode", "Passed nếu tạo link. Failed nếu 500."),
            ("IT_PAYMENT_02", "Quote", "POST /api/student/payments/quote {\"courseIds\":[<id>]}", "HTTP 200 có totalAmount", "Passed nếu 200."),
            ("IT_PAYMENT_03", "Webhook thiếu chữ ký (negative)", "POST /api/payos/webhook body giả {}", "HTTP 400 (từ chối)", "Passed nếu bị từ chối. Failed nếu 404."),
            ("IT_PAYMENT_04/05", "Manager xem orders", "GET /api/manager/payments/orders (token MANAGER)", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "6. Online Learning (IT - Course)",
        "role": "Public + LEARNER",
        "account": "Không token (public) / Learner",
        "goal": "Catalog công khai; content/progress cần đã mua/enroll.",
        "tool": "Postman",
        "cases": [
            ("IT_COURSE_01", "Danh sách khóa public", "GET /api/online-courses (không token)", "HTTP 200", "Passed nếu 200."),
            ("IT_COURSE_02", "Chi tiết khóa", "GET /api/online-courses/{slugOrId}", "HTTP 200", "Passed nếu 200."),
            ("IT_COURSE_03/06", "Xem content learner", "GET /api/student/online-courses/{id}/content + Bearer", "200 nếu đã enroll; 400 nếu chưa", "Passed nếu đã enroll và 200. N/A nếu chưa enroll (400 đúng business)."),
            ("IT_COURSE_04", "Cập nhật progress", "PATCH /api/student/online-courses/{id}/lessons/{lessonId}/progress {\"completed\":true}", "HTTP 200 nếu đủ quyền", "Passed/N/A tương tự."),
            ("IT_COURSE_05", "Rating", "POST /api/student/online-courses/{id}/rating {\"score\":5,\"comment\":\"ok\"}", "HTTP 200/201", "Passed nếu gửi được. N/A nếu thiếu điều kiện."),
        ],
    },
    {
        "name": "7. Course Discussion (IT - Discuss)",
        "role": "LEARNER + CONTENT_MANAGER",
        "account": "Learner + content.manager@englishlab.vn",
        "goal": "List thảo luận; tạo/report cần enroll; CM xem reports.",
        "tool": "Postman",
        "cases": [
            ("IT_DISCUSS_01", "Tạo thảo luận", "POST /api/student/online-courses/{courseId}/discussions {title, content}", "200 nếu đã enroll", "Passed nếu tạo được. N/A nếu 400 chưa enroll."),
            ("IT_DISCUSS_02", "List thảo luận", "GET /api/online-courses/{courseId}/discussions", "HTTP 200", "Passed nếu 200."),
            ("IT_DISCUSS_03/04", "Report thread", "POST /api/student/online-courses/discussions/{threadId}/reports", "HTTP 200", "Passed nếu report được. N/A nếu không có thread."),
            ("IT_DISCUSS_05", "CM moderation", "GET /api/content-manager/discussion-reports (token CM)", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "8. CM Online Courses (IT - Content)",
        "role": "CONTENT_MANAGER",
        "account": "content.manager@englishlab.vn / Password123!",
        "goal": "CM quản lý khóa học online.",
        "tool": "Postman",
        "cases": [
            ("IT_CONTENT_01..04", "List khóa CM", "GET /api/content-manager/online-courses", "HTTP 200 danh sách", "Passed nếu 200. (Create/publish chi tiết: làm thêm POST/PUT theo UI nếu cô yêu cầu sâu)."),
        ],
    },
    {
        "name": "9. Packages & Bundles (IT - Package)",
        "role": "CONTENT_MANAGER",
        "account": "content.manager@englishlab.vn",
        "goal": "Quản lý package/bundle.",
        "tool": "Postman",
        "cases": [
            ("IT_PACKAGE_01..03", "List packages", "GET /api/content-manager/packages", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "10. Curriculum & Banks (IT - Curriculum)",
        "role": "CONTENT_MANAGER",
        "account": "content.manager@englishlab.vn",
        "goal": "Chương trình, ngân hàng bài, rubric.",
        "tool": "Postman",
        "cases": [
            ("IT_CURRICULUM_01/05", "Curriculum programs", "GET /api/content-manager/curriculum-programs", "HTTP 200", "Passed nếu 200."),
            ("IT_CURRICULUM_02", "Exercise/Assessment bank", "GET /api/content-manager/exercise-bank hoặc /assessment-bank", "HTTP 200", "Passed nếu 200."),
            ("IT_CURRICULUM_03", "Learning paths", "GET /api/content-manager/learning-paths", "HTTP 200 hoặc N/A nếu chưa có API", "Ghi đúng thực tế."),
            ("IT_CURRICULUM_04", "Rubrics", "GET /api/content-manager/rubrics", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "11. Enrollment Requests (IT - EnrollReq)",
        "role": "LEARNER + STAFF",
        "account": "Learner + staff@englishlab.vn",
        "goal": "HV gửi form tư vấn; Staff xem danh sách.",
        "tool": "Postman",
        "cases": [
            ("IT_ENROLLREQ_01/04", "HV tạo request", "1) GET /api/course-offerings lấy id\n2) POST /api/student/course-enrollment-requests\nBody: courseOfferingId, contactName, contactEmail, contactPhone, consultationTrack\n3) GET /api/student/course-enrollment-requests/my", "200 nếu tạo mới được", "Passed nếu tạo + list /my OK. N/A nếu báo đã có form đang xử lý."),
            ("IT_ENROLLREQ_02/03/05", "Staff list", "GET /api/staff/enrollment-requests (token STAFF)", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "12. TM Classroom Ops (IT - Classroom)",
        "role": "TRAINING_MANAGER (+ public)",
        "account": "training.manager@englishlab.vn / Password123!",
        "goal": "Offerings, registrations, waitlist reorder.",
        "tool": "Postman",
        "cases": [
            ("IT_CLASS_01", "Public offerings", "GET /api/classroom-offerings", "HTTP 200", "Passed nếu 200."),
            ("IT_CLASS_02/08", "TM list lớp", "GET /api/training-manager/classrooms", "HTTP 200", "Passed nếu 200."),
            ("IT_CLASS_03", "TM chi tiết", "GET /api/training-manager/classrooms/{id}", "HTTP 200", "Passed nếu 200."),
            ("IT_CLASS_04/06", "Registrations", "GET /api/training-manager/classrooms/registrations?classroomOfferingId={id}", "HTTP 200", "Passed nếu 200."),
            ("IT_CLASS_05", "Reorder waitlist", "1) GET registrations?status=WAITLIST&classroomOfferingId=\n2) Nếu ≥2 HV: PUT /api/training-manager/classrooms/{id}/waitlist/order {\"enrollmentIds\":[...]}", "HTTP 200", "Passed nếu reorder OK. N/A nếu <2 HV waitlist."),
            ("IT_CLASS_07", "Xem lớp trước khi gán GV", "GET classroom detail", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "13. Learner Classroom (IT - LearnerCls)",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com (cần đã được gán lớp)",
        "goal": "Lớp của tôi, session, homework, materials, gradebook.",
        "tool": "Postman",
        "cases": [
            ("IT_LEARNERCLS_01", "My classrooms", "GET /api/student/classrooms/my-classrooms", "HTTP 200 (có thể [] nếu chưa gán)", "Passed nếu 200."),
            ("IT_LEARNERCLS_02", "Sessions", "GET /api/student/classrooms/{id}/sessions", "HTTP 200", "Passed nếu 200. N/A nếu không có lớp."),
            ("IT_LEARNERCLS_03/05", "Homework", "GET /api/student/classrooms/{id}/homework", "HTTP 200", "Passed nếu 200."),
            ("IT_LEARNERCLS_04", "Materials", "GET /api/student/classrooms/{id}/materials", "HTTP 200", "Passed nếu 200."),
            ("IT_LEARNERCLS_06", "Gradebook của tôi", "GET /api/student/classrooms/{id}/gradebook/me", "HTTP 200/204", "Passed nếu 200/204."),
        ],
    },
    {
        "name": "14. Teacher Operations (IT - Teacher)",
        "role": "TEACHER",
        "account": "classroom.teacher1@englishlab.vn / Password123!",
        "goal": "Lớp được assign, homework, điểm danh, gradebook, change request.",
        "tool": "Postman",
        "cases": [
            ("IT_TEACH_01/06", "Lớp assigned", "GET /api/teacher/classrooms/assigned", "HTTP 200", "Passed nếu 200."),
            ("IT_TEACH_02", "Homework", "GET /api/teacher/classrooms/{id}/homework", "HTTP 200", "Passed nếu 200."),
            ("IT_TEACH_03", "Điểm danh theo session", "1) GET .../sessions\n2) GET /api/teacher/classrooms/sessions/{sessionId}/attendance", "HTTP 200", "Passed nếu 200. N/A nếu chưa có session."),
            ("IT_TEACH_04", "Gradebook lớp", "GET /api/teacher/classrooms/{id}/gradebook", "HTTP 200", "Passed nếu 200."),
            ("IT_TEACH_05", "Change requests của tôi", "GET /api/teacher/classrooms/requests/mine", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "15. Classroom Quiz (IT - Quiz)",
        "role": "TEACHER + LEARNER",
        "account": "Teacher + Learner",
        "goal": "List quiz; delete destructive → thường N/A trên demo.",
        "tool": "Postman",
        "cases": [
            ("IT_QUIZ_01/02", "Teacher list quiz", "GET /api/teacher/classrooms/{offeringId}/quizzes", "HTTP 200", "Passed nếu 200."),
            ("IT_QUIZ_03", "Learner list quiz", "GET /api/student/classrooms/quizzes", "HTTP 200", "Passed nếu 200."),
            ("IT_QUIZ_04", "Xóa quiz", "DELETE /api/teacher/quizzes/{id} (chỉ quiz test tự tạo)", "HTTP 204", "N/A nếu không muốn xóa data demo."),
        ],
    },
    {
        "name": "16. Assessment & Placement (IT - Assess)",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com",
        "goal": "Placement, mock test; assessment khóa cần enroll.",
        "tool": "Postman",
        "cases": [
            ("IT_ASSESS_01", "Placement hiện tại", "GET /api/student/placement-tests/current", "HTTP 200", "Passed nếu 200."),
            ("IT_ASSESS_02", "Submit placement thiếu đáp án (negative)", "POST /api/student/placement-tests/current/submit {\"answers\":[]}", "HTTP 400 hoặc 200 tùy trạng thái", "Passed nếu hành vi hợp lý (reject/accept). Failed nếu 500."),
            ("IT_ASSESS_03/05", "Assessments theo khóa", "GET /api/student/online-courses/{courseId}/assessments", "200 nếu enroll", "Passed nếu 200. N/A nếu 400 chưa enroll."),
            ("IT_ASSESS_04/06", "Mock tests", "GET /api/student/mock-tests", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "17. Support Tickets (IT - Support)",
        "role": "LEARNER + MANAGER",
        "account": "Learner + classroom.manager@englishlab.vn",
        "goal": "Tạo/list ticket; staff xem.",
        "tool": "Postman",
        "cases": [
            ("IT_SUPPORT_01", "Tạo ticket", "POST /api/student/support-tickets\nBody: subject (≥5 ký tự), category=TECHNICAL|ACCOUNT|..., message (≥10 ký tự)", "HTTP 200/201", "Passed nếu tạo được."),
            ("IT_SUPPORT_02", "List của tôi", "GET /api/student/support-tickets", "HTTP 200", "Passed nếu 200."),
            ("IT_SUPPORT_03", "Manager list", "GET /api/manager/support-tickets", "HTTP 200", "Passed nếu 200."),
            ("IT_SUPPORT_04", "Body rỗng (negative)", "POST {} ", "HTTP 400", "Passed nếu validation lỗi."),
        ],
    },
    {
        "name": "18. Administration (IT - Admin)",
        "role": "ADMIN",
        "account": "classroom.admin@englishlab.vn / Password123!",
        "goal": "Users, audit, config.",
        "tool": "Postman",
        "cases": [
            ("IT_ADMIN_01/02", "List users", "GET /api/admin/users", "HTTP 200", "Passed nếu 200."),
            ("IT_ADMIN_03", "Audit logs", "GET /api/admin/audit-logs", "HTTP 200", "Passed nếu 200."),
            ("IT_ADMIN_04", "System config", "GET /api/admin/system/config", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "19. Lark Meetings (IT - Lark)",
        "role": "Public webhook + TM",
        "account": "TM cho sync",
        "goal": "Webhook challenge; sync recording.",
        "tool": "Postman",
        "cases": [
            ("IT_LARK_01/02", "Webhook", "POST /api/lark/events body challenge/url_verification", "200 nếu challenge đúng format hệ thống; có thể 400 nếu thiếu chữ ký", "Passed nếu 200. N/A nếu 400 do cấu hình Lark."),
            ("IT_LARK_03", "Sync Lark session", "POST /api/training-manager/recordings/sessions/{sessionId}/sync-lark", "200 nếu session thật", "N/A nếu session không tồn tại."),
        ],
    },
    {
        "name": "20. Infrastructure (IT - Infra)",
        "role": "TRAINING_MANAGER",
        "account": "training.manager@englishlab.vn",
        "goal": "Campus, room, session template.",
        "tool": "Postman",
        "cases": [
            ("IT_INFRA_01", "Campuses", "GET /api/training-manager/infrastructure/campuses", "HTTP 200", "Passed nếu 200."),
            ("IT_INFRA_02", "Rooms", "GET /api/training-manager/infrastructure/rooms", "HTTP 200", "Passed nếu 200."),
            ("IT_INFRA_03", "Session templates", "GET /api/training-manager/infrastructure/session-templates", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "21. Reports & Revenue (IT - Report)",
        "role": "TM + CM",
        "account": "training.manager@englishlab.vn + content.manager@englishlab.vn",
        "goal": "Dashboard / doanh thu.",
        "tool": "Postman",
        "cases": [
            ("IT_REPORT_01", "TM dashboard", "GET /api/training-manager/dashboard", "HTTP 200", "Passed nếu 200."),
            ("IT_REPORT_02", "Revenue analytics", "GET /api/content-manager/revenue/analytics", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "22. Classroom Proposals (IT - Proposal)",
        "role": "STAFF",
        "account": "staff@englishlab.vn",
        "goal": "Đề xuất mở lớp.",
        "tool": "Postman",
        "cases": [
            ("IT_PROPOSAL_01..03", "List proposals", "GET /api/staff/classroom-proposals", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "23. Attendance Disputes (IT - Dispute)",
        "role": "LEARNER + TEACHER",
        "account": "Learner + Teacher",
        "goal": "Khiếu nại điểm danh.",
        "tool": "Postman",
        "cases": [
            ("IT_DISPUTE_01", "HV xem disputes", "GET /api/student/attendance/disputes", "HTTP 200 (có thể [])", "Passed nếu 200."),
            ("IT_DISPUTE_02/03", "GV pending", "GET /api/teacher/attendance-disputes/pending", "HTTP 200", "Passed nếu 200."),
        ],
    },
    {
        "name": "24. Learning Notes (IT - Notes)",
        "role": "LEARNER",
        "account": "0386852628z@gmail.com",
        "goal": "Ghi chú bài học.",
        "tool": "Postman",
        "cases": [
            ("IT_NOTES_01/02", "List notes", "GET /api/student/learning/notes", "HTTP 200", "Passed nếu 200."),
        ],
    },
]


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HƯỚNG DẪN TỰ TEST INTEGRATION TEST\nDỰ ÁN ENGLISHLAB (SEP490_G23)")
    set_run_font(r, size=18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Dành cho sinh viên tự thực hiện và ghi kết quả vào file Excel Integration Test")
    set_run_font(r, size=12)

    p(doc, "Người thực hiện / Tester: phongdx", bold=True)
    p(doc, "Môi trường: Backend http://localhost:8080 · Tool: Postman (hoặc Thunder Client) · File ghi kết quả: SEP490_G23_Report5.2_Integration Test_COMPLETED_HONEST_FORMATTED.xlsx")

    add_heading_vn(doc, "A. Mục đích tài liệu", 1)
    p(doc, "Tài liệu này hướng dẫn bạn TỰ CHẠY từng module Integration Test theo đúng sheet Excel, rồi tự ghi Passed / Failed / N/A vào cột Round 1–3. Không cần viết code; chủ yếu dùng Postman gọi API.")

    add_heading_vn(doc, "B. Chuẩn bị trước khi test", 1)
    numbered(doc, "Bật backend Spring Boot (cổng 8080) và PostgreSQL đã seed dữ liệu demo.")
    numbered(doc, "Cài Postman → tạo Environment biến baseUrl = http://localhost:8080")
    numbered(doc, "Mở file Excel Integration Test (bản HONEST_FORMATTED).")
    numbered(doc, "Chuẩn bị tài khoản demo (mật khẩu thường là Password123!):")
    bullet(doc, "LEARNER: 0386852628z@gmail.com")
    bullet(doc, "TEACHER: classroom.teacher1@englishlab.vn")
    bullet(doc, "TRAINING_MANAGER: training.manager@englishlab.vn")
    bullet(doc, "STAFF: staff@englishlab.vn")
    bullet(doc, "MANAGER: classroom.manager@englishlab.vn")
    bullet(doc, "CONTENT_MANAGER: content.manager@englishlab.vn")
    bullet(doc, "ADMIN: classroom.admin@englishlab.vn")

    add_heading_vn(doc, "C. Quy trình test 1 test case (áp dụng mọi module)", 1)
    numbered(doc, "Mở đúng sheet module trong Excel (ví dụ IT - Auth).")
    numbered(doc, "Đọc Pre-conditions → chuẩn bị role/token/data.")
    numbered(doc, "Làm lần lượt theo Procedure trong tài liệu này (hoặc cột Procedure trên Excel).")
    numbered(doc, "So sánh với Expected Results.")
    numbered(doc, "Ghi vào Excel:")
    bullet(doc, "Round 1 (và 2, 3 nếu retest): Passed / Failed / N/A")
    bullet(doc, "Test date: ngày bạn chạy")
    bullet(doc, "Tester: tên bạn (phongdx)")
    bullet(doc, "Note: HTTP status hoặc lý do N/A ngắn gọn")

    add_heading_vn(doc, "D. Cách chấm Passed / Failed / N/A", 1)
    add_table(
        doc,
        ["Kết quả", "Khi nào ghi"],
        [
            ["Passed", "Đúng expected: happy-path thành công HOẶC negative test bị từ chối đúng (401/400)."],
            ["Failed", "Sai expected, lỗi 500, endpoint không tồn tại, dữ liệu sai."],
            ["N/A", "Thiếu điều kiện môi trường (chưa enroll khóa, thiếu OTP/mail, không muốn xóa data demo, waitlist < 2 HV…)."],
        ],
    )
    p(doc, "Lưu ý với cô/giảng viên: N/A không phải giấu Fail. Đó là case chưa đủ precondition để khẳng định full happy-path.")

    add_heading_vn(doc, "E. Mẹo dùng Postman nhanh", 1)
    bullet(doc, "Tạo request Login → Tests script lưu token: pm.environment.set('token', pm.response.json().accessToken);")
    bullet(doc, "Request sau: Header Authorization = Bearer {{token}}")
    bullet(doc, "Body chọn raw JSON.")
    bullet(doc, "Với OTP verify/reset: có thể xem bảng auth_tokens trong PostgreSQL (cột token, type).")

    add_heading_vn(doc, "F. Hướng dẫn chi tiết theo từng MODULE", 1)
    p(doc, "Có 24 module tương ứng 24 sheet IT - …. Làm lần lượt; mỗi buổi nên làm 1–2 module.")

    for m in MODULES:
        add_heading_vn(doc, m["name"], 2)
        p(doc, f"Vai trò: {m['role']}", bold=True)
        p(doc, f"Tài khoản: {m['account']}")
        p(doc, f"Mục tiêu: {m['goal']}")
        p(doc, f"Công cụ: {m['tool']}")
        p(doc, "Các case cần làm:", bold=True)
        rows = []
        for case in m["cases"]:
            cid, ten, cac_buoc, mong_doi, cham = case
            rows.append([cid, ten, cac_buoc, mong_doi, cham])
        add_table(doc, ["Mã case", "Tên kiểm thử", "Các bước làm (Postman)", "Kết quả mong đợi", "Cách chấm"], rows)

    add_heading_vn(doc, "G. Thứ tự nên tự test (khuyến nghị)", 1)
    numbered(doc, "Auth → User → Notif (nền tảng)")
    numbered(doc, "Classroom → LearnerCls → Teacher → Quiz → Dispute → Notes")
    numbered(doc, "Commerce → Payment")
    numbered(doc, "Course → Discuss → Assess")
    numbered(doc, "EnrollReq → Content/Package/Curriculum → Infra → Report → Proposal → Support → Admin → Lark")

    add_heading_vn(doc, "H. Checklist trước khi nộp / thuyết trình", 1)
    bullet(doc, "Đã tự chạy và ghi Round tối thiểu cho module cô yêu cầu demo (Auth + Classroom + 1 module khác).")
    bullet(doc, "Tester = tên thật của bạn.")
    bullet(doc, "Note có ghi HTTP hoặc lý do N/A.")
    bullet(doc, "Test Statistics mở bằng Excel để công thức tự tính Passed/Failed/N/A.")
    bullet(doc, "Biết giải thích 1 case Passed và 1 case N/A.")

    add_heading_vn(doc, "I. Câu trả lời ngắn nếu cô hỏi", 1)
    p(doc, "Em tự test bằng Postman theo Procedure trên từng sheet IT, đối chiếu Expected Results, rồi ghi Round 1–3. Passed là đạt; N/A là thiếu precondition trên môi trường demo; Failed là sai expected hoặc lỗi hệ thống.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    PROJ.mkdir(parents=True, exist_ok=True)
    import shutil
    shutil.copy2(OUT, PROJ / OUT.name)
    # also markdown mirror for quick read
    md = PROJ / "Huong_dan_tu_test_Integration_Test_EnglishLab.md"
    lines = [
        "# Hướng dẫn tự test Integration Test – EnglishLab\n",
        "File Word đầy đủ nằm tại Downloads. Dưới đây là bản rút gọn.\n",
        "## Quy trình 1 case\n1. Đọc Pre-condition\n2. Gọi API Postman\n3. So Expected\n4. Ghi Round/Tester/Note\n",
        "## Chấm điểm\n- Passed: đúng expected\n- Failed: sai/500\n- N/A: thiếu precondition\n",
    ]
    for m in MODULES:
        lines.append(f"## {m['name']}\n")
        lines.append(f"- Role: {m['role']}\n- Account: {m['account']}\n- Goal: {m['goal']}\n")
        for cid, ten, buoc, mong, cham in m["cases"]:
            lines.append(f"### {cid} – {ten}\n")
            lines.append(f"**Bước:**\n```\n{buoc}\n```\n")
            lines.append(f"**Mong đợi:** {mong}\n")
            lines.append(f"**Chấm:** {cham}\n")
    md.write_text("\n".join(lines), encoding="utf-8")
    print("WORD", OUT)
    print("MD", md)


if __name__ == "__main__":
    build()
