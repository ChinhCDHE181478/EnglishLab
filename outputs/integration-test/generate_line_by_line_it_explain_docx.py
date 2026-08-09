# -*- coding: utf-8 -*-
"""Word: giải thích từng dòng code của 26 file Integration Test (src/test/.../it)."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

JAVA_DIR = Path(r"D:\EngLishLab\EnglishLab\backend\src\test\java\fu\sap490\g23\backend\it")
OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Giai_thich_tung_dong_code_26_file_IT.docx"

# Thứ tự đọc: helper trước, Auth mẫu, rồi các module
ORDER = [
    "ItSupport.java",
    "AuthIT.java",
    "AuthOtpIT.java",
    "UserIT.java",
    "NotificationIT.java",
    "CommerceIT.java",
    "PaymentIT.java",
    "OnlineCourseIT.java",
    "DiscussionIT.java",
    "ContentManagerCourseIT.java",
    "PackageIT.java",
    "CurriculumIT.java",
    "EnrollmentRequestIT.java",
    "TrainingManagerClassroomIT.java",
    "StudentClassroomIT.java",
    "TeacherClassroomIT.java",
    "ClassroomQuizIT.java",
    "AssessmentIT.java",
    "SupportTicketIT.java",
    "AdminIT.java",
    "LarkIT.java",
    "InfrastructureIT.java",
    "ReportIT.java",
    "ClassroomProposalIT.java",
    "AttendanceDisputeIT.java",
    "LearningNotesIT.java",
]


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(2)


def explain_line(code: str, fname: str) -> str:
    s = code.strip()
    if not s:
        return "(dòng trống — chỉ để dễ đọc)"
    if s.startswith("//") or s.startswith("*") or s.startswith("/**") or s == "*/":
        return "Comment / ghi chú cho người đọc — không chạy."
    if s.startswith("package "):
        return "Khai báo package: class thuộc package it (Integration Test)."
    if s.startswith("import static "):
        if "ItSupport" in s:
            return "Import static helper (PASSWORD, LEARNER, login, bearer…) — gọi ngắn không cần ItSupport."
        if "MockMvcRequestBuilders" in s:
            return "Import sẵn get/post/put/patch/delete để viết mockMvc.perform(post(...))."
        if "MockMvcResultMatchers" in s:
            return "Import sẵn status(), jsonPath()… để andExpect(...)."
        return f"Import static: dùng thẳng tên method từ class khác. ({s[14:60]}…)" if len(s) > 74 else f"Import static: {s[14:]}"
    if s.startswith("import "):
        mapping = [
            ("DisplayName", "Annotation đặt tên test = mã IT_* trên Excel."),
            ("Test", "Annotation đánh dấu method là 1 test case JUnit."),
            ("Autowired", "Spring tự tiêm bean (MockMvc, JdbcTemplate…) vào field."),
            ("AutoConfigureMockMvc", "Bật MockMvc trong @SpringBootTest (Spring Boot 4)."),
            ("SpringBootTest", "Chạy gần như cả app Spring — bean Controller/Service/Repo thật nối nhau."),
            ("MediaType", "Kiểu nội dung JSON khi gửi body."),
            ("MockMvc", "Công cụ giả lập HTTP gọi vào Controller (không cần Postman)."),
            ("MvcResult", "Giữ response để đọc body (ví dụ lấy accessToken)."),
            ("JdbcTemplate", "Chạy SQL trực tiếp (đọc OTP từ auth_tokens)."),
            ("JsonNode", "Đọc JSON response (Jackson)."),
            ("ObjectMapper", "Parse chuỗi JSON ↔ object."),
            ("UUID", "Tạo email ngẫu nhiên tránh trùng khi register."),
            ("Assertions", "assertTrue / assertEquals để tự khẳng định điều kiện."),
        ]
        for key, msg in mapping:
            if key in s:
                return msg
        return "Import thư viện cần dùng trong file."

    if s.startswith("@SpringBootTest"):
        return "Bật Integration Test: load Spring context → xác minh Controller–Service–Repository."
    if s.startswith("@AutoConfigureMockMvc"):
        return "Cấu hình sẵn MockMvc để test HTTP trong context Spring."
    if s.startswith("@Autowired"):
        return "Spring inject dependency vào field ngay dưới (MockMvc / JdbcTemplate)."
    if s.startswith("@DisplayName"):
        m = re.search(r'"([^"]+)"', s)
        return f"Tên hiển thị khi chạy test = map Excel: « {m.group(1) if m else s} »."
    if s.startswith("@Test"):
        return "Đánh dấu method này là 1 Integration Test case — JUnit sẽ chạy."

    if s.startswith("public final class ItSupport") or s.startswith("public class ItSupport"):
        return "Class helper dùng chung — không có @Test nên không tự chạy."
    if re.match(r"public class \w+IT", s) or re.match(r"class \w+IT", s):
        return f"Class IT của module — chứa các @Test map sheet Excel. ({fname})"
    if s == "private ItSupport() {":
        return "Constructor private: không cho new ItSupport() — chỉ dùng static."
    if s == "}":
        return "Đóng khối lệnh (class/method/if)."
    if s == "};" or s.endswith("};"):
        return "Kết thúc khai báo (text block / statement)."

    if "PASSWORD" in s and "static final" in s:
        return "Hằng mật khẩu demo dùng chung trong IT."
    if "static final String" in s and "=" in s:
        return "Hằng email tài khoản demo theo role (LEARNER/TEACHER/TM…)."
    if "ObjectMapper MAPPER" in s:
        return "ObjectMapper dùng parse JSON login response."

    if "String login(MockMvc" in s:
        return "Hàm login: gọi POST /api/auth/login → trả accessToken (chuỗi JWT)."
    if "bearer(String" in s:
        return "Ghép header Authorization dạng « Bearer <token> »."
    if "mapper()" in s and "ObjectMapper" in s:
        return "Cho class IT khác lấy chung ObjectMapper."

    if "mockMvc.perform" in s:
        return "Gửi HTTP request giả lập vào Controller (điểm vào Integration)."
    if s.startswith(".contentType"):
        return "Khai báo body là JSON."
    if s.startswith(".content("):
        return "Gắn chuỗi JSON body vào request."
    if s.startswith(".header(\"Authorization\""):
        return "Gắn JWT vào header — request đi qua Security filter."
    if s.startswith(".andExpect(status()"):
        if "2xx" in s or "isOk" in s or "isCreated" in s:
            return "Khẳng định HTTP thành công (Expected Excel)."
        if "4xx" in s or "isUnauthorized" in s or "isForbidden" in s or "isBadRequest" in s:
            return "Khẳng định HTTP lỗi client — dùng cho negative test (vẫn có thể Passed)."
        return "Khẳng định HTTP Status theo Expected."
    if s.startswith(".andExpect(jsonPath"):
        return "Khẳng định field JSON trong response (ví dụ email, checkoutUrl)."
    if s.startswith(".andReturn()"):
        return "Lấy MvcResult để đọc body thủ công (token, id…)."
    if "getResponse().getContentAsString" in s:
        return "Đọc body response dạng chuỗi JSON."
    if "getResponse().getStatus" in s:
        return "Đọc mã HTTP Status dạng số."
    if "json.get(\"accessToken\")" in s or "accessToken" in s and "asText" in s:
        return "Lấy JWT từ JSON login."
    if "UUID.randomUUID()" in s:
        return "Tạo phần ngẫu nhiên — email register không trùng."
    if "formatted(" in s or ".formatted(" in s:
        return "Điền biến vào text block JSON (email/password…)."
    if s.startswith('String body = """') or s == 'String body = """':
        return "Bắt đầu text block JSON (Java 15+)."
    if s.startswith('"""'):
        return "Kết thúc / mở text block."
    if s.startswith("String token = login"):
        return "Login lấy JWT trước khi gọi API cần quyền."
    if s.startswith("String email ="):
        return "Chuẩn bị email dùng trong case."
    if "jdbcTemplate.queryForObject" in s or "latestOtp" in s:
        return "Đọc OTP từ DB (auth_tokens) — phục vụ IT_AUTH_03/09."
    if "JdbcTemplate" in s and "private" in s:
        return "Field SQL helper — Spring inject."
    if "MockMvc mockMvc" in s:
        return "Field MockMvc — Spring inject; dùng trong mọi @Test."
    if s.startswith("void it") or re.match(r"void \w+\(", s):
        return "Method test: 1 method ≈ 1 Test Case ID trên Excel."
    if "assertTrue" in s or "Assertions." in s:
        return "Tự kiểm tra điều kiện bổ sung (ví dụ status < 500)."
    if "firstCourseId" in s or "items.get(0)" in s:
        return "Lấy id khóa học đầu tiên từ API public để dùng tiếp."
    if "status < 500" in s:
        return "Chấp nhận 2xx/4xx nhưng Fail nếu server 5xx."
    if s.startswith("return "):
        return "Trả kết quả ra caller."
    if "MediaType.APPLICATION_JSON" in s:
        return "Content-Type: application/json."

    # HTTP path hints
    if "/api/" in s and ("post(" in s or "get(" in s or "put(" in s or "patch(" in s or "delete(" in s):
        return "Chọn HTTP method + đường dẫn API — khớp Procedure Excel « Call … via MockMvc »."

    if s.startswith("{"):
        return "Mở khối lệnh."
    if "throws Exception" in s:
        return "Cho phép ném Exception (MockMvc/IO) — JUnit sẽ bắt nếu fail."

    return "Dòng lệnh Java trong test — xem ngữ cảnh method phía trên."


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.6)
        s.right_margin = Cm(1.6)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "GIẢI THÍCH TỪNG DÒNG CODE\n"
        "26 FILE INTEGRATION TEST (MockMvc / Spring Boot Test)\n"
        "backend/src/test/java/fu/sap490/g23/backend/it/\n"
        "EnglishLab – SEP490_G23"
    )
    font(r, size=14, bold=True)

    p(doc, "Cách đọc: mỗi file → từng dòng code (trái) → ý nghĩa tiếng Việt (phải/ bên dưới).", bold=True)
    p(
        doc,
        "Nhắc nhanh: @SpringBootTest + MockMvc = IT xác minh Controller–Service–Repository. "
        "ItSupport không có @Test. Mỗi @DisplayName(IT_*) map Excel.",
    )

    h(doc, "0. Thuật ngữ gặp lại nhiều lần", 1)
    rows = [
        ("@SpringBootTest", "Load context Spring — bean thật nối nhau"),
        ("@AutoConfigureMockMvc", "Cho phép MockMvc"),
        ("mockMvc.perform", "Gửi HTTP vào Controller"),
        ("andExpect", "Khẳng định Expected"),
        ("@DisplayName", "Mã IT trên Excel"),
        ("Bearer token", "JWT sau login"),
        ("Negative test", "Cố tình sai — expect 4xx vẫn Passed"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    for i, hd in enumerate(["Thuật ngữ", "Nghĩa"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(hd)
        font(rr, 10, True)
    for ri, (a, b) in enumerate(rows):
        for ci, val in enumerate((a, b)):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(val)
            font(rr, 9)
    doc.add_paragraph()

    files = []
    for name in ORDER:
        path = JAVA_DIR / name
        if path.exists():
            files.append(path)
    for path in sorted(JAVA_DIR.glob("*.java")):
        if path not in files:
            files.append(path)

    h(doc, "1. Giải thích từng file / từng dòng", 1)
    p(doc, f"Tổng số file: {len(files)}.")

    for fi, path in enumerate(files, 1):
        text = path.read_text(encoding="utf-8")
        lines = text.splitlines()
        h(doc, f"FILE {fi}/{len(files)}: {path.name}", 2)
        p(doc, f"Đường dẫn: backend/src/test/java/fu/sap490/g23/backend/it/{path.name}", bold=True)

        if path.name == "ItSupport.java":
            p(
                doc,
                "Vai trò: helper dùng chung (login, tài khoản demo). KHÔNG chạy file này bằng Run Test.",
                bold=True,
            )
        elif path.name == "AuthIT.java":
            p(
                doc,
                "Vai trò: mẫu chuẩn IT Auth — đọc kỹ file này rồi các file *IT khác cùng pattern.",
                bold=True,
            )
        else:
            p(doc, "Vai trò: class Integration Test map sheet Excel tương ứng; chạy bằng mvnw -Dtest=TênClass.")

        # line-by-line table in chunks to keep Word usable
        # Use paragraphs: "Lxx | code" then explanation — tables of 100+ rows are heavy
        for ln, code in enumerate(lines, 1):
            para = doc.add_paragraph()
            r1 = para.add_run(f"L{ln}: ")
            font(r1, 9, True, color=RGBColor(0x1F, 0x4E, 0x79))
            r2 = para.add_run(code if code.strip() else "␣")
            font(r2, 8, name="Consolas")
            para.paragraph_format.space_after = Pt(0)

            exp = explain_line(code, path.name)
            para2 = doc.add_paragraph()
            r3 = para2.add_run("→ " + exp)
            font(r3, 9)
            para2.paragraph_format.space_after = Pt(4)
            para2.paragraph_format.left_indent = Cm(0.3)

        p(doc, "—" * 36)

    h(doc, "2. Cách chạy lại sau khi đã hiểu code", 1)
    p(doc, "cd backend rồi:")
    para = doc.add_paragraph()
    rr = para.add_run('.\\mvnw.cmd "-Dtest=AuthIT" test\n.\\mvnw.cmd "-Dtest=*IT" test')
    font(rr, 9, name="Consolas")

    h(doc, "3. Nói với cô (sau khi đọc code)", 1)
    p(
        doc,
        "Mỗi class *IT dùng @SpringBootTest + MockMvc; từng dòng mockMvc.perform là điểm vào Controller, "
        "Spring nối Service–Repository. @DisplayName khớp Test Case ID Excel. ItSupport chỉ hỗ trợ login.",
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    shutil.copy2(out, PROJ / OUT_NAME)
    print(out, "files", len(files), "approx_lines", sum(len(p.read_text(encoding='utf-8').splitlines()) for p in files))
    return out


if __name__ == "__main__":
    build()
