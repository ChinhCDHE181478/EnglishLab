# -*- coding: utf-8 -*-
"""Sinh file kịch bản thuyết trình (MD + DOCX) cho từng Test Case ID."""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
from full_modules import MODULES  # noqa: E402
from generate_testcases_word_per_module import VN_GOAL  # noqa: E402

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
EXCEL = HERE / "SEP490_G23_Report5.2_Integration Test_COMPLETED_HONEST_FORMATTED.xlsx"
MD_NAME = "Kich_ban_thuyet_trinh_tung_Test_Case.md"
DOCX_NAME = "Kich_ban_thuyet_trinh_tung_Test_Case.docx"

# sheet Excel -> class Java trong src/test/.../it
SHEET_TO_CLASS = {
    "IT - Auth": "AuthIT (riêng IT_AUTH_03/09 nằm ở AuthOtpIT)",
    "IT - User": "UserIT",
    "IT - Notif": "NotificationIT",
    "IT - Commerce": "CommerceIT",
    "IT - Payment": "PaymentIT",
    "IT - Course": "OnlineCourseIT",
    "IT - Discuss": "DiscussionIT",
    "IT - Content": "ContentManagerCourseIT",
    "IT - Package": "PackageIT",
    "IT - Curriculum": "CurriculumIT",
    "IT - EnrollReq": "EnrollmentRequestIT",
    "IT - Classroom": "TrainingManagerClassroomIT",
    "IT - LearnerCls": "StudentClassroomIT",
    "IT - Teacher": "TeacherClassroomIT",
    "IT - Quiz": "ClassroomQuizIT",
    "IT - Assess": "AssessmentIT",
    "IT - Support": "SupportTicketIT",
    "IT - Admin": "AdminIT",
    "IT - Lark": "LarkIT",
    "IT - Infra": "InfrastructureIT",
    "IT - Report": "ReportIT",
    "IT - Proposal": "ClassroomProposalIT",
    "IT - Dispute": "AttendanceDisputeIT",
    "IT - Notes": "LearningNotesIT",
}

ROLE_BY_SHEET = {
    "IT - Auth": "không cần token (trừ /me)",
    "IT - User": "LEARNER",
    "IT - Notif": "LEARNER",
    "IT - Commerce": "LEARNER",
    "IT - Payment": "LEARNER và MANAGER",
    "IT - Course": "public và LEARNER",
    "IT - Discuss": "LEARNER và CONTENT_MANAGER",
    "IT - Content": "CONTENT_MANAGER",
    "IT - Package": "CONTENT_MANAGER",
    "IT - Curriculum": "CONTENT_MANAGER",
    "IT - EnrollReq": "LEARNER và STAFF",
    "IT - Classroom": "TRAINING_MANAGER",
    "IT - LearnerCls": "LEARNER",
    "IT - Teacher": "TEACHER",
    "IT - Quiz": "TEACHER và LEARNER",
    "IT - Assess": "LEARNER",
    "IT - Support": "LEARNER và MANAGER",
    "IT - Admin": "ADMIN",
    "IT - Lark": "webhook công khai và TRAINING_MANAGER",
    "IT - Infra": "TRAINING_MANAGER",
    "IT - Report": "TRAINING_MANAGER và CONTENT_MANAGER",
    "IT - Proposal": "STAFF",
    "IT - Dispute": "LEARNER và TEACHER",
    "IT - Notes": "LEARNER",
}

PRE_VI = [
    (r"Database is available", "database đang chạy"),
    (r"mail sender is stubbed", "mail được stub"),
    (r"email is unused", "email chưa được dùng"),
    (r"A verified LEARNER account exists", "đã có tài khoản LEARNER đã xác thực"),
    (r"A verified user exists", "đã có user đã xác thực"),
    (r"An unverified user exists", "đã có user chưa xác thực"),
    (r"Valid JWT", "đã có JWT hợp lệ"),
    (r"valid JWT is available", "đã có JWT hợp lệ"),
    (r"Spring Security configuration is active", "Spring Security đang bật"),
    (r"already exists", "đã tồn tại sẵn"),
]


def vi_pre(pre: str) -> str:
    """Rút gọn Pre-condition sang tiếng Việt dễ đọc."""
    text = " ".join(x.strip() for x in str(pre).split("\n") if x.strip())
    for en, vi in PRE_VI:
        text = re.sub(en, vi, text, flags=re.IGNORECASE)
    text = text.rstrip(".")
    return text if text else "không có điều kiện đặc biệt"


def api_calls(proc: str) -> list[str]:
    """Lấy danh sách 'METHOD /api/...' xuất hiện trong Procedure."""
    found = []
    for m in re.finditer(r"\b(GET|POST|PUT|PATCH|DELETE)\s+(/api/[^\s,)]+)", proc, re.I):
        call = f"{m.group(1).upper()} {m.group(2).rstrip('.')}"
        if call not in found:
            found.append(call)
    return found


def java_classes(proc: str, components: str) -> list[str]:
    """Lấy tên Controller/Service/Repository nhắc trong Procedure + Components."""
    names = []
    for text in (proc, components):
        for m in re.finditer(r"\b([A-Z][A-Za-z0-9]*(?:Controller|Service|Repository|Filter))\b", str(text)):
            if m.group(1) not in names:
                names.append(m.group(1))
    return names


def flow_sentence(proc: str, components: str) -> str:
    """Câu mô tả luồng tích hợp Controller → Service → Repository."""
    names = java_classes(proc, components)
    controllers = [n for n in names if n.endswith("Controller")]
    services = [n for n in names if n.endswith("Service")]
    repos = [n for n in names if n.endswith("Repository")]
    filters = [n for n in names if n.endswith("Filter")]

    chain = []
    if controllers:
        chain.append(" / ".join(controllers[:2]))
    if services:
        chain.append(" / ".join(services[:2]))
    if repos:
        chain.append(" / ".join(repos[:2]))

    if not chain:
        return "request đi vào Controller rồi xuống Service và Repository trong cùng Spring context"

    flow = " → ".join(chain)
    if filters:
        flow += f" (đi qua {filters[0]} trước khi vào Controller)"
    return flow


def vi_expected(exp: str) -> list[str]:
    """Giữ nguyên câu tiếng Anh trên Excel để khớp tài liệu nộp cho cô."""
    out = []
    for line in str(exp).split("\n"):
        s = line.strip().rstrip(".")
        if s:
            out.append(s)
    return out[:4]


def vi_expected_gist(exp: str) -> str:
    """Một câu tiếng Việt tóm tắt kỳ vọng để nói miệng."""
    text = str(exp)
    low = text.lower()
    status = re.search(r"\b([1-5]xx|\d{3}(?:/\d{3})?)\b", text)
    code = status.group(1) if status else ""

    if re.search(r"\b(4xx|401|403|409|400|404|422)\b", text):
        head = f"API phải chặn và trả lỗi {code}" if code else "API phải chặn và trả lỗi client"
        tail = "dữ liệu trong database không bị thay đổi"
        if "unchanged" not in low and "no " not in low and "remains" not in low:
            tail = "không tạo ra dữ liệu sai"
        return f"{head}, {tail} — đây là case negative nên bị chặn đúng chính là Passed."

    head = f"API trả về {code} thành công" if code else "API trả về thành công"
    if re.search(r"insert|persist|create|saved|becomes true|updated", low):
        return f"{head} và dữ liệu tương ứng được ghi/cập nhật đúng trong database."
    if re.search(r"list|page|pagination|contains", low):
        return f"{head} và body chứa đúng dữ liệu cần lấy (kèm thông tin phân trang nếu có)."
    return f"{head} và nội dung phản hồi khớp với nghiệp vụ mong đợi."


def load_results() -> dict[str, tuple]:
    res: dict[str, tuple] = {}
    if not EXCEL.exists():
        return res
    wb = openpyxl.load_workbook(EXCEL, read_only=True, data_only=True)
    for sn in wb.sheetnames:
        if not sn.startswith("IT - "):
            continue
        ws = wb[sn]
        for row in ws.iter_rows(min_row=11, max_col=14, values_only=True):
            if row[0] and str(row[0]).startswith("IT_"):
                res[str(row[0])] = (row[5], row[8], row[11])
    wb.close()
    return res


def result_sentence(case_id: str, results: dict) -> str:
    r = results.get(case_id)
    if not r:
        return "Em ghi kết quả vào cột Round trên Excel sau khi chạy."
    final = r[2] or r[0]
    if final == "Passed":
        return "Kết quả trên Excel: **Passed** — actual khớp expected."
    if final == "N/A":
        return (
            "Kết quả trên Excel: **N/A** — môi trường demo thiếu tiền điều kiện nên chưa kết luận đủ, "
            "em ghi trung thực chứ không tính là Passed."
        )
    if final == "Failed":
        return (
            "Kết quả trên Excel: **Failed** ở vòng chạy cũ. Sau đó em bổ sung `AuthOtpIT` đọc OTP thật từ bảng "
            "`auth_tokens` nên case này chạy lại được; nếu chạy lại đạt thì cập nhật Round thành Passed."
        )
    return f"Kết quả trên Excel: {final}."


def short_line(case_id: str, goal: str, calls: list[str], flow: str, exp: str, results: dict) -> str:
    call = calls[0] if calls else "API tương ứng"
    status = re.search(r"\b([1-5]xx|\d{3}(?:/\d{3})?)\b", str(exp))
    expect = f" Mong đợi {status.group(1)}." if status else ""
    final = (results.get(case_id) or (None, None, None))[2] or "chưa chạy"
    tail = {
        "Passed": "Kết quả Passed.",
        "N/A": "Kết quả N/A vì môi trường demo thiếu tiền điều kiện.",
        "Failed": "Vòng cũ Failed, sau đó em bổ sung AuthOtpIT đọc OTP thật từ DB để chạy lại.",
    }.get(final, f"Kết quả {final}.")
    return f"{case_id}: {goal} Em gọi `{call}`, luồng {flow}.{expect} {tail}"


def case_block(case, sheet, components, results) -> list[str]:
    cid = case["id"]
    goal = VN_GOAL.get(cid, case["desc"])
    calls = api_calls(case["proc"])
    flow = flow_sentence(case["proc"], components)
    lines = [
        f"### {cid}",
        "",
        "**Nói ngắn (15 giây):**",
        "",
        f"> {short_line(cid, goal, calls, flow, case['exp'], results)}",
        "",
        "**Nói đầy đủ (45–60 giây):**",
        "",
        f"1. **Mục tiêu:** {goal}",
        f"2. **Tiền điều kiện:** {vi_pre(case['pre'])}.",
    ]
    if calls:
        lines.append(f"3. **Bước thực hiện:** gọi {', '.join('`' + c + '`' for c in calls[:3])} bằng MockMvc trong `@SpringBootTest`.")
    else:
        lines.append("3. **Bước thực hiện:** gọi API của chức năng này bằng MockMvc trong `@SpringBootTest`.")
    lines.append(f"4. **Luồng tích hợp:** {flow}.")
    lines.append(f"5. **Kết quả mong đợi (nói bằng tiếng Việt):** {vi_expected_gist(case['exp'])}")
    exps = vi_expected(case["exp"])
    if exps:
        lines.append("   Nguyên văn trên Excel:")
        for e in exps:
            lines.append(f"   - {e}")
    lines.append(f"6. **Kết quả thực tế:** {result_sentence(cid, results)}")
    lines.append("")
    return lines


def build_md(results: dict) -> str:
    parts: list[str] = []
    a = parts.append

    a("# Kịch bản thuyết trình Integration Test — từng Test Case\n\n")
    a("Dự án: **EnglishLab (SEP490_G23)** · Người trình bày: **phongdx**\n\n")
    a("Tài liệu này để **học thuộc ý** rồi nói với cô. Mỗi test case có 2 mức:\n\n")
    a("- **Nói ngắn (15 giây)** — dùng khi cô hỏi lướt nhiều case.\n")
    a("- **Nói đầy đủ (45–60 giây)** — dùng khi cô yêu cầu giải thích sâu 1 case.\n\n")
    a("---\n\n")

    a("## Phần A — Mở đầu (thuộc nguyên văn, 30 giây)\n\n")
    a(
        "> Dạ em trình bày phần Integration Test. Integration Test là kiểm thử **tương tác giữa các thành phần**: "
        "Controller gọi Service, Service gọi Repository và các Service gọi lẫn nhau. "
        "Em thiết kế test case trên file Excel Report 5.2 gồm 24 module và 111 case, "
        "và thực thi bằng **@SpringBootTest kết hợp MockMvc** trong `backend/src/test/java/fu/sap490/g23/backend/it/`. "
        "Postman em chỉ dùng để quan sát thủ công, không dùng thay Integration Test.\n\n"
    )

    a("## Phần B — 6 câu nền tảng (cô hay hỏi chen ngang)\n\n")
    a("| Cô hỏi | Em trả lời (thuộc ý) |\n|---|---|\n")
    a("| Integration Test là gì? | Kiểm thử tích hợp nhiều tầng chạy thật với nhau: Controller → Service → Repository, không mock hết như Unit Test. |\n")
    a("| Vì sao không chỉ dùng Postman? | Postman không nạp Spring context, không chạy trong JUnit nên không chứng minh được các bean thật nối nhau. |\n")
    a("| MockMvc là gì? | Công cụ của Spring Test giả lập HTTP đi thẳng vào Controller thật trong context test. |\n")
    a("| @SpringBootTest để làm gì? | Khởi động gần như toàn bộ ứng dụng để Controller, Service, Repository là bean thật. |\n")
    a("| Code map với Excel thế nào? | Mỗi method test có `@DisplayName` chứa đúng mã `IT_...` trên Excel. |\n")
    a("| Passed / Failed / N/A khác nhau? | Passed là actual khớp expected, kể cả negative bị chặn đúng; Failed là sai expected hoặc lỗi 500; N/A là thiếu tiền điều kiện nên chưa kết luận. |\n\n")

    a("## Phần C — Cách chạy để demo tại chỗ\n\n")
    a("```powershell\ncd D:\\EngLishLab\\EnglishLab\\backend\n.\\mvnw.cmd \"-Dtest=AuthIT\" test\n.\\mvnw.cmd \"-Dtest=*IT\" test\n```\n\n")
    a("> Em chạy `mvnw -Dtest=AuthIT test`; Maven in ra số test chạy và kết quả, em đối chiếu với cột Round trên Excel.\n\n")
    a("---\n\n")

    a("## Phần D — Kịch bản theo từng Test Case\n\n")

    total = 0
    for idx, m in enumerate(MODULES, 1):
        sheet = m["sheet"]
        cls = SHEET_TO_CLASS.get(sheet, "")
        a(f"## Module {idx}. {m['name']} — sheet `{sheet}`\n\n")
        a("**Câu dẫn cho cả module:**\n\n")
        a(
            f"> Module {m['name']} nằm ở sheet `{sheet}`, chạy bằng class `{cls}`. "
            f"Vai trò sử dụng: {ROLE_BY_SHEET.get(sheet, 'theo pre-condition trên Excel')}. "
            f"Thành phần tích hợp chính: {m.get('components', '')}.\n\n"
        )
        for g in m["groups"]:
            a(f"**Nhóm: {g['name']}**\n\n")
            for c in g["cases"]:
                total += 1
                a("\n".join(case_block(c, sheet, m.get("components", ""), results)))
                a("\n")
        a("---\n\n")

    a("## Phần E — Kết thúc (thuộc nguyên văn, 20 giây)\n\n")
    a(
        "> Tóm lại, bộ Integration Test của em gồm 111 test case trên 24 module, thiết kế trên Excel và thực thi bằng "
        "Spring Boot Test với MockMvc. Các case Passed là actual khớp expected, các case N/A là môi trường demo thiếu "
        "tiền điều kiện và em ghi nhận trung thực. Em sẵn sàng chạy trực tiếp case nào cô muốn xem ạ.\n\n"
    )

    a("## Phần F — Mẹo học thuộc\n\n")
    a("1. Thuộc **Phần A** và **Phần E** nguyên văn (chỉ ~50 giây).\n")
    a("2. Thuộc **6 câu ở Phần B** — đây là phần cô hỏi nhiều nhất.\n")
    a("3. Với 111 case, chỉ cần thuộc dạng **Nói ngắn**; theo công thức: *mã case → mục tiêu → API → luồng Controller-Service-Repository → kết quả*.\n")
    a("4. Chọn 3 case demo sâu: một happy-path (`IT_AUTH_05`), một negative (`IT_AUTH_06`), một N/A (ví dụ case cần enroll khóa học).\n")
    a("5. Khi bí: quay lại công thức *gọi API gì → đi qua tầng nào → mong đợi status nào*.\n\n")
    a(f"_Tổng số test case trong tài liệu: {total}._\n")

    return "".join(parts)


def font(run, size=11, bold=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def md_to_docx(md_text: str, path: Path) -> None:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.9)
        s.right_margin = Cm(1.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("KỊCH BẢN THUYẾT TRÌNH INTEGRATION TEST\nTHEO TỪNG TEST CASE\nEnglishLab – SEP490_G23")
    font(r, 15, True)

    in_code = False
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            para = doc.add_paragraph()
            font(para.add_run(line if line else " "), 9, name="Consolas")
            para.paragraph_format.space_after = Pt(0)
            continue
        if not line.strip() or set(line.strip()) == {"-"}:
            continue
        clean = line.replace("**", "").replace("`", "")
        if clean.startswith("### "):
            hp = doc.add_heading(clean[4:], level=3)
            for run in hp.runs:
                font(run, 12, True)
        elif clean.startswith("## "):
            hp = doc.add_heading(clean[3:], level=2)
            for run in hp.runs:
                font(run, 13, True)
        elif clean.startswith("# "):
            hp = doc.add_heading(clean[2:], level=1)
            for run in hp.runs:
                font(run, 15, True)
        elif clean.startswith("> "):
            para = doc.add_paragraph()
            run = para.add_run(clean[2:])
            font(run, 11, True)
            para.paragraph_format.left_indent = Cm(0.6)
        elif clean.startswith("- ") or clean.startswith("   - "):
            para = doc.add_paragraph(style="List Bullet")
            font(para.add_run(clean.lstrip(" -")), 11)
        elif re.match(r"^\d+\.\s", clean):
            para = doc.add_paragraph(style="List Number")
            font(para.add_run(re.sub(r"^\d+\.\s", "", clean)), 11)
        elif clean.startswith("|"):
            para = doc.add_paragraph()
            font(para.add_run(clean.strip("|").replace("|", "  —  ")), 10)
        else:
            para = doc.add_paragraph()
            font(para.add_run(clean), 11)

    doc.save(path)


def main() -> None:
    results = load_results()
    md = build_md(results)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    md_path = OUT_DIR / MD_NAME
    md_path.write_text(md, encoding="utf-8")
    shutil.copy2(md_path, HERE / MD_NAME)

    docx_path = OUT_DIR / DOCX_NAME
    md_to_docx(md, docx_path)
    shutil.copy2(docx_path, HERE / DOCX_NAME)

    print("MD  ", md_path)
    print("DOCX", docx_path)
    print("cases with results:", len(results))


if __name__ == "__main__":
    main()
