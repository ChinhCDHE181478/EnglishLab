# -*- coding: utf-8 -*-
"""Generate Vietnamese Word guide for Integration Code Tests per module."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT = OUT_DIR / "Huong_dan_CODE_TEST_Integration_Test_EnglishLab.docx"


def font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        font(run, 16 if level == 1 else 13 if level == 2 else 12, True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size, bold)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15


def bullets(doc, items):
    for t in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(t)
        font(run, 11)


def nums(doc, items):
    for t in items:
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(t)
        font(run, 11)


def code_block(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, 9)
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(8)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        font(run, 10, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            font(run, 9)
    doc.add_paragraph()


MODULES = [
    ("AUTH", "IT - Auth", "AuthIT", "fu.sap490.g23.backend.it.AuthIT",
     "Register/Login/Verify/Forgot/Reset",
     [
         ("IT_AUTH_01", "register_createsUser", "POST /api/auth/register → status 200/201"),
         ("IT_AUTH_02", "register_duplicateEmail_rejected", "cùng email → 4xx"),
         ("IT_AUTH_03", "verifyEmail_withValidOtp", "OTP từ DB/auth_tokens → 200"),
         ("IT_AUTH_04", "verifyEmail_invalidOtp_rejected", "OTP sai → 4xx"),
         ("IT_AUTH_05", "login_returnsJwt_andMeOk", "login + GET /api/user/me 200"),
         ("IT_AUTH_06", "login_wrongPassword_rejected", "password sai → 401"),
         ("IT_AUTH_07", "me_withoutToken_unauthorized", "không Bearer → 401/403"),
         ("IT_AUTH_08", "forgotPassword_accepted", "POST forgot-password → 200"),
         ("IT_AUTH_09", "resetPassword_withValidCode", "code từ DB → 200"),
         ("IT_AUTH_10", "resetPassword_invalidCode_rejected", "code sai → 4xx"),
     ]),
    ("USER", "IT - User", "UserIT", "fu.sap490.g23.backend.it.UserIT",
     "Profile / password / avatar",
     [
         ("IT_USER_01", "getMe_ok", "GET /api/user/me"),
         ("IT_USER_02", "updateMe_ok", "PUT /api/user/me"),
         ("IT_USER_03", "changePassword_wrongCurrent_rejected", "PUT password sai current"),
         ("IT_USER_04", "uploadAvatar_okOrSkip", "POST multipart avatar"),
         ("IT_USER_05", "updateMe_unauthorized", "không token → 401/403"),
     ]),
    ("NOTIF", "IT - Notif", "NotificationIT", "fu.sap490.g23.backend.it.NotificationIT",
     "Preferences + student notifications",
     [
         ("IT_NOTIF_01", "getPreferences", "GET notification-preferences"),
         ("IT_NOTIF_02", "updatePreferences", "PUT preferences"),
         ("IT_NOTIF_03", "updatePreferences_invalidBody", "PUT {} → 400"),
         ("IT_NOTIF_04", "listNotifications", "GET /api/student/notifications"),
         ("IT_NOTIF_05", "unreadCount", "GET unread-count"),
     ]),
    ("COMMERCE", "IT - Commerce", "CommerceIT", "fu.sap490.g23.backend.it.CommerceIT",
     "Cart / wishlist",
     [
         ("IT_COMMERCE_01", "addToCart_andGetCart", "POST cart/{id} + GET cart"),
         ("IT_COMMERCE_02", "wishlist_moveToCart", "wishlist + move-to-cart"),
         ("IT_COMMERCE_03", "clearCart", "DELETE cart"),
         ("IT_COMMERCE_04", "addToCart_again", "add lại sau clear"),
     ]),
    ("PAYMENT", "IT - Payment", "PaymentIT", "fu.sap490.g23.backend.it.PaymentIT",
     "Quote / PayOS / webhook / orders",
     [
         ("IT_PAYMENT_01", "createPayosLink", "POST /payos/link {courseIds}"),
         ("IT_PAYMENT_02", "quotePayment", "POST /quote {courseIds}"),
         ("IT_PAYMENT_03", "webhook_unsigned_rejected", "POST /api/payos/webhook → 4xx"),
         ("IT_PAYMENT_04", "managerListOrders", "GET manager payments orders"),
         ("IT_PAYMENT_05", "managerListOrders_alias", "tương tự / alias CM"),
     ]),
    ("COURSE", "IT - Course", "OnlineCourseIT", "fu.sap490.g23.backend.it.OnlineCourseIT",
     "Public catalog + learner content",
     [
         ("IT_COURSE_01", "publicCatalog", "GET /api/online-courses"),
         ("IT_COURSE_02", "publicDetail", "GET /api/online-courses/{slug}"),
         ("IT_COURSE_03", "learnerContent_whenEnrolled", "GET content (cần enroll)"),
         ("IT_COURSE_04", "updateProgress_whenEnrolled", "PATCH progress"),
         ("IT_COURSE_05", "rating_whenAllowed", "POST rating"),
         ("IT_COURSE_06", "learnerContent_again", "tương tự 03"),
     ]),
    ("DISCUSS", "IT - Discuss", "DiscussionIT", "fu.sap490.g23.backend.it.DiscussionIT",
     "Discussion + report + CM moderation",
     [
         ("IT_DISCUSS_01", "createDiscussion", "POST student discussions"),
         ("IT_DISCUSS_02", "listDiscussions", "GET /api/online-courses/{id}/discussions"),
         ("IT_DISCUSS_03", "reportThread", "POST .../reports"),
         ("IT_DISCUSS_04", "reportReply", "report reply nếu có"),
         ("IT_DISCUSS_05", "cmListReports", "GET /api/content-manager/discussion-reports"),
     ]),
    ("CONTENT", "IT - Content", "ContentManagerCourseIT", "fu.sap490.g23.backend.it.ContentManagerCourseIT",
     "CM online courses",
     [
         ("IT_CONTENT_01", "listCmCourses", "GET /api/content-manager/online-courses"),
         ("IT_CONTENT_02", "publishFlow_surface", "list/publish endpoint reachable"),
         ("IT_CONTENT_03", "versionSurface", "versions API"),
         ("IT_CONTENT_04", "listAgain", "list lại"),
     ]),
    ("PACKAGE", "IT - Package", "PackageIT", "fu.sap490.g23.backend.it.PackageIT",
     "CM packages",
     [
         ("IT_PACKAGE_01", "listPackages", "GET /api/content-manager/packages"),
         ("IT_PACKAGE_02", "bundleItemsSurface", "PUT bundle-items nếu có id"),
         ("IT_PACKAGE_03", "listAgain", "list packages"),
     ]),
    ("CURRICULUM", "IT - Curriculum", "CurriculumIT", "fu.sap490.g23.backend.it.CurriculumIT",
     "Curriculum / bank / rubric",
     [
         ("IT_CURRICULUM_01", "listPrograms", "GET curriculum-programs"),
         ("IT_CURRICULUM_02", "exerciseOrAssessmentBank", "GET exercise-bank/assessment-bank"),
         ("IT_CURRICULUM_03", "learningPaths", "GET learning-paths"),
         ("IT_CURRICULUM_04", "rubrics", "GET /api/content-manager/rubrics"),
         ("IT_CURRICULUM_05", "listProgramsAgain", "list programs"),
     ]),
    ("ENROLLREQ", "IT - EnrollReq", "EnrollmentRequestIT", "fu.sap490.g23.backend.it.EnrollmentRequestIT",
     "Student submit + staff list",
     [
         ("IT_ENROLLREQ_01", "studentSubmitAndListMine", "POST enroll + GET /my"),
         ("IT_ENROLLREQ_02", "staffList", "GET /api/staff/enrollment-requests"),
         ("IT_ENROLLREQ_03", "staffListAgain", "staff list"),
         ("IT_ENROLLREQ_04", "studentSubmitAgain", "submit (có thể conflict)"),
         ("IT_ENROLLREQ_05", "staffProcessSurface", "staff list/process"),
     ]),
    ("CLASS", "IT - Classroom", "TrainingManagerClassroomIT", "fu.sap490.g23.backend.it.TrainingManagerClassroomIT",
     "TM classrooms / registrations / waitlist",
     [
         ("IT_CLASS_01", "publicOfferings", "GET /api/classroom-offerings"),
         ("IT_CLASS_02", "tmList", "GET /api/training-manager/classrooms"),
         ("IT_CLASS_03", "tmDetail", "GET .../classrooms/{id}"),
         ("IT_CLASS_04", "registrations", "GET .../registrations"),
         ("IT_CLASS_05", "reorderWaitlist", "PUT .../waitlist/order"),
         ("IT_CLASS_06", "registrationsAgain", "registrations"),
         ("IT_CLASS_07", "assignTeacherPrecheck", "detail trước assign"),
         ("IT_CLASS_08", "tmListAgain", "tm list"),
     ]),
    ("LEARNERCLS", "IT - LearnerCls", "StudentClassroomIT", "fu.sap490.g23.backend.it.StudentClassroomIT",
     "Learner classroom access",
     [
         ("IT_LEARNERCLS_01", "myClassrooms", "GET my-classrooms"),
         ("IT_LEARNERCLS_02", "sessions", "GET .../sessions"),
         ("IT_LEARNERCLS_03", "homework", "GET .../homework"),
         ("IT_LEARNERCLS_04", "materials", "GET .../materials"),
         ("IT_LEARNERCLS_05", "homeworkAgain", "homework"),
         ("IT_LEARNERCLS_06", "gradebookMe", "GET .../gradebook/me"),
     ]),
    ("TEACH", "IT - Teacher", "TeacherClassroomIT", "fu.sap490.g23.backend.it.TeacherClassroomIT",
     "Teacher assigned ops",
     [
         ("IT_TEACH_01", "assignedClassrooms", "GET /api/teacher/classrooms/assigned"),
         ("IT_TEACH_02", "homework", "GET .../homework"),
         ("IT_TEACH_03", "sessionAttendance", "GET sessions/{id}/attendance"),
         ("IT_TEACH_04", "gradebook", "GET .../gradebook"),
         ("IT_TEACH_05", "myChangeRequests", "GET requests/mine"),
         ("IT_TEACH_06", "assignedAgain", "assigned"),
     ]),
    ("QUIZ", "IT - Quiz", "ClassroomQuizIT", "fu.sap490.g23.backend.it.ClassroomQuizIT",
     "Teacher/student quiz",
     [
         ("IT_QUIZ_01", "teacherListQuizzes", "GET teacher .../quizzes"),
         ("IT_QUIZ_02", "teacherListAgain", "list quizzes"),
         ("IT_QUIZ_03", "studentListQuizzes", "GET /api/student/classrooms/quizzes"),
         ("IT_QUIZ_04", "deleteQuiz_optional", "DELETE quiz (cẩn thận data)"),
     ]),
    ("ASSESS", "IT - Assess", "AssessmentIT", "fu.sap490.g23.backend.it.AssessmentIT",
     "Placement / assessments / mock",
     [
         ("IT_ASSESS_01", "placementCurrent", "GET placement-tests/current"),
         ("IT_ASSESS_02", "placementSubmit_invalidOrOk", "POST submit"),
         ("IT_ASSESS_03", "courseAssessments", "GET online-courses/{id}/assessments"),
         ("IT_ASSESS_04", "mockTests", "GET /api/student/mock-tests"),
         ("IT_ASSESS_05", "courseAssessmentsAgain", "assessments"),
         ("IT_ASSESS_06", "mockTestsAgain", "mock tests"),
     ]),
    ("SUPPORT", "IT - Support", "SupportTicketIT", "fu.sap490.g23.backend.it.SupportTicketIT",
     "Support tickets",
     [
         ("IT_SUPPORT_01", "createTicket", "POST support-tickets (subject/category/message)"),
         ("IT_SUPPORT_02", "listMine", "GET student support-tickets"),
         ("IT_SUPPORT_03", "managerList", "GET manager support-tickets"),
         ("IT_SUPPORT_04", "createInvalid_rejected", "POST {} → 400"),
     ]),
    ("ADMIN", "IT - Admin", "AdminIT", "fu.sap490.g23.backend.it.AdminIT",
     "Admin users/audit/config",
     [
         ("IT_ADMIN_01", "listUsers", "GET /api/admin/users"),
         ("IT_ADMIN_02", "usersSurface", "users API"),
         ("IT_ADMIN_03", "auditLogs", "GET /api/admin/audit-logs"),
         ("IT_ADMIN_04", "systemConfig", "GET /api/admin/system/config"),
     ]),
    ("LARK", "IT - Lark", "LarkIT", "fu.sap490.g23.backend.it.LarkIT",
     "Lark webhook / sync",
     [
         ("IT_LARK_01", "webhookChallenge", "POST /api/lark/events"),
         ("IT_LARK_02", "webhookEvent", "events payload"),
         ("IT_LARK_03", "syncRecording", "POST recordings/.../sync-lark"),
     ]),
    ("INFRA", "IT - Infra", "InfrastructureIT", "fu.sap490.g23.backend.it.InfrastructureIT",
     "Campuses/rooms/templates",
     [
         ("IT_INFRA_01", "listCampuses", "GET .../infrastructure/campuses"),
         ("IT_INFRA_02", "listRooms", "GET .../infrastructure/rooms"),
         ("IT_INFRA_03", "listTemplates", "GET .../infrastructure/session-templates"),
     ]),
    ("REPORT", "IT - Report", "ReportIT", "fu.sap490.g23.backend.it.ReportIT",
     "Dashboard / revenue",
     [
         ("IT_REPORT_01", "tmDashboard", "GET /api/training-manager/dashboard"),
         ("IT_REPORT_02", "revenueAnalytics", "GET content-manager/revenue/analytics"),
     ]),
    ("PROPOSAL", "IT - Proposal", "ClassroomProposalIT", "fu.sap490.g23.backend.it.ClassroomProposalIT",
     "Staff proposals",
     [
         ("IT_PROPOSAL_01", "listProposals", "GET /api/staff/classroom-proposals"),
         ("IT_PROPOSAL_02", "listAgain", "list"),
         ("IT_PROPOSAL_03", "listAgain2", "list"),
     ]),
    ("DISPUTE", "IT - Dispute", "AttendanceDisputeIT", "fu.sap490.g23.backend.it.AttendanceDisputeIT",
     "Attendance disputes",
     [
         ("IT_DISPUTE_01", "studentListDisputes", "GET /api/student/attendance/disputes"),
         ("IT_DISPUTE_02", "teacherPending", "GET /api/teacher/attendance-disputes/pending"),
         ("IT_DISPUTE_03", "teacherPendingAgain", "pending"),
     ]),
    ("NOTES", "IT - Notes", "LearningNotesIT", "fu.sap490.g23.backend.it.LearningNotesIT",
     "Learner notes",
     [
         ("IT_NOTES_01", "listNotes", "GET /api/student/learning/notes"),
         ("IT_NOTES_02", "listNotesAgain", "list notes"),
     ]),
]


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HƯỚNG DẪN CODE TEST (JUnit + MockMvc)\nINTEGRATION TEST – ENGLISHLAB (SEP490_G23)")
    font(r, 18, True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Viết và chạy test Java tự động cho từng module, map với các mã IT_* trên Excel")
    font(r, 12)

    p(doc, "Tester: phongdx", True)
    p(doc, "Vị trí code test gợi ý: backend/src/test/java/fu/sap490/g23/backend/it/")
    p(doc, "Chạy: cd backend → .\\mvnw.cmd test   hoặc Run Test trong IntelliJ/VS Code")

    h(doc, "A. Code test là gì?", 1)
    p(doc, "Code test Integration = viết class Java dùng JUnit 5 + MockMvc (hoặc @SpringBootTest) để gọi API/service giống Postman, nhưng máy tự chạy và tự assert.")
    table(doc, ["Cách", "Ai chạy", "Khi nào dùng"], [
        ["Postman", "Bạn bấm tay", "Tự test demo với cô, ghi Excel"],
        ["Code test", "Maven/IDE", "Tự động hóa, chạy lại nhanh, CI"],
    ])

    h(doc, "B. Hai kiểu viết trong project này", 1)
    p(doc, "1) Mockito + @ExtendWith(MockitoExtension) — test Service với mock Repository (nhanh, không cần DB). Project đã có nhiều file kiểu này (LoginTest, WaitlistTest…).", True)
    p(doc, "2) @SpringBootTest + @AutoConfigureMockMvc — Integration thật hơn: lên context Spring, gọi HTTP qua MockMvc, có thể chạm DB. Phù hợp map Excel IT_*.", True)
    p(doc, "Tài liệu này ưu tiên kiểu (2) cho Integration Test. Nếu DB test khó setup, có thể bắt đầu bằng (1) rồi nâng dần.")

    h(doc, "C. Cấu trúc thư mục đề xuất", 1)
    code_block(doc, """backend/src/test/java/fu/sap490/g23/backend/it/
  ItSupport.java              // helper login lấy JWT
  AuthIT.java
  UserIT.java
  NotificationIT.java
  CommerceIT.java
  PaymentIT.java
  OnlineCourseIT.java
  ... (mỗi module 1 class)
""")

    h(doc, "D. Template lớp hỗ trợ (copy dùng chung)", 1)
    p(doc, "Tạo file ItSupport.java:")
    code_block(doc, """package fu.sap490.g23.backend.it;

import com.fasterxml.jackson.databind.JsonNode;
import com.fasterxml.jackson.databind.ObjectMapper;
import org.springframework.http.MediaType;
import org.springframework.test.web.servlet.MockMvc;
import org.springframework.test.web.servlet.MvcResult;

import static org.springframework.test.web.servlet.request.MockMvcRequestBuilders.post;
import static org.springframework.test.web.servlet.result.MockMvcResultMatchers.status;

public final class ItSupport {
    private ItSupport() {}
    private static final ObjectMapper MAPPER = new ObjectMapper();

    public static String login(MockMvc mvc, String email, String password) throws Exception {
        String body = \"\"\"
            {"email":"%s","password":"%s"}
            \"\"\".formatted(email, password);
        MvcResult result = mvc.perform(post("/api/auth/login")
                .contentType(MediaType.APPLICATION_JSON)
                .content(body))
            .andExpect(status().isOk())
            .andReturn();
        JsonNode json = MAPPER.readTree(result.getResponse().getContentAsString());
        return json.get("accessToken").asText();
    }
}
""")

    h(doc, "E. Template 1 class Integration (Auth mẫu đầy đủ)", 1)
    p(doc, "Tạo AuthIT.java — đây là mẫu để clone sang module khác:")
    code_block(doc, """package fu.sap490.g23.backend.it;

import org.junit.jupiter.api.DisplayName;
import org.junit.jupiter.api.Test;
import org.springframework.beans.factory.annotation.Autowired;
import org.springframework.boot.test.autoconfigure.web.servlet.AutoConfigureMockMvc;
import org.springframework.boot.test.context.SpringBootTest;
import org.springframework.http.MediaType;
import org.springframework.test.web.servlet.MockMvc;

import java.util.UUID;

import static org.springframework.test.web.servlet.request.MockMvcRequestBuilders.*;
import static org.springframework.test.web.servlet.result.MockMvcResultMatchers.*;

@SpringBootTest
@AutoConfigureMockMvc
class AuthIT {

    @Autowired MockMvc mockMvc;

    private static final String LEARNER = "0386852628z@gmail.com";
    private static final String PASS = "Password123!";

    @Test
    @DisplayName("IT_AUTH_01 register")
    void register_createsUser() throws Exception {
        String email = "it." + UUID.randomUUID() + "@englishlab-it.test";
        String body = \"\"\"
            {"email":"%s","password":"%s","fullName":"IT User"}
            \"\"\".formatted(email, PASS);
        mockMvc.perform(post("/api/auth/register")
                .contentType(MediaType.APPLICATION_JSON)
                .content(body))
            .andExpect(status().is2xxSuccessful());
    }

    @Test
    @DisplayName("IT_AUTH_05 login + me")
    void login_returnsJwt_andMeOk() throws Exception {
        String token = ItSupport.login(mockMvc, LEARNER, PASS);
        mockMvc.perform(get("/api/user/me")
                .header("Authorization", "Bearer " + token))
            .andExpect(status().isOk())
            .andExpect(jsonPath("$.email").value(LEARNER));
    }

    @Test
    @DisplayName("IT_AUTH_06 wrong password")
    void login_wrongPassword_rejected() throws Exception {
        mockMvc.perform(post("/api/auth/login")
                .contentType(MediaType.APPLICATION_JSON)
                .content(\"\"\"{"email":"%s","password":"Wrong!"} \"\"\".formatted(LEARNER)))
            .andExpect(status().is4xxClientError());
    }

    @Test
    @DisplayName("IT_AUTH_07 me without token")
    void me_withoutToken_unauthorized() throws Exception {
        mockMvc.perform(get("/api/user/me"))
            .andExpect(status().is4xxClientError());
    }
}
""")

    h(doc, "F. Cách chạy code test", 1)
    nums(doc, [
        "Mở terminal tại thư mục backend.",
        "Chạy tất cả: .\\mvnw.cmd test",
        "Chạy 1 class: .\\mvnw.cmd -Dtest=AuthIT test",
        "Chạy 1 method: .\\mvnw.cmd -Dtest=AuthIT#login_returnsJwt_andMeOk test",
        "Hoặc trong IDE: click biểu tượng Run bên cạnh @Test / class.",
    ])
    p(doc, "Kết quả xanh = Passed; đỏ = Failed. Ghi vào Excel sheet tương ứng (Round 1/2/3) + Note = tên method test.")

    h(doc, "G. Quy tắc map Excel ↔ Code", 1)
    bullets(doc, [
        "Mỗi IT_* = một @Test (hoặc một method rõ ràng).",
        "@DisplayName(\"IT_AUTH_05 ...\") để đối chiếu Excel.",
        "Passed Excel khi test xanh và đúng expected của case.",
        "Failed Excel khi test đỏ (assert sai / 500).",
        "N/A Excel khi chưa viết được do thiếu seed/precondition — ghi Note: chưa automate.",
    ])

    h(doc, "H. Account demo dùng trong code", 1)
    table(doc, ["Role", "Email", "Password"], [
        ["LEARNER", "0386852628z@gmail.com", "Password123!"],
        ["TEACHER", "classroom.teacher1@englishlab.vn", "Password123!"],
        ["TM", "training.manager@englishlab.vn", "Password123!"],
        ["STAFF", "staff@englishlab.vn", "Password123!"],
        ["MANAGER", "classroom.manager@englishlab.vn", "Password123!"],
        ["CM", "content.manager@englishlab.vn", "Password123!"],
        ["ADMIN", "classroom.admin@englishlab.vn", "Password123!"],
    ])

    h(doc, "I. Hướng dẫn CODE TEST theo từng MODULE", 1)
    p(doc, "Với mỗi module: tạo 1 class Java, viết các @Test theo bảng. Clone từ AuthIT rồi đổi URL/token role.")

    for code, sheet, clazz, fqn, goal, cases in MODULES:
        h(doc, f"{code} — Sheet Excel: {sheet}", 2)
        p(doc, f"Class đề xuất: {clazz}.java", True)
        p(doc, f"Package đầy đủ: {fqn}")
        p(doc, f"Mục tiêu: {goal}")
        p(doc, "Các method cần viết:")
        rows = [(cid, method, assert_) for cid, method, assert_ in cases]
        table(doc, ["Mã Excel", "Tên method Java (@Test)", "Assert chính"], rows)
        # short snippet tip per module type
        if code == "TEACH":
            p(doc, "Gợi ý: token TEACHER; trước hết GET /api/teacher/classrooms/assigned lấy id lớp.")
        if code == "CLASS":
            p(doc, "Gợi ý: token TM; waitlist reorder cần ≥2 enrollment WAITLIST trong DB seed.")
        if code == "PAYMENT":
            p(doc, "Gợi ý: body bắt buộc {\"courseIds\":[id]}; webhook unsigned kỳ vọng 4xx.")
        if code == "SUPPORT":
            p(doc, "Gợi ý: body tạo ticket dùng field message (không phải body), category enum TECHNICAL.")
        if code == "ENROLLREQ":
            p(doc, "Gợi ý: list của HV là GET /api/student/course-enrollment-requests/my (không phải path gốc).")
        if code == "INFRA":
            p(doc, "Gợi ý: path đúng là /api/training-manager/infrastructure/campuses|rooms|session-templates.")
        if code == "NOTES":
            p(doc, "Gợi ý: GET /api/student/learning/notes")
        if code == "DISPUTE":
            p(doc, "Gợi ý: student /api/student/attendance/disputes ; teacher /api/teacher/attendance-disputes/pending")

    h(doc, "J. Ví dụ nhanh module khác (Teacher)", 1)
    code_block(doc, """@Test
@DisplayName("IT_TEACH_01 assigned")
void assignedClassrooms() throws Exception {
    String token = ItSupport.login(mockMvc,
        "classroom.teacher1@englishlab.vn", "Password123!");
    mockMvc.perform(get("/api/teacher/classrooms/assigned")
            .header("Authorization", "Bearer " + token))
        .andExpect(status().isOk());
}
""")

    h(doc, "K. Ví dụ negative + happy-path Support", 1)
    code_block(doc, """@Test
@DisplayName("IT_SUPPORT_01 create")
void createTicket() throws Exception {
    String token = ItSupport.login(mockMvc, LEARNER, PASS);
    String body = \"\"\"
      {"subject":"IT ticket code test","category":"TECHNICAL",
       "message":"Noi dung ticket integration code test"}
      \"\"\";
    mockMvc.perform(post("/api/student/support-tickets")
            .header("Authorization", "Bearer " + token)
            .contentType(MediaType.APPLICATION_JSON)
            .content(body))
        .andExpect(status().isOk());
}

@Test
@DisplayName("IT_SUPPORT_04 invalid")
void createInvalid_rejected() throws Exception {
    String token = ItSupport.login(mockMvc, LEARNER, PASS);
    mockMvc.perform(post("/api/student/support-tickets")
            .header("Authorization", "Bearer " + token)
            .contentType(MediaType.APPLICATION_JSON)
            .content("{}"))
        .andExpect(status().isBadRequest());
}
""")

    h(doc, "L. Lưu ý kỹ thuật khi chạy @SpringBootTest", 1)
    bullets(doc, [
        "Cần DB PostgreSQL englishlab đang chạy (giống app thật), vì project chưa dùng Testcontainers trong pom.",
        "Nếu security chặn: luôn gắn header Authorization Bearer.",
        "Tránh xóa dữ liệu demo trong test (quiz delete, user delete) trừ khi tạo data riêng rồi dọn.",
        "OTP verify/reset: đọc auth_tokens bằng JdbcTemplate trong test hoặc seed cố định.",
        "Test flaky (rate-limit OTP 15s): dùng @Disabled hoặc Thread.sleep có kiểm soát — ghi Note trên Excel.",
        "Project đã có test kiểu Mockito ở service/*Test.java — có thể dùng song song cho phần logic thuần.",
    ])

    h(doc, "M. Checklist nộp / thuyết trình code test", 1)
    nums(doc, [
        "Có package it/ với ít nhất AuthIT + 1 module nghiệp vụ (Classroom/Teacher).",
        "Mỗi @Test có @DisplayName chứa mã IT_*.",
        "Chạy .\\mvnw.cmd -Dtest=AuthIT test ra xanh (hoặc giải thích fail thật).",
        "Excel Round được cập nhật theo kết quả code test / Postman.",
        "Biết phân biệt: Postman = thủ công; Code test = tự động.",
    ])

    h(doc, "N. Câu nói với cô", 1)
    p(doc, "Ngoài Postman, em viết Integration Test bằng JUnit 5 + MockMvc (@SpringBootTest). Mỗi test method map 1 mã IT_* trên Excel. Em chạy bằng Maven, method xanh thì ghi Passed trên sheet tương ứng.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    PROJ.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT, PROJ / OUT.name)
    print("WORD", OUT)


if __name__ == "__main__":
    build()
