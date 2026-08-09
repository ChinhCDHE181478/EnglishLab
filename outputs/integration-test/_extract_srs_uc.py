# -*- coding: utf-8 -*-
from pathlib import Path
import re
from docx import Document

srs = Path(r"C:\Users\phong\Downloads\TÀI LIỆU ĐỒ ÁN ENGLISHLAB SRS\REPORT\SEP490_G23_Report3_Software Requirement Specification_New.docx")
d = Document(str(srs))
out = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test\_srs_uc_headings.txt")

lines = []
for i, p in enumerate(d.paragraphs):
    t = (p.text or "").strip()
    if not t:
        continue
    style = (p.style.name or "") if p.style else ""
    if style.startswith("Heading") or re.match(r"^2\.\d+", t) or re.match(r"^UC\b", t, re.I) or "Use Case" in t or "use case" in t.lower():
        lines.append(f"{i}|{style}|{t[:200]}")

out.write_text("\n".join(lines), encoding="utf-8")
print("paragraphs", len(d.paragraphs), "matched", len(lines), "->", out)
for x in lines[:100]:
    print(x)
