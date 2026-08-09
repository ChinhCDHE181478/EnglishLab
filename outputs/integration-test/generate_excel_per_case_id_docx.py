# -*- coding: utf-8 -*-
"""
Word: Hướng dẫn Integration Test với Excel — CHI TIẾT TỪNG Test Case ID.
"""
from __future__ import annotations

import importlib.util
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
from full_modules import MODULES  # noqa: E402

# load detailed machine steps from existing generator
_spec = importlib.util.spec_from_file_location(
    "it_beginner_steps", HERE / "generate_testcases_beginner_steps_word.py"
)
_beg = importlib.util.module_from_spec(_spec)
assert _spec.loader is not None
_spec.loader.exec_module(_beg)
get_steps = _beg.get_steps

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = HERE
OUT_NAME = "Huong_dan_Excel_IT_CHI_TIET_tung_Test_Case_ID.docx"
EXCEL = "SEP490_G23_Report5.2_Integration Test_COMPLETED_HONEST_FORMATTED.xlsx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True)


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.1


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    font(run, size=10)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    font(run, size=10)


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=9, name="Consolas")
    para.paragraph_format.left_indent = Cm(0.35)
    para.paragraph_format.space_after = Pt(4)


def field(doc, label, value):
    para = doc.add_paragraph()
    r1 = para.add_run(label + " ")
    font(r1, size=10, bold=True)
    r2 = para.add_run(str(value))
    font(r2, size=10)


ACCOUNTS = {
    "Auth": "Public / Learner 0386852628z@gmail.com",
    "User": "LEARNER 0386852628z@gmail.com",
    "Notif": "LEARNER",
    "Commerce": "LEARNER",
    "Payment": "LEARNER + MANAGER classroom.manager@englishlab.vn",
    "Course": "Public + LEARNER",
    "Discuss": "LEARNER + CM content.manager@englishlab.vn",
    "Content": "CM content.manager@englishlab.vn",
    "Package": "CM",
    "Curriculum": "CM",
    "EnrollReq": "LEARNER + STAFF staff@englishlab.vn",
    "Classroom": "TM training.manager@englishlab.vn (+ public)",
    "LearnerCls": "LEARNER (nên đã gán lớp)",
    "Teacher": "TEACHER classroom.teacher1@englishlab.vn",
    "Quiz": "TEACHER + LEARNER",
    "Assess": "LEARNER",
    "Support": "LEARNER + MANAGER",
    "Admin": "ADMIN classroom.admin@englishlab.vn",
    "Lark": "Webhook + TM",
    "Infra": "TM",
    "Report": "TM + CM",
    "Proposal": "STAFF",
    "Dispute": "LEARNER + TEACHER",
    "Notes": "LEARNER",
}


def sheet_account(sheet: str) -> str:
    key = sheet.replace("IT - ", "")
    return ACCOUNTS.get(key, "Xem Pre-conditions trên Excel; password demo Password123!")


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.6)
        s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(1.7)
        s.right_margin = Cm(1.7)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "HƯỚNG DẪN INTEGRATION TEST VỚI FILE EXCEL\n"
        "CHI TIẾT TỪNG TEST CASE ID\n"
        "EnglishLab – SEP490_G23"
    )
    font(r, size=15, bold=True)

    p(doc, f"File Excel: {EXCEL}", bold=True)
    p(doc, "Cách dùng: mở Excel đúng sheet → tìm đúng Test Case ID → làm theo mục “Bước trên máy” → ghi Round trên Excel.")
    p(doc, "Password demo chung: Password123!")

    h(doc, "A. Nhắc nhanh cột Excel (mọi case đều dùng)", 1)
    bullet(doc, "Test Case ID — mã cần tìm (vd IT_AUTH_01)")
    bullet(doc, "Test Case Description — mục tiêu tích hợp")
    bullet(doc, "Test Case Procedure — luồng Controller→Service→Repository")
    bullet(doc, "Expected Results — đối chiếu sau khi chạy")
    bullet(doc, "Pre-conditions — thiếu thì cân nhắc N/A")
    bullet(doc, "Round 1/2/3 + Test date + Tester — ô bạn điền")

    h(doc, "B. Chuẩn bị 1 lần trước khi làm các ID", 1)
    numbered(doc, "Bật PostgreSQL + backend cổng 8080.")
    numbered(doc, "Mở Excel bản HONEST_FORMATTED.")
    numbered(doc, "Mở Postman, Environment EnglishLab-Local, biến baseUrl=http://localhost:8080")
    numbered(doc, "Login Learner lưu {{token}}; các role khác lưu teacherToken/tmToken/… khi cần.")
    code(
        doc,
        'POST {{baseUrl}}/api/auth/login\n'
        '{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!"\n}',
    )

    h(doc, "C. Cách ghi Round sau mỗi ID", 1)
    bullet(doc, "Passed — đúng Expected (kể cả negative bị chặn đúng)")
    bullet(doc, "Failed — sai Expected / 500")
    bullet(doc, "N/A — thiếu Pre-condition (OTP, enroll, data…)")
    bullet(doc, "Tester = tên bạn · Test date = ngày chạy · Note ngắn HTTP …")

    h(doc, "D. CHI TIẾT TỪNG TEST CASE ID", 1)

    total = 0
    for mi, m in enumerate(MODULES, 1):
        sheet = m["sheet"]
        h(doc, f"MODULE {mi}. Sheet Excel: {sheet} — {m['name']}", 2)
        field(doc, "Feature / Controller:", m.get("function", ""))
        field(doc, "Tài khoản gợi ý:", sheet_account(sheet))
        field(doc, "Components:", m.get("components", ""))

        for g in m["groups"]:
            p(doc, f"Nhóm trên Excel: {g['name']}", bold=True)
            for c in g["cases"]:
                total += 1
                cid = c["id"]
                para = doc.add_paragraph()
                rr = para.add_run(f"{cid}")
                font(rr, size=12, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

                p(doc, "1) Trên file Excel — làm đúng các click sau:", bold=True)
                numbered(doc, f"Mở sheet « {sheet} » (thanh sheet dưới cùng Excel).")
                numbered(doc, f"Dùng Ctrl+F tìm: {cid}")
                numbered(doc, "Đọc cột Pre-conditions — chuẩn bị đủ điều kiện.")
                numbered(doc, "Đọc cột Test Case Description — biết mục tiêu.")
                numbered(doc, "Đọc cột Test Case Procedure — biết hàm nào tích hợp với hàm nào.")
                numbered(doc, "Đọc cột Expected Results — nhớ Status / dữ liệu cần có.")

                p(doc, "2) Mục tiêu (tóm tắt):", bold=True)
                bullet(doc, c["desc"])

                p(doc, "3) Pre-conditions (từ Excel):", bold=True)
                for line in str(c["pre"]).split("\n"):
                    if line.strip():
                        bullet(doc, line.strip())

                p(doc, "4) Procedure tích hợp hàm (đối chiếu cột Excel):", bold=True)
                for line in str(c["proc"]).split("\n"):
                    if line.strip():
                        bullet(doc, line.strip())

                p(doc, "5) Bước thực hiện trên máy (chi tiết):", bold=True)
                steps, body = get_steps(cid, c["proc"])
                for st in steps:
                    numbered(doc, st)
                if body:
                    p(doc, "Body JSON mẫu (copy/paste vào Postman):", bold=True)
                    code(doc, body)

                p(doc, "6) Expected Results (đối chiếu cột Excel):", bold=True)
                for line in str(c["exp"]).split("\n"):
                    if line.strip():
                        bullet(doc, line.strip())

                p(doc, "7) Ghi vào Excel ngay sau khi chạy:", bold=True)
                bullet(doc, f"Dòng {cid} → cột Round 1 (hoặc 2/3): Passed / Failed / N/A")
                bullet(doc, "Cột Test date cạnh Round: ngày hôm nay")
                bullet(doc, "Cột Tester: tên bạn")
                bullet(doc, "Nếu N/A/Failed: ghi Note ngắn (vd HTTP 400 — chưa enroll)")

                field(
                    doc,
                    "Checklist nhanh:",
                    f"☐ Đã tìm thấy {cid} trên {sheet}   ☐ Đã chạy   ☐ Đã ghi Round",
                )
                p(doc, "—" * 34)

    h(doc, "E. Mục lục Test Case ID (tra cứu nhanh)", 1)
    for m in MODULES:
        ids = [c["id"] for g in m["groups"] for c in g["cases"]]
        p(doc, f"{m['sheet']}: " + ", ".join(ids), size=10)

    h(doc, "F. Nói với cô khi cầm Excel + Word này", 1)
    p(
        doc,
        "Em hướng dẫn/thực hiện Integration Test theo từng Test Case ID trên Excel. "
        "Mỗi ID có Pre-condition, Procedure tích hợp hàm, Expected; em chạy rồi ghi Round. "
        "Postman chỉ là tool thực thi.",
    )
    p(doc, f"Tổng số Test Case ID trong tài liệu: {total}.", bold=True)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    shutil.copy2(out, PROJ / OUT_NAME)
    print(out)
    print("cases", total)
    return out


if __name__ == "__main__":
    build()
