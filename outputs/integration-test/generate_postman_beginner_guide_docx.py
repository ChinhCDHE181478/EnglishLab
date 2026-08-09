# -*- coding: utf-8 -*-
"""
Word guide: Postman Integration Test cho người mới bắt đầu
+ phần giải thích thuyết trình với giảng viên.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(r"C:\Users\phong\Downloads\intergration test")
PROJ = Path(r"D:\EngLishLab\EnglishLab\outputs\integration-test")
OUT_NAME = "Huong_dan_POSTMAN_cho_nguoi_moi_va_thuyet_trinh.docx"


def font(run, size=11, bold=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        font(r, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True)
    return x


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.15
    return para


def tip(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("💡 Mẹo: " + text)
    font(run, size=11, bold=False, color=RGBColor(0x1F, 0x4E, 0x79))
    return para


def warn(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("⚠ Lưu ý: " + text)
    font(run, size=11, bold=False, color=RGBColor(0xC0, 0x39, 0x2B))
    return para


def say(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("🎤 Nói với cô: " + text)
    font(run, size=11, bold=False, color=RGBColor(0x27, 0x6E, 0x49))
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    font(run, size=11)
    return para


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    font(run, size=11)
    return para


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    font(run, size=9, name="Consolas")
    para.paragraph_format.left_indent = Cm(0.4)
    para.paragraph_format.space_after = Pt(6)
    return para


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hd)
        font(r, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            font(r, size=9)
    doc.add_paragraph()


# ---------- module mini cases (beginner style) ----------
def add_case(doc, cid, name, why, clicks, body, expect, excel, talk):
    h(doc, f"{cid} – {name}", 3)
    p(doc, f"Mục đích case này: {why}")
    p(doc, "Làm từng click như sau:", bold=True)
    for i, c in enumerate(clicks, 1):
        numbered(doc, c)
    if body:
        p(doc, "Body (nếu có) — dán vào raw JSON:", bold=True)
        code(doc, body)
    p(doc, f"Kết quả đúng: {expect}", bold=True)
    p(doc, f"Ghi vào Excel: {excel}")
    say(doc, talk)


MODULES = [
    {
        "title": "MODULE 1 – Authentication (đăng ký / đăng nhập)",
        "sheet": "IT - Auth",
        "why": "Đây là cửa vào hệ thống. Không login được thì không test được hầu hết module khác.",
        "role": "Công khai (chưa cần token), trừ case /me.",
        "talk_module": "Module Auth kiểm tra chuỗi AuthController → AuthService → database: đăng ký, OTP, JWT login, quên mật khẩu.",
        "cases": [
            (
                "IT_AUTH_01",
                "Đăng ký tài khoản mới",
                "Kiểm tra hệ thống tạo được user mới.",
                [
                    "New → HTTP Request → đặt tên AUTH_01_register.",
                    "Method chọn POST (mũi tên cạnh GET).",
                    "URL: {{baseUrl}}/api/auth/register",
                    "Tab Body → chọn raw → góc phải chọn JSON.",
                    "Dán body bên dưới (đổi email nếu bị trùng).",
                    "Bấm Send. Nhìn dòng Status (màu xanh = ổn).",
                ],
                '{\n  "email": "it.reg.demo01@englishlab-it.test",\n  "password": "Password123!",\n  "fullName": "IT Register User"\n}',
                "Status 200 hoặc 201.",
                "Round = Passed nếu 2xx. Failed nếu 500. Note ghi HTTP …",
                "Case đăng ký happy-path: Controller nhận JSON, Service lưu user + OTP vào DB.",
            ),
            (
                "IT_AUTH_02",
                "Đăng ký trùng email (negative)",
                "Kiểm tra hệ thống từ chối email đã tồn tại — đây là test âm (negative).",
                [
                    "Copy request AUTH_01 → đổi tên AUTH_02_duplicate.",
                    "Đổi email thành 0386852628z@gmail.com (đã có sẵn).",
                    "Send → Status phải là 400/409, không được 200.",
                ],
                '{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!",\n  "fullName": "Dup"\n}',
                "Status 400 hoặc 409.",
                "Passed nếu bị từ chối. Failed nếu vẫn tạo được.",
                "Negative test cũng là Passed khi hệ thống chặn đúng — không phải Fail.",
            ),
            (
                "IT_AUTH_03",
                "Xác thực email bằng OTP",
                "Sau đăng ký thường phải nhập mã OTP mới kích hoạt tài khoản.",
                [
                    "Chạy AUTH_01 với email mới.",
                    "Mở PostgreSQL (pgAdmin/DBeaver) → bảng auth_tokens → xem dòng mới nhất, cột token/code.",
                    "Tạo POST {{baseUrl}}/api/auth/verify-email",
                    "Body có email + code = OTP vừa copy (field tên là code).",
                    "Send → 200 thì OK. Thử login email đó.",
                ],
                '{\n  "email": "it.reg.demo01@englishlab-it.test",\n  "code": "123456"\n}',
                "Status 200; login được sau đó.",
                "Passed nếu verify OK. N/A nếu không đọc được OTP/mail.",
                "Nếu không lấy được OTP em ghi N/A vì thiếu precondition, không phải Fail.",
            ),
            (
                "IT_AUTH_04",
                "OTP sai (negative)",
                "Mã OTP giả không được kích hoạt tài khoản.",
                [
                    "POST {{baseUrl}}/api/auth/verify-email",
                    "code cố định 000000.",
                    "Send → kỳ vọng 4xx.",
                ],
                '{\n  "email": "0386852628z@gmail.com",\n  "code": "000000"\n}',
                "Thường 400. Nếu user đã verify sẵn có thể khác — ghi Note thực tế.",
                "Passed nếu từ chối. Failed nếu 500.",
                "Case negative: Service kiểm tra token trong DB và từ chối OTP sai.",
            ),
            (
                "IT_AUTH_05",
                "Login lấy JWT + gọi /me",
                "Login trả accessToken; dùng token gọi API hồ sơ.",
                [
                    "POST {{baseUrl}}/api/auth/login với email/password Learner.",
                    "Trong Response copy accessToken (chuỗi dài).",
                    "Hoặc tab Tests dán script lưu {{token}} (xem Phần D).",
                    "Tạo GET {{baseUrl}}/api/user/me",
                    "Authorization → Bearer Token → {{token}} → Send.",
                ],
                '{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!"\n}',
                "Login 200 có accessToken; /me 200 đúng email.",
                "Passed nếu cả 2 bước OK.",
                "JWT chứng minh Security Filter + UserController tích hợp đúng.",
            ),
            (
                "IT_AUTH_06",
                "Sai mật khẩu (negative)",
                "Sai password không được cấp token.",
                [
                    "POST login với password WrongPass999!",
                    "Send → 401/400, không có accessToken.",
                ],
                '{\n  "email": "0386852628z@gmail.com",\n  "password": "WrongPass999!"\n}',
                "4xx, không token.",
                "Passed nếu không cấp token.",
                "Negative login bảo vệ tài khoản người dùng.",
            ),
            (
                "IT_AUTH_07",
                "Gọi /me không token (negative)",
                "API bảo vệ phải chặn khi thiếu Authorization.",
                [
                    "GET {{baseUrl}}/api/user/me",
                    "Xóa hết Authorization.",
                    "Send → 401/403.",
                ],
                "",
                "401 hoặc 403.",
                "Passed nếu bị chặn.",
                "Chứng minh Spring Security chặn trước khi vào Controller.",
            ),
            (
                "IT_AUTH_08",
                "Quên mật khẩu",
                "Hệ thống nhận yêu cầu gửi OTP reset.",
                [
                    "POST {{baseUrl}}/api/auth/forgot-password",
                    "Body chỉ có email.",
                    "Nếu báo chờ vài giây → đợi rồi Send lại.",
                ],
                '{\n  "email": "0386852628z@gmail.com"\n}',
                "200 hoặc thông báo generic; không 500.",
                "Passed nếu không 5xx.",
                "Forgot-password tạo auth_tokens loại PASSWORD_RESET.",
            ),
            (
                "IT_AUTH_09",
                "Đặt lại mật khẩu bằng OTP",
                "Dùng OTP reset đổi mật khẩu.",
                [
                    "Chạy AUTH_08 trước.",
                    "Đọc OTP trong auth_tokens (type liên quan PASSWORD).",
                    "POST {{baseUrl}}/api/auth/reset-password với code + newPassword.",
                    "Nên đặt newPassword = Password123! để không phá demo.",
                    "Login lại xác nhận.",
                ],
                '{\n  "email": "0386852628z@gmail.com",\n  "code": "123456",\n  "newPassword": "Password123!"\n}',
                "200; login mật khẩu mới được.",
                "Passed nếu OK. N/A nếu không lấy OTP.",
                "Luồng đầy đủ: forgot → OTP DB → reset → login.",
            ),
            (
                "IT_AUTH_10",
                "Reset OTP sai (negative)",
                "OTP sai không đổi được mật khẩu.",
                [
                    "POST reset-password với code 000000.",
                    "Send → 400.",
                ],
                '{\n  "email": "0386852628z@gmail.com",\n  "code": "000000",\n  "newPassword": "Password123!"\n}',
                "400.",
                "Passed nếu từ chối.",
                "Negative reset: hash password trong DB không đổi.",
            ),
        ],
    },
]


# Shorter modules for remaining 2-24 — still beginner but slightly denser
SHORT = [
    (
        "MODULE 2 – Hồ sơ người dùng (IT - User)",
        "IT - User",
        "LEARNER 0386852628z@gmail.com",
        "Sau khi có token, kiểm tra xem/sửa hồ sơ.",
        "Module User: UserController + UserService + bảng users.",
        [
            ("IT_USER_01", "GET hồ sơ", "GET {{baseUrl}}/api/user/me + Bearer {{token}}", "", "200", "Passed nếu 200", "Lấy đúng user theo JWT."),
            ("IT_USER_02", "Cập nhật hồ sơ", "PUT {{baseUrl}}/api/user/me", '{\n  "fullName": "Learner IT Update",\n  "phoneNumber": "0901234567",\n  "targetExam": "TOEIC",\n  "targetScore": 700,\n  "studyGoal": "IT"\n}', "200", "Passed nếu 200", "Service lưu field vào DB."),
            ("IT_USER_03", "Đổi MK sai current", "PUT {{baseUrl}}/api/user/me/password", '{\n  "currentPassword": "Wrong!",\n  "newPassword": "Password123!"\n}', "400", "Passed nếu từ chối", "Negative: không đổi hash khi sai MK cũ."),
            ("IT_USER_04", "Upload avatar", "POST {{baseUrl}}/api/user/me/avatar — Body form-data key file=ảnh", "", "200/201", "Passed/N/A theo thực tế", "Upload file thật, không JSON."),
            ("IT_USER_05", "Sửa không token", "PUT /api/user/me không Authorization", '{"fullName":"x"}', "401/403", "Passed nếu chặn", "Security chặn."),
        ],
    ),
    (
        "MODULE 3 – Thông báo (IT - Notif)",
        "IT - Notif",
        "LEARNER",
        "Preference + danh sách thông báo.",
        "Kiểm tra API preference và list notification của học viên.",
        [
            ("IT_NOTIF_01", "Lấy preference", "GET {{baseUrl}}/api/user/me/notification-preferences + Bearer", "", "200", "Passed nếu 200", "Đọc cấu hình thông báo."),
            ("IT_NOTIF_02", "Tắt/bật in-app", "PUT .../notification-preferences", '{\n  "inAppEnabled": false,\n  "emailEnabled": true,\n  "larkEnabled": false\n}', "200", "Passed nếu 200; nhớ bật lại", "Đổi rồi trả demo về trạng thái tốt."),
            ("IT_NOTIF_03", "Body thiếu", "PUT với {}", "{}", "400", "Passed nếu validation lỗi", "Negative validation."),
            ("IT_NOTIF_04/05", "List + unread", "GET /api/student/notifications và .../unread-count", "", "200", "Passed nếu 200", "List + đếm chưa đọc."),
        ],
    ),
    (
        "MODULE 4 – Giỏ hàng / Wishlist (IT - Commerce)",
        "IT - Commerce",
        "LEARNER",
        "Thêm/xóa khóa trong giỏ.",
        "Commerce: cart/wishlist gắn học viên đã login.",
        [
            ("IT_COMMERCE_01", "Thêm giỏ", "1) GET /api/online-courses lấy id\n2) DELETE /api/student/commerce/cart\n3) POST /api/student/commerce/cart/{{courseId}}\n4) GET cart", "", "200 + thấy khóa", "Passed nếu thêm được", "Luồng cart đầy đủ."),
            ("IT_COMMERCE_02", "Wishlist→cart", "POST wishlist rồi move-to-cart", "", "200 hoặc đã có", "Passed/N/A", "Có thể N/A nếu đã trong giỏ."),
            ("IT_COMMERCE_03", "Xóa giỏ", "DELETE /api/student/commerce/cart", "", "200/204", "Passed nếu xóa được", "Clear cart."),
            ("IT_COMMERCE_04", "Thêm lại", "Lặp COMMERCE_01", "", "200", "Passed nếu OK", "Retest thêm giỏ."),
        ],
    ),
    (
        "MODULE 5 – Thanh toán PayOS (IT - Payment)",
        "IT - Payment",
        "LEARNER + MANAGER",
        "Tạo quote/link thanh toán; không cần trả tiền thật.",
        "Payment tích hợp PayOS + quản lý orders.",
        [
            ("IT_PAYMENT_01", "Tạo link PayOS", "POST /api/student/payments/payos/link", '{\n  "courseIds": [1]\n}', "200 có checkoutUrl", "Passed nếu có link; Failed nếu 500", "Chỉ tạo link, không bắt buộc mở PayOS trả thật."),
            ("IT_PAYMENT_02", "Quote", "POST /api/student/payments/quote", '{\n  "courseIds": [1]\n}', "200 có totalAmount", "Passed nếu 200", "Tính giá trước khi thanh toán."),
            ("IT_PAYMENT_03", "Webhook giả", "POST /api/payos/webhook body {}", "{}", "400", "Passed nếu từ chối", "Negative: thiếu chữ ký."),
            ("IT_PAYMENT_04/05", "Manager orders", "GET /api/manager/payments/orders + managerToken", "", "200", "Passed nếu 200", "Manager xem đơn."),
        ],
    ),
    (
        "MODULE 6 – Khóa học online (IT - Course)",
        "IT - Course",
        "Public + LEARNER",
        "Catalog công khai; học cần enroll.",
        "Public API không token; student API cần mua/enroll.",
        [
            ("IT_COURSE_01", "List public", "GET /api/online-courses (không token)", "", "200", "Passed nếu 200", "Ai cũng xem catalog được."),
            ("IT_COURSE_02", "Chi tiết", "GET /api/online-courses/{{slugOrId}}", "", "200", "Passed nếu 200", "Chi tiết khóa."),
            ("IT_COURSE_03/06", "Content HV", "GET /api/student/online-courses/{{courseId}}/content + token", "", "200 enroll / 400 chưa", "Passed nếu enroll; N/A nếu chưa enroll", "N/A là hợp lệ khi thiếu precondition."),
            ("IT_COURSE_04", "Progress", "PATCH .../lessons/{{lessonId}}/progress", '{\n  "completed": true\n}', "200 nếu đủ quyền", "Passed/N/A", "Cập nhật tiến độ."),
            ("IT_COURSE_05", "Rating", "POST .../rating", '{\n  "score": 5,\n  "comment": "ok"\n}', "200/201", "Passed/N/A", "Đánh giá khóa."),
        ],
    ),
    (
        "MODULE 7 – Thảo luận (IT - Discuss)",
        "IT - Discuss",
        "LEARNER + CM",
        "Tạo/list/report thảo luận.",
        "Discussion gắn khóa; CM xem report.",
        [
            ("IT_DISCUSS_01", "Tạo thảo luận", "POST /api/student/online-courses/{{courseId}}/discussions", '{\n  "title": "IT discuss",\n  "content": "Noi dung thao luan test"\n}', "200 nếu enroll", "Passed/N/A", "Cần enroll mới tạo được."),
            ("IT_DISCUSS_02", "List", "GET /api/online-courses/{{courseId}}/discussions", "", "200", "Passed nếu 200", "List thread."),
            ("IT_DISCUSS_03/04", "Report", "POST .../discussions/{{threadId}}/reports", '{\n  "reason": "Spam IT"\n}', "200", "Passed/N/A", "Báo cáo nội dung."),
            ("IT_DISCUSS_05", "CM reports", "GET /api/content-manager/discussion-reports + cmToken", "", "200", "Passed nếu 200", "Moderation."),
        ],
    ),
    (
        "MODULE 8 – CM khóa học (IT - Content)",
        "IT - Content",
        "content.manager@englishlab.vn",
        "CM xem danh sách khóa quản lý.",
        "Content Manager quản trị khóa online.",
        [
            ("IT_CONTENT_01..04", "List khóa CM", "GET /api/content-manager/online-courses + cmToken", "", "200", "Passed nếu 200", "CM list courses."),
        ],
    ),
    (
        "MODULE 9 – Package (IT - Package)",
        "IT - Package",
        "CONTENT_MANAGER",
        "List package/bundle.",
        "Gói khóa học.",
        [
            ("IT_PACKAGE_01..03", "List packages", "GET /api/content-manager/packages + cmToken", "", "200", "Passed nếu 200", "CM packages."),
        ],
    ),
    (
        "MODULE 10 – Curriculum (IT - Curriculum)",
        "IT - Curriculum",
        "CONTENT_MANAGER",
        "Chương trình, ngân hàng, rubric.",
        "Ngân hàng nội dung phục vụ giảng dạy.",
        [
            ("IT_CURRICULUM_01/05", "Programs", "GET /api/content-manager/curriculum-programs", "", "200", "Passed nếu 200", "Curriculum programs."),
            ("IT_CURRICULUM_02", "Banks", "GET /api/content-manager/exercise-bank (hoặc assessment-bank)", "", "200", "Passed nếu 200", "Exercise bank."),
            ("IT_CURRICULUM_03", "Learning paths", "GET /api/content-manager/learning-paths", "", "200 hoặc N/A", "Ghi đúng thực tế", "Có thể chưa có API."),
            ("IT_CURRICULUM_04", "Rubrics", "GET /api/content-manager/rubrics", "", "200", "Passed nếu 200", "Rubric chấm."),
        ],
    ),
    (
        "MODULE 11 – Đăng ký tư vấn (IT - EnrollReq)",
        "IT - EnrollReq",
        "LEARNER + STAFF",
        "HV gửi form; Staff xem danh sách.",
        "Enrollment request: học viên ↔ staff.",
        [
            ("IT_ENROLLREQ_01/04", "HV tạo form", "GET /api/course-offerings → POST /api/student/course-enrollment-requests → GET .../my", '{\n  "courseOfferingId": 1,\n  "contactName": "HV IT",\n  "contactEmail": "0386852628z@gmail.com",\n  "contactPhone": "0901111222",\n  "consultationTrack": "TOEIC"\n}', "200 nếu tạo mới", "Passed/N/A nếu đã có form đang xử lý", "Happy-path tạo yêu cầu tư vấn."),
            ("IT_ENROLLREQ_02/03/05", "Staff list", "GET /api/staff/enrollment-requests + staffToken", "", "200", "Passed nếu 200", "Staff tiếp nhận."),
        ],
    ),
    (
        "MODULE 12 – Lớp học TM (IT - Classroom)",
        "IT - Classroom",
        "training.manager@englishlab.vn",
        "TM xem lớp, đăng ký, waitlist.",
        "Trái tim offline classroom: Offering + Enrollment.",
        [
            ("IT_CLASS_01", "Public offerings", "GET /api/classroom-offerings (không token)", "", "200", "Passed nếu 200", "Catalog lớp công khai."),
            ("IT_CLASS_02/08", "TM list", "GET /api/training-manager/classrooms + tmToken", "", "200; copy classroomId", "Passed nếu 200", "TM quản lý lớp."),
            ("IT_CLASS_03", "Chi tiết", "GET /api/training-manager/classrooms/{{classroomId}}", "", "200", "Passed nếu 200", "Detail lớp."),
            ("IT_CLASS_04/06", "Registrations", "GET .../registrations?classroomOfferingId={{classroomId}}", "", "200", "Passed nếu 200", "Danh sách đăng ký."),
            ("IT_CLASS_05", "Reorder waitlist", "Nếu ≥2 WAITLIST: PUT .../waitlist/order", '{\n  "enrollmentIds": [101, 102]\n}', "200 hoặc N/A", "N/A nếu <2 HV / 404 endpoint", "Có thể N/A nếu môi trường thiếu data."),
            ("IT_CLASS_07", "Xem lớp", "GET detail lớp", "", "200", "Passed nếu 200", "Xem trước khi gán GV."),
        ],
    ),
    (
        "MODULE 13 – HV trong lớp (IT - LearnerCls)",
        "IT - LearnerCls",
        "LEARNER (đã gán lớp càng tốt)",
        "Lớp của tôi, buổi học, BT, tài liệu, điểm.",
        "Góc nhìn học viên đã vào lớp.",
        [
            ("IT_LEARNERCLS_01", "My classrooms", "GET /api/student/classrooms/my-classrooms", "", "200 (có thể [])", "Passed nếu 200", "[] vẫn Passed nếu Status 200."),
            ("IT_LEARNERCLS_02", "Sessions", "GET /api/student/classrooms/{{id}}/sessions", "", "200", "Passed/N/A nếu không có lớp", "Buổi học."),
            ("IT_LEARNERCLS_03/05", "Homework", "GET .../homework", "", "200", "Passed/N/A", "Bài tập."),
            ("IT_LEARNERCLS_04", "Materials", "GET .../materials", "", "200", "Passed/N/A", "Tài liệu."),
            ("IT_LEARNERCLS_06", "Gradebook me", "GET .../gradebook/me", "", "200/204", "Passed nếu 200/204", "Điểm cá nhân."),
        ],
    ),
    (
        "MODULE 14 – Giáo viên (IT - Teacher)",
        "IT - Teacher",
        "classroom.teacher1@englishlab.vn",
        "Lớp được assign, điểm danh, gradebook.",
        "Teacher operations trên lớp được giao.",
        [
            ("IT_TEACH_01/06", "Assigned", "GET /api/teacher/classrooms/assigned", "", "200", "Passed nếu 200", "Lớp GV dạy."),
            ("IT_TEACH_02", "Homework", "GET /api/teacher/classrooms/{{id}}/homework", "", "200", "Passed nếu 200", "BT của lớp."),
            ("IT_TEACH_03", "Attendance", "GET sessions → GET .../sessions/{{sessionId}}/attendance", "", "200", "Passed/N/A", "Điểm danh theo buổi."),
            ("IT_TEACH_04", "Gradebook", "GET .../gradebook", "", "200", "Passed nếu 200", "Sổ điểm lớp."),
            ("IT_TEACH_05", "Requests mine", "GET /api/teacher/classrooms/requests/mine", "", "200", "Passed nếu 200", "Yêu cầu đổi lịch của GV."),
        ],
    ),
    (
        "MODULE 15 – Quiz lớp (IT - Quiz)",
        "IT - Quiz",
        "TEACHER + LEARNER",
        "List quiz; tránh xóa data demo.",
        "Quiz gắn classroom offering.",
        [
            ("IT_QUIZ_01/02", "GV list quiz", "GET /api/teacher/classrooms/{{offeringId}}/quizzes", "", "200", "Passed nếu 200", "GV xem quiz."),
            ("IT_QUIZ_03", "HV list quiz", "GET /api/student/classrooms/quizzes", "", "200", "Passed nếu 200", "HV xem quiz."),
            ("IT_QUIZ_04", "Xóa quiz", "DELETE /api/teacher/quizzes/{{id}} — chỉ quiz tự tạo", "", "204 hoặc bỏ qua", "N/A nếu không muốn xóa demo", "Destructive → thường N/A."),
        ],
    ),
    (
        "MODULE 16 – Assessment (IT - Assess)",
        "IT - Assess",
        "LEARNER",
        "Placement / mock / assessment khóa.",
        "Đánh giá đầu vào và bài trên khóa.",
        [
            ("IT_ASSESS_01", "Placement current", "GET /api/student/placement-tests/current", "", "200", "Passed nếu 200", "Bài placement hiện tại."),
            ("IT_ASSESS_02", "Submit rỗng", "POST .../current/submit", '{\n  "answers": []\n}', "400 hoặc hợp lý", "Passed nếu không 500", "Negative submit."),
            ("IT_ASSESS_03/05", "Assessments khóa", "GET /api/student/online-courses/{{courseId}}/assessments", "", "200/400", "Passed/N/A", "Cần enroll."),
            ("IT_ASSESS_04/06", "Mock tests", "GET /api/student/mock-tests", "", "200", "Passed nếu 200", "Danh sách mock."),
        ],
    ),
    (
        "MODULE 17 – Support ticket (IT - Support)",
        "IT - Support",
        "LEARNER + MANAGER",
        "Tạo ticket hỗ trợ.",
        "Kênh hỗ trợ kỹ thuật/tài khoản.",
        [
            ("IT_SUPPORT_01", "Tạo ticket", "POST /api/student/support-tickets", '{\n  "subject": "IT Support ticket",\n  "category": "TECHNICAL",\n  "message": "Mo ta chi tiet loi integration test"\n}', "200/201", "Passed nếu tạo được", "subject≥5, message≥10 ký tự."),
            ("IT_SUPPORT_02", "List của tôi", "GET /api/student/support-tickets", "", "200", "Passed nếu 200", "HV xem ticket."),
            ("IT_SUPPORT_03", "Manager list", "GET /api/manager/support-tickets", "", "200", "Passed nếu 200", "Manager xử lý."),
            ("IT_SUPPORT_04", "Body rỗng", "POST {}", "{}", "400", "Passed nếu validation", "Negative."),
        ],
    ),
    (
        "MODULE 18 – Admin (IT - Admin)",
        "IT - Admin",
        "classroom.admin@englishlab.vn",
        "Quản trị user / audit / config.",
        "Admin system-level APIs.",
        [
            ("IT_ADMIN_01/02", "Users", "GET /api/admin/users + adminToken", "", "200", "Passed nếu 200", "Danh sách user."),
            ("IT_ADMIN_03", "Audit", "GET /api/admin/audit-logs", "", "200", "Passed nếu 200", "Nhật ký."),
            ("IT_ADMIN_04", "Config", "GET /api/admin/system/config", "", "200", "Passed nếu 200", "Cấu hình hệ thống."),
        ],
    ),
    (
        "MODULE 19 – Lark (IT - Lark)",
        "IT - Lark",
        "Webhook + TM",
        "Tích hợp họp trực tuyến — phụ thuộc cấu hình.",
        "Nhiều case có thể N/A nếu chưa cấu hình Lark.",
        [
            ("IT_LARK_01/02", "Webhook", "POST /api/lark/events", '{\n  "type": "url_verification",\n  "challenge": "it-challenge-123"\n}', "200 hoặc 400", "Passed nếu 200; N/A nếu thiếu cấu hình", "Webhook challenge."),
            ("IT_LARK_03", "Sync session", "POST /api/training-manager/recordings/sessions/{{sessionId}}/sync-lark", "", "200 hoặc N/A", "N/A nếu session/Lark thiếu", "Đồng bộ recording."),
        ],
    ),
    (
        "MODULE 20 – Cơ sở vật chất (IT - Infra)",
        "IT - Infra",
        "TRAINING_MANAGER",
        "Campus / phòng / template buổi học.",
        "Hạ tầng phục vụ xếp lớp.",
        [
            ("IT_INFRA_01", "Campuses", "GET /api/training-manager/infrastructure/campuses", "", "200", "Passed nếu 200", "Cơ sở."),
            ("IT_INFRA_02", "Rooms", "GET /api/training-manager/infrastructure/rooms", "", "200", "Passed nếu 200", "Phòng học."),
            ("IT_INFRA_03", "Templates", "GET /api/training-manager/infrastructure/session-templates", "", "200", "Passed nếu 200", "Mẫu buổi học."),
        ],
    ),
    (
        "MODULE 21 – Báo cáo (IT - Report)",
        "IT - Report",
        "TM + CM",
        "Dashboard và doanh thu.",
        "Báo cáo vận hành / revenue.",
        [
            ("IT_REPORT_01", "TM dashboard", "GET /api/training-manager/dashboard", "", "200", "Passed nếu 200", "TM overview."),
            ("IT_REPORT_02", "Revenue", "GET /api/content-manager/revenue/analytics", "", "200", "Passed nếu 200", "CM doanh thu."),
        ],
    ),
    (
        "MODULE 22 – Đề xuất lớp (IT - Proposal)",
        "IT - Proposal",
        "staff@englishlab.vn",
        "Staff xem đề xuất mở lớp.",
        "Classroom proposal workflow.",
        [
            ("IT_PROPOSAL_01..03", "List proposals", "GET /api/staff/classroom-proposals", "", "200", "Passed nếu 200", "Staff proposals."),
        ],
    ),
    (
        "MODULE 23 – Khiếu nại điểm danh (IT - Dispute)",
        "IT - Dispute",
        "LEARNER + TEACHER",
        "HV/GV xem dispute.",
        "Attendance dispute.",
        [
            ("IT_DISPUTE_01", "HV disputes", "GET /api/student/attendance/disputes", "", "200", "Passed nếu 200", "[] vẫn OK."),
            ("IT_DISPUTE_02/03", "GV pending", "GET /api/teacher/attendance-disputes/pending", "", "200", "Passed nếu 200", "GV duyệt khiếu nại."),
        ],
    ),
    (
        "MODULE 24 – Ghi chú học (IT - Notes)",
        "IT - Notes",
        "LEARNER",
        "List learning notes.",
        "Ghi chú bài học của HV.",
        [
            ("IT_NOTES_01/02", "List notes", "GET /api/student/learning/notes", "", "200", "Passed nếu 200", "Notes API."),
        ],
    ),
]


def add_short_module(doc, title, sheet, account, goal, talk, cases):
    h(doc, title, 2)
    p(doc, f"Sheet Excel: {sheet}", bold=True)
    p(doc, f"Tài khoản: {account}")
    p(doc, f"Mục tiêu đơn giản: {goal}")
    say(doc, talk)
    p(doc, "Trước khi làm: Login đúng role → lưu token vào Environment → Authorization Bearer.", bold=True)
    for cid, name, how, body, expect, excel, talk_c in cases:
        h(doc, f"{cid} – {name}", 3)
        p(doc, "Cách làm:", bold=True)
        for line in how.split("\n"):
            bullet(doc, line)
        if body:
            p(doc, "Body mẫu:")
            code(doc, body)
        p(doc, f"Kỳ vọng: {expect}")
        p(doc, f"Ghi Excel: {excel}")
        say(doc, talk_c)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "HƯỚNG DẪN POSTMAN CHO NGƯỜI MỚI BẮT ĐẦU\n"
        "+ CÁCH GIẢI THÍCH VỚI CÔ GIÁO\n"
        "INTEGRATION TEST – ENGLISHLAB (SEP490_G23)"
    )
    font(r, size=17, bold=True)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Đọc từ đầu đến cuối là làm được · Không cần biết lập trình")
    font(r, size=12)

    p(doc, "Sinh viên / Tester: phongdx", bold=True)
    p(doc, "Công cụ: Postman · Backend: http://localhost:8080 · Ghi kết quả vào Excel Integration Test (bản HONEST).")

    # ===== PHẦN 0 =====
    h(doc, "PHẦN 0 – Đọc phần này trước (giải thích bằng tiếng Việt dễ hiểu)", 1)

    h(doc, "0.1. Integration Test là gì?", 2)
    p(
        doc,
        "Unit Test = kiểm tra 1 hàm nhỏ. Integration Test = kiểm tra nhiều tầng chạy thật với nhau: "
        "Postman gọi API → Controller → Service → Database. Em không cần viết code; em gọi API bằng Postman "
        "và xem hệ thống có trả đúng không.",
    )
    say(
        doc,
        "Em làm Integration Test theo hướng black-box qua API: kiểm tra luồng Controller–Service–Repository "
        "bằng cách gọi endpoint thật trên môi trường local, rồi ghi Passed/Failed/N/A vào Excel.",
    )

    h(doc, "0.2. Postman là gì?", 2)
    p(
        doc,
        "Postman giống “trình duyệt cho API”. Thay vì mở trang web, em gửi yêu cầu HTTP (GET/POST/PUT/DELETE) "
        "tới server và xem Status + JSON trả về.",
    )
    table(
        doc,
        ["Từ", "Nghĩa đơn giản"],
        [
            ["API", "Cửa giao tiếp của backend (đường dẫn /api/...)"],
            ["Request", "Em gửi đi (method + URL + body + header)"],
            ["Response", "Server trả lại (Status + JSON)"],
            ["Status 200", "Thành công"],
            ["Status 400/401/403", "Lỗi phía client / chưa đăng nhập / không quyền"],
            ["Status 500", "Lỗi server → thường ghi Failed"],
            ["JWT / accessToken", "Thẻ tạm sau login; gắn vào Header để chứng minh đã đăng nhập"],
            ["Bearer Token", "Cách gắn JWT: Authorization: Bearer <token>"],
            ["Environment", "Chỗ lưu biến {{baseUrl}}, {{token}} để khỏi gõ lại"],
        ],
    )

    h(doc, "0.3. Passed / Failed / N/A nghĩa là gì?", 2)
    table(
        doc,
        ["Ghi Excel", "Ý nghĩa", "Ví dụ nói với cô"],
        [
            ["Passed", "Đúng mong đợi", "Login đúng thì 200; login sai thì 401 — cả hai đều có thể Passed"],
            ["Failed", "Sai mong đợi / lỗi hệ thống", "Lẽ ra 200 mà ra 500, hoặc API không tồn tại khi đáng lẽ có"],
            ["N/A", "Chưa đủ điều kiện để kết luận", "Chưa enroll khóa nên không vào content — thiếu precondition, không phải giấu Fail"],
        ],
    )
    warn(doc, "Negative test (cố tình sai) mà hệ thống chặn đúng → ghi Passed, không ghi Failed.")

    # ===== PHẦN 1 =====
    h(doc, "PHẦN 1 – Chuẩn bị máy (checklist)", 1)
    numbered(doc, "PostgreSQL đang chạy, database tên englishlab có dữ liệu demo.")
    numbered(doc, "Backend Spring Boot chạy cổng 8080 (terminal thấy Started BackendApplication).")
    numbered(doc, "Mở trình duyệt thử: http://localhost:8080/swagger-ui.html (nếu mở được là server sống).")
    numbered(doc, "Cài Postman Desktop từ trang postman.com (bản miễn phí).")
    numbered(doc, "Mở file Excel Integration Test bản HONEST_FORMATTED để ghi kết quả.")
    tip(doc, "Nếu Postman báo Could not get response → backend chưa chạy hoặc sai cổng.")

    h(doc, "1.1. Tài khoản demo (password chung: Password123!)", 2)
    table(
        doc,
        ["Vai trò", "Email", "Dùng module nào"],
        [
            ["LEARNER (học viên)", "0386852628z@gmail.com", "Auth, User, Commerce, Course, Support…"],
            ["TEACHER", "classroom.teacher1@englishlab.vn", "Teacher, Quiz, Dispute"],
            ["TRAINING_MANAGER", "training.manager@englishlab.vn", "Classroom, Infra, Report, Lark"],
            ["STAFF", "staff@englishlab.vn", "EnrollReq, Proposal"],
            ["MANAGER", "classroom.manager@englishlab.vn", "Payment orders, Support"],
            ["CONTENT_MANAGER", "content.manager@englishlab.vn", "Content, Package, Curriculum, Discuss"],
            ["ADMIN", "classroom.admin@englishlab.vn", "Admin"],
        ],
    )

    # ===== PHẦN 2 =====
    h(doc, "PHẦN 2 – Cài Postman Environment (làm đúng 1 lần)", 1)
    numbered(doc, "Mở Postman → góc trái chọn Environments (hoặc icon mắt/envi).")
    numbered(doc, "Bấm + Create Environment → đặt tên: EnglishLab-Local.")
    numbered(doc, "Thêm các biến như bảng dưới (cột Initial value).")
    numbered(doc, "Bấm Save.")
    numbered(doc, "Góc trên bên phải Postman: chọn Environment = EnglishLab-Local (rất quan trọng!).")

    table(
        doc,
        ["VARIABLE", "INITIAL VALUE", "Giải thích"],
        [
            ["baseUrl", "http://localhost:8080", "Địa chỉ server"],
            ["token", "(để trống)", "JWT học viên — tự điền sau login"],
            ["teacherToken", "(để trống)", "JWT giáo viên"],
            ["tmToken", "(để trống)", "JWT đào tạo"],
            ["staffToken", "(để trống)", "JWT staff"],
            ["managerToken", "(để trống)", "JWT manager"],
            ["cmToken", "(để trống)", "JWT content manager"],
            ["adminToken", "(để trống)", "JWT admin"],
            ["courseId", "(điền sau)", "Id khóa học copy từ API"],
            ["classroomId", "(điền sau)", "Id lớp"],
        ],
    )
    tip(doc, "Khi gõ URL nhớ dùng {{baseUrl}}/... — hai dấu ngoặc nhọn. Nếu Postman gạch đỏ biến → chưa chọn Environment.")

    # ===== PHẦN 3 =====
    h(doc, "PHẦN 3 – Tạo request Login và tự lưu token (quan trọng nhất)", 1)
    p(doc, "3.1. Tạo Collection", bold=True)
    numbered(doc, "Collections → + Create Collection → tên EnglishLab-IT.")
    numbered(doc, "Bấm ⋯ trên collection → Add request → tên LOGIN_LEARNER.")

    p(doc, "3.2. Cấu hình Login", bold=True)
    numbered(doc, "Method = POST.")
    numbered(doc, "URL = {{baseUrl}}/api/auth/login")
    numbered(doc, "Tab Body → raw → JSON.")
    code(
        doc,
        '{\n  "email": "0386852628z@gmail.com",\n  "password": "Password123!"\n}',
    )
    numbered(doc, "Bấm Send.")
    numbered(doc, "Thấy Status 200 và trong Body có accessToken là đúng.")

    p(doc, "3.3. Script tự lưu token (làm 1 lần)", bold=True)
    numbered(doc, "Vẫn request LOGIN_LEARNER → tab Scripts (hoặc Tests, tùy bản Postman).")
    numbered(doc, "Chọn Post-response.")
    numbered(doc, "Dán đoạn sau rồi Save:")
    code(
        doc,
        "if (pm.response.code === 200) {\n"
        "  const json = pm.response.json();\n"
        '  pm.environment.set("token", json.accessToken);\n'
        "}",
    )
    numbered(doc, "Send lại lần nữa → mở Environment → thấy token đã có chuỗi dài.")

    p(doc, "3.4. Login các role khác", bold=True)
    p(doc, "Duplicate request LOGIN_LEARNER, đổi email + đổi tên biến trong script:")
    bullet(doc, 'Teacher: set("teacherToken", ...)')
    bullet(doc, 'TM: set("tmToken", ...)')
    bullet(doc, 'Staff: set("staffToken", ...)')
    bullet(doc, 'Manager: set("managerToken", ...)')
    bullet(doc, 'CM: set("cmToken", ...)')
    bullet(doc, 'Admin: set("adminToken", ...)')

    h(doc, "3.5. Gắn token vào request bảo vệ", 2)
    numbered(doc, "Mở request cần login → tab Authorization.")
    numbered(doc, "Type = Bearer Token.")
    numbered(doc, "Ô Token gõ: {{token}} (hoặc {{tmToken}} tùy role).")
    numbered(doc, "Send.")
    tip(doc, "Quên gắn token → thường ra 401. Đó là đúng nếu đang test negative; sai nếu đang test happy-path.")

    # ===== PHẦN 4 =====
    h(doc, "PHẦN 4 – Quy trình làm 1 case (in ra dán bàn cũng được)", 1)
    numbered(doc, "Mở Excel → đúng sheet (ví dụ IT - Auth).")
    numbered(doc, "Tìm mã case (ví dụ IT_AUTH_05).")
    numbered(doc, "Đọc Pre-condition → login đúng role.")
    numbered(doc, "Trong Postman tạo/ gửi request theo hướng dẫn module.")
    numbered(doc, "Nhìn Status + vài field JSON quan trọng.")
    numbered(doc, "Ghi Excel ngay:")
    bullet(doc, "Round 1: Passed / Failed / N/A")
    bullet(doc, "Test date: hôm nay")
    bullet(doc, "Tester: phongdx (hoặc tên thật)")
    bullet(doc, "Note: HTTP 200 hoặc “N/A: chưa enroll khóa”")

    # ===== PHẦN 5 demo walkthrough =====
    h(doc, "PHẦN 5 – Demo mẫu 5 phút (làm trước mặt cô cũng được)", 1)
    p(doc, "Làm live 3 case này là đủ thể hiện hiểu Integration Test:", bold=True)
    numbered(doc, "IT_AUTH_05: Login + /me → Passed (happy-path).")
    numbered(doc, "IT_AUTH_06: Sai mật khẩu → Passed (negative đúng).")
    numbered(doc, "IT_COURSE_03: Content khi chưa enroll → có thể N/A (thiếu precondition).")
    say(
        doc,
        "Em demo 1 Passed happy-path, 1 Passed negative, 1 N/A thiếu điều kiện để chứng minh em hiểu cách chấm trung thực.",
    )

    # ===== MODULE 1 detailed =====
    h(doc, "PHẦN 6 – Chi tiết MODULE 1 (Auth) — làm kỹ từng click", 1)
    m = MODULES[0]
    p(doc, f"Sheet Excel: {m['sheet']}", bold=True)
    p(doc, f"Vì sao làm module này trước: {m['why']}")
    say(doc, m["talk_module"])
    for c in m["cases"]:
        add_case(doc, *c)

    # ===== SHORT MODULES =====
    h(doc, "PHẦN 7 – Các MODULE 2 → 24 (làm theo từng bước ngắn)", 1)
    p(
        doc,
        "Mỗi module dưới đây: login đúng role → gửi request → ghi Excel. "
        "Nếu Status đúng expected thì Passed; thiếu data thì N/A và viết Note rõ.",
    )
    for item in SHORT:
        add_short_module(doc, *item)

    # ===== OTP =====
    h(doc, "PHẦN 8 – Cách lấy OTP khi không có email", 1)
    p(doc, "Mở tool DB (pgAdmin/DBeaver) → database englishlab → Query:")
    code(
        doc,
        "SELECT id, user_id, type, token, code, expires_at, created_at\n"
        "FROM auth_tokens\n"
        "ORDER BY id DESC\n"
        "LIMIT 20;",
    )
    bullet(doc, "Copy giá trị code/token mới nhất.")
    bullet(doc, "Dán vào field code của verify-email hoặc reset-password.")
    warn(doc, "Sau khi reset mật khẩu demo, hãy đặt lại Password123! để cả nhóm còn dùng chung.")

    # ===== THUYẾT TRÌNH =====
    h(doc, "PHẦN 9 – Kịch bản thuyết trình với cô (học thuộc ý, không cần thuộc từng chữ)", 1)

    h(doc, "9.1. Mở đầu 30 giây", 2)
    say(
        doc,
        "Dạ em trình bày phần tự thực hiện Integration Test bằng Postman. "
        "Em gọi API thật trên localhost, đối chiếu Expected Results trên Excel, "
        "ghi Passed/Failed/N/A theo đúng kết quả quan sát được.",
    )

    h(doc, "9.2. Giải thích phương pháp", 2)
    say(
        doc,
        "Em chọn cách black-box qua HTTP vì Integration Test ở mức API phản ánh đúng luồng "
        "Controller → Service → Repository/DB mà người dùng thật cũng đi qua (web cũng gọi API này).",
    )

    h(doc, "9.3. Demo 1 case Passed", 2)
    say(
        doc,
        "Ví dụ IT_AUTH_05: em POST /api/auth/login, nhận accessToken, rồi GET /api/user/me kèm Bearer. "
        "Cả hai bước 200 và email khớp → Passed.",
    )

    h(doc, "9.4. Demo 1 case Negative vẫn Passed", 2)
    say(
        doc,
        "Ví dụ IT_AUTH_06: em cố tình nhập sai mật khẩu, hệ thống trả 401 và không cấp token. "
        "Đúng expected nên vẫn Passed — negative test thành công.",
    )

    h(doc, "9.5. Giải thích N/A", 2)
    say(
        doc,
        "Một số case cần precondition (đã enroll khóa, waitlist ≥ 2, OTP từ mail…). "
        "Khi môi trường demo thiếu điều kiện, em ghi N/A và Note rõ lý do — "
        "đây là ghi nhận trung thực, không phải bỏ sót Fail.",
    )

    h(doc, "9.6. Câu hỏi cô có thể hỏi + câu trả lời ngắn", 2)
    table(
        doc,
        ["Cô hỏi", "Em trả lời"],
        [
            ["Vì sao dùng Postman mà không viết JUnit?", "Postman đủ để kiểm thử tích hợp API end-to-end trên môi trường chạy thật; nhanh, trực quan, khớp Excel. Code test JUnit/MockMvc là hướng bổ sung nếu cô yêu cầu."],
            ["Làm sao biết đúng Service/DB?", "Qua hành vi: Status + dữ liệu JSON; negative bị chặn; với Auth còn đối chiếu bảng auth_tokens/users khi cần OTP."],
            ["N/A có được tính hoàn thành không?", "N/A là kết quả hợp lệ khi thiếu precondition; em vẫn đã thực hiện bước gọi API và ghi nhận trung thực."],
            ["JWT là gì?", "Access token sau login; gắn Authorization Bearer để Security nhận diện user."],
            ["Failed khác N/A thế nào?", "Failed = hệ thống sai expected/lỗi 500. N/A = chưa đủ điều kiện để kết luận pass/fail đầy đủ."],
            ["Em đã test mấy module?", "Liệt kê module đã tự chạy (Auth, Classroom, …) và mở Excel Round tương ứng cho cô xem."],
        ],
    )

    h(doc, "9.7. Kết thúc 20 giây", 2)
    say(
        doc,
        "Tóm lại em đã tự thực hiện Integration Test bằng Postman theo từng mã IT trên Excel, "
        "chấm điểm trung thực gồm cả Passed negative và N/A khi thiếu precondition. Em sẵn sàng demo thêm case cô muốn xem.",
    )

    # ===== CHECKLIST =====
    h(doc, "PHẦN 10 – Checklist trước khi gặp cô", 1)
    bullet(doc, "Backend đang chạy, Postman chọn đúng Environment EnglishLab-Local.")
    bullet(doc, "Login Learner thành công, biến token có giá trị.")
    bullet(doc, "Excel đã ghi Round + Note cho ít nhất Auth + 1–2 module khác.")
    bullet(doc, "Thuộc 3 câu: IT là gì / Passed negative / N/A là gì.")
    bullet(doc, "Chuẩn bị sẵn request AUTH_05, AUTH_06 để demo live.")
    bullet(doc, "Không xóa data demo; không thanh toán PayOS thật.")

    h(doc, "PHẦN 11 – Lỗi thường gặp & cách xử lý", 1)
    table(
        doc,
        ["Hiện tượng", "Nguyên nhân", "Cách xử lý"],
        [
            ["Could not get response", "Backend tắt / sai URL", "Chạy lại Spring Boot; kiểm tra baseUrl"],
            ["401 Unauthorized", "Quên token / token hết hạn", "Login lại; gắn Bearer đúng biến"],
            ["403 Forbidden", "Sai role", "Login đúng tài khoản module"],
            ["400 Bad Request", "Sai JSON / thiếu field", "So body mẫu; nhớ field OTP là code"],
            ["404 Not Found", "Sai path", "Đối chiếu Swagger"],
            ["500", "Lỗi server", "Ghi Failed + Note; chụp response"],
            ["{{token}} không thay", "Chưa chọn Environment", "Chọn EnglishLab-Local góc phải"],
            ["No tests found (IDE)", "Nhầm sang code Java", "Phần này dùng Postman, không cần Run Test IDE"],
        ],
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / OUT_NAME
    doc.save(out)
    PROJ.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out, PROJ / OUT_NAME)
    print(f"DOCX {out}")
    return out


if __name__ == "__main__":
    build()
