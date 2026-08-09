# -*- coding: utf-8 -*-
"""Chuyển các file markdown giải thích code sang bản Word dễ đọc."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS_ROOT = Path(r"C:\Users\phong\Downloads\intergration test\giai thich code")
MD_DIR = DOCS_ROOT / "file md"
DOCX_DIR = DOCS_ROOT

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BRAND = RGBColor(0x73, 0x00, 0x14)
CODE_SHADE = "F4F4F6"
QUOTE_SHADE = "FFF4F5"

CODE_REF = re.compile(r"^```(\d+):(\d+):(\S+)$")
FENCE = re.compile(r"^```(\w*)$")
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def style_run(run, size=11, bold=False, italic=False, mono=False, color=None):
    name = CODE_FONT if mono else BODY_FONT
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), name)


def shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def add_inline(paragraph, text, size=11, base_bold=False):
    """Ghi text có hỗ trợ **đậm** và `code`."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            style_run(paragraph.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            style_run(paragraph.add_run(part[1:-1]), size=size - 1, mono=True, color=BRAND)
        else:
            style_run(paragraph.add_run(part), size=size, bold=base_bold)


def add_heading(doc, text, level):
    sizes = {1: 17, 2: 14, 3: 12.5}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level < 3 else 10)
    para.paragraph_format.space_after = Pt(5)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text.replace("**", "").replace("`", ""))
    style_run(run, size=sizes.get(level, 12), bold=True, color=BRAND if level < 3 else None)
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_block(doc, lines, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_after = Pt(1)
        style_run(cap.add_run(caption), size=8.5, italic=True, color=BRAND)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), CODE_SHADE)
    cell.text = ""
    for index, line in enumerate(lines):
        para = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        style_run(para.add_run(line if line.strip() else " "), size=8.5, mono=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        shade(cell._tc.get_or_add_tcPr(), "F2DEE1")
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], text, size=9.5, base_bold=True)
    for row in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(header)]):
            cells[idx].paragraphs[0].text = ""
            add_inline(cells[idx].paragraphs[0], text, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_quote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), QUOTE_SHADE)
    cell.paragraphs[0].text = ""
    add_inline(cell.paragraphs[0], text, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def split_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped or stripped in {"---", "***"}:
            i += 1
            continue

        ref = CODE_REF.match(stripped)
        fence = FENCE.match(stripped)
        if ref or fence:
            caption = None
            if ref:
                caption = f"{ref.group(3)} — dòng {ref.group(1)}–{ref.group(2)}"
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i].rstrip())
                i += 1
            i += 1
            add_code_block(doc, block, caption)
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].strip()) <= set("|-: "):
            rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            add_heading(doc, stripped[level:].strip(), min(level, 3))
            i += 1
            continue

        if stripped.startswith("> "):
            add_quote(doc, stripped[2:])
            i += 1
            continue

        bullet = re.match(r"^([-*+])\s+(.*)$", stripped)
        number = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if bullet or number:
            indent = len(line) - len(line.lstrip(" "))
            style = "List Number" if number else "List Bullet"
            para = doc.add_paragraph(style=style)
            para.paragraph_format.left_indent = Cm(0.75 + 0.55 * (indent // 2))
            para.paragraph_format.space_after = Pt(2)
            add_inline(para, (number or bullet).group(2), size=11)
            i += 1
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(5)
        para.paragraph_format.line_spacing = 1.15
        add_inline(para, stripped, size=11)
        i += 1

    doc.save(docx_path)


def main() -> None:
    md_files = sorted(MD_DIR.glob("*.md"))
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    done = 0
    for md in md_files:
        docx = DOCX_DIR / f"{md.stem}.docx"
        try:
            convert(md, docx)
        except PermissionError:
            print(f"BỎ QUA {docx.name} — file đang mở trong Word, hãy đóng rồi chạy lại")
            continue
        done += 1
        print(f"{docx.name}  ({round(docx.stat().st_size / 1024)} KB)")
    print(f"Đã tạo {done}/{len(md_files)} file Word trong {DOCX_DIR}")


if __name__ == "__main__":
    main()
